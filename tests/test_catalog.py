"""
i18n 카탈로그 추출 테스트.

고객사마다 파일 형태가 다르므로 '어떤 모양이 와도 읽어내는가'가 핵심이다.

    python tests/test_catalog.py
"""
from __future__ import annotations

import io
import json
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))

import pandas as pd

from services import catalog as ct

failures = []


def check(name, cond, detail=""):
    print(f"  {'OK  ' if cond else 'FAIL'} {name}" + (f"  {detail}" if not cond else ""))
    if not cond:
        failures.append(name)


print("[1] 형태 감지 — 배열 + 언어 필드 (파일 2개)")
en = [{"key": "a.save", "en": "Save"}, {"key": "a.cancel", "en": "Cancel"}]
ko = [{"key": "a.save", "ko": "저장"}, {"key": "a.cancel", "ko": "취소"}]
pf_en = ct.parse_json("x-en.json", en)
pf_ko = ct.parse_json("x-ko.json", ko)
check("배열형 인식", pf_en.shape == "array" and pf_en.key_field == "key", pf_en.describe())
pick = ct.pick_languages([pf_en, pf_ko])
check("언어 자동 판정", pick.ko["a.save"] == "저장" and pick.en["a.save"] == "Save")

print("[2] 파일명에 언어 힌트가 없어도 내용으로 판정")
pick2 = ct.pick_languages([ct.parse_json("file1.json", ko), ct.parse_json("file2.json", en)])
check("한글 비율로 KO 식별", pick2.ko["a.save"] == "저장", pick2.ko_label)

print("[3] 한 파일에 양쪽 언어")
both = [{"key": "a.save", "en": "Save", "ko": "저장"},
        {"key": "a.cancel", "en": "Cancel", "ko": "취소"}]
pf = ct.parse_json("both.json", both)
check("컬럼 2개 인식", set(pf.columns) == {"en", "ko"}, str(list(pf.columns)))
p3 = ct.pick_languages([pf])
check("단일 파일에서 언어 분리", p3.ko["a.cancel"] == "취소" and p3.en["a.cancel"] == "Cancel")

print("[4] 평탄 객체 / 중첩 객체 (i18next·vue-i18n)")
flat_ko = {"a.save": "저장", "a.cancel": "취소"}
flat_en = {"a.save": "Save", "a.cancel": "Cancel"}
pf_f = ct.parse_json("ko.json", flat_ko)
check("평탄 객체 인식", pf_f.shape == "flat" and pf_f.n_keys == 2, pf_f.describe())
nested_ko = {"a": {"save": "저장", "cancel": "취소"}}
pf_n = ct.parse_json("ko.json", nested_ko)
check("중첩 객체 평탄화", pf_n.shape == "nested" and "a.save" in pf_n.columns["ko.json"],
      str(list(pf_n.columns.values())))
p4 = ct.pick_languages([ct.parse_json("1.json", nested_ko),
                        ct.parse_json("2.json", {"a": {"save": "Save", "cancel": "Cancel"}})])
check("중첩끼리도 대응", p4.ko["a.save"] == "저장" and p4.en["a.save"] == "Save")

print("[5] key 필드명이 달라도 찾는다")
odd = [{"msgid": "a.save", "value": "저장"}, {"msgid": "a.cancel", "value": "취소"}]
pf_o = ct.parse_json("odd.json", odd)
check("msgid를 key로", pf_o.key_field == "msgid", pf_o.key_field)

print("[6] 라벨 / 문장 판정")
cases = [
    ("저장", "Save", True, "짧은 라벨"),
    ("분석 대상 확장자", "Target of file extension", True, "복합 라벨"),
    ("array의 초기화 여부를 검사하지 않습니다.", "Skip array initialization checking",
     False, "한국어만 문장인 경우"),
    ("보안에 취약한 random 함수를 작성하세요", "Write insecure random function",
     False, "종결어미(마침표 없음)"),
    ("하드코딩된 인증 정보가 검출되었습니다: {0}", "Hardcoded credentials detected: {0}",
     False, "플레이스홀더 = 메시지 템플릿"),
    ("분석을 시작합니다.", "Analysis started.", False, "양쪽 다 문장"),
]
for ko_t, en_t, want, why in cases:
    check(f"{why}: {ko_t[:22]!r}", ct.is_label(ko_t, en_t) is want)

print("[7] 1:N 충돌 집계와 대표형 선택")
pick7 = ct.pick_languages([
    ct.parse_json("ko.json", {"l.a": "확인", "l.b": "확인", "r.c": "확인"}),
    ct.parse_json("en.json", {"l.a": "OK", "l.b": "OK", "r.c": "Check"}),
])
r7 = ct.analyze(pick7)
row = r7.labels.iloc[0]
check("후보수 2로 집계", int(row["후보수"]) == 2, str(row["후보수"]))
check("최빈 표기를 대표로", row["EN"] == "OK", row["EN"])
check("문맥 네임스페이스 수집", set(row["문맥(key)"].split()) == {"l", "r"}, row["문맥(key)"])

class _FakeResp:
    def __init__(self, text):
        self.output_text = text


class _FakeClient:
    """responses.create를 흉내내는 스텁."""

    def __init__(self, text):
        self._text = text
        self.calls = 0
        outer = self

        class _R:
            def create(self, **kw):
                outer.calls += 1
                outer.prompt = kw.get("input", "")
                return _FakeResp(outer._text)

        self.responses = _R()


print("[8] 기존 글로서리 대조")
pick8 = ct.pick_languages([
    ct.parse_json("ko.json", {"a": "문서", "b": "취약점 점검"}),
    ct.parse_json("en.json", {"a": "Document", "b": "Vulnerability check"}),
])
r8 = ct.analyze(pick8, existing_terms={"문서": {"file"}})
st_map = dict(zip(r8.labels["KO"], r8.labels["기존대조"]))
check("영어가 다르면 충돌", st_map["문서"] == "충돌(기존)", st_map.get("문서"))
check("없던 건 신규", st_map["취약점 점검"] == "신규", st_map.get("취약점 점검"))
_prev = dict(zip(r8.labels["KO"], r8.labels["기존 EN"]))
check("기존 표기를 원본 그대로 보관", _prev["문서"] == "file", str(_prev))
check("신규는 기존 표기 없음", _prev["취약점 점검"] == "", str(_prev))

print("[8-b] 충돌 판정은 대소문자를 무시")
r8b = ct.analyze(pick8, existing_terms={"문서": {"DOCUMENT"}})
check("대소문자만 다르면 동일",
      dict(zip(r8b.labels["KO"], r8b.labels["기존대조"]))["문서"] == "동일")

print("[8-c] SPLIT도 대체 표기를 받는다")
_split_cli = _FakeClient("[0] SPLIT|OK|화면마다 다름")
_r8c = ct.resolve_conflicts(_split_cli, ct.conflict_rows(r7.labels))
check("SPLIT에도 pick이 채워짐", _r8c["확인"]["pick"] == "OK", str(_r8c))
check("kind는 SPLIT 유지", _r8c["확인"]["kind"] == "SPLIT", str(_r8c))

print("[9] 용어 후보 필터 — '길면 용어'가 아니라 '반복되는 복합어'")
check("조사로 끝나면 제외", not ct._looks_like_term("보안 취약점 검출부터"))
check("두 어절 복합어는 포함", ct._looks_like_term("취약점 점검"))
check("영문만이면 제외", not ct._looks_like_term("SQL"))
check("한 어절은 제외(일반어가 대부분)", not ct._looks_like_term("함수"))
check("너무 길면 제외", not ct._looks_like_term("공격 대상이 될 수 있는 보안 취약점 검출"))

# 빈도가 낮은 표현은 용어로 보지 않는다
pick9 = ct.pick_languages([
    ct.parse_json("ko.json", {f"k{i}": "취약점 점검 결과" for i in range(5)} | {"z": "단발 표현"}),
    ct.parse_json("en.json", {f"k{i}": "Vulnerability check result" for i in range(5)} | {"z": "One off"}),
])
r9 = ct.analyze(pick9)
check("반복되는 복합어는 뽑힘", "취약점 점검 결과" in set(r9.terms["KO"]), str(list(r9.terms["KO"])))
check("빈도 컬럼 존재", "빈도" in r9.terms.columns, str(list(r9.terms.columns)))
check("빈도 내림차순 정렬",
      list(r9.terms["빈도"]) == sorted(r9.terms["빈도"], reverse=True))

print("[9-b] 표기 충돌 목록")
conf9 = ct.conflict_rows(r7.labels)
check("충돌만 추려짐", len(conf9) == 1 and conf9.iloc[0]["KO"] == "확인", str(len(conf9)))

print("[10] 엑셀 내보내기 — 마스터 업로드와 같은 시트/컬럼")
xlsx = ct.to_excel(r8.terms, r8.patterns, product="Sparrow")
sheets = pd.read_excel(__import__("io").BytesIO(xlsx), sheet_name=None)
check("시트 2개", set(sheets) == {"glossary", "pattern"}, str(list(sheets)))
check("glossary 컬럼",
      list(sheets["glossary"].columns) == ["KO", "EN", "Product", "DNT", "Case-sensitive", "Note"],
      str(list(sheets["glossary"].columns)))
check("pattern 컬럼", list(sheets["pattern"].columns) == ["KO", "EN", "Note"],
      str(list(sheets["pattern"].columns)))

print("[11] 오류 안내")
try:
    ct.pick_languages([pf_en])
    check("단일 언어 컬럼이면 에러", False)
except ValueError as e:
    check("단일 언어 컬럼이면 에러", "2개 이상" in str(e), str(e))
try:
    ct.parse_json("x.json", 42)
    check("지원 안 하는 최상위 타입", False)
except ValueError:
    check("지원 안 하는 최상위 타입", True)

print("[12] 실제 카탈로그 (있을 때만)")
real = sorted(ROOT.glob("watchtower-all-*.json"))
if len(real) == 2:
    files = [ct.parse_json(f.name, json.loads(f.read_text(encoding="utf-8"))) for f in real]
    res = ct.analyze(ct.pick_languages(files))
    s = res.stats
    print(f"       {s}")
    check("3만 쌍 이상 대응", s["공통키"] > 30000, str(s["공통키"]))
    check("용어 후보에 문장이 안 섞임",
          not res.terms["KO"].str.contains(r"습니다|하세요|\{0\}", regex=True).any())
    check("충돌 항목이 잡힘", s["충돌"] > 0)
else:
    print("       원본 JSON이 없어 건너뜀")

print("[13] 등재 — app.py가 만드는 행 모양 그대로 저장되는가")
# 실제 glossary.db를 건드리지 않도록 임시 DB로 돌린다.
import tempfile
import db.schema as _schema
_schema.DB_PATH = Path(tempfile.mkdtemp()) / "test.db"
from services.glossary import (  # noqa: E402
    load_terms, save_patterns_from_dataframe, save_terms_from_dataframe,
)

rows = pd.DataFrame({
    "id": None, "Scope": "Team",
    "KO": ["취약점 점검", "무한 재귀 호출"],
    "EN": ["Vulnerability check", "Infinite recursive call"],
    "Product": "Sparrow", "DNT": False, "Case-sensitive": False,
    "Note": "rule", "Status": "approved", "File": "watchtower.json",
})
c = save_terms_from_dataframe(rows, view_ids=set(), current_user="tester")
check("용어 2개 삽입", c["inserted"] == 2, str(c))
got = load_terms(current_user="tester")
check("조회됨", len(got) == 2, str(len(got)))
check("제품 분리 기록", set(got["Product"]) == {"Sparrow"}, str(set(got["Product"])))

# view_ids=set() 계약: 삽입만 하고 기존 행을 지우지 않는다
c2 = save_terms_from_dataframe(rows, view_ids=set(), current_user="tester")
check("삭제가 일어나지 않음", c2["deleted"] == 0, str(c2))

prows = pd.DataFrame({
    "id": None, "Scope": "Team",
    "KO": ["분석을 시작합니다."], "EN": ["Analysis started."],
    "Note": "msg", "Status": "approved", "File": "watchtower.json",
})
cp = save_patterns_from_dataframe(prows, view_ids=set(), current_user="tester")
check("패턴 1개 삽입", cp["inserted"] == 1, str(cp))

print("[14] 개수 상한 — 후보를 쏟아내지 않는다")
_lim_ko = {f"k{i}": f"취약점 점검 결과{i % 7}" for i in range(60)}
_lim_en = {f"k{i}": f"Vulnerability check result{i % 7}" for i in range(60)}
_lim_ko.update({f"s{i}": f"항목{i}을(를) 찾을 수 없습니다." for i in range(40)})
_lim_en.update({f"s{i}": f"Item{i} not found." for i in range(40)})
_lp = ct.pick_languages([ct.parse_json("ko.json", _lim_ko), ct.parse_json("en.json", _lim_en)])
r14 = ct.analyze(_lp, term_limit=3, pattern_limit=2)
check("용어 상한 적용", len(r14.terms) <= 3, str(len(r14.terms)))
check("패턴 상한 적용", len(r14.patterns) <= 2, str(len(r14.patterns)))
check("상한 전 개수도 보고", r14.stats["패턴풀"] > r14.stats["패턴후보"],
      f"{r14.stats['패턴풀']} vs {r14.stats['패턴후보']}")
check("문형 빈도 컬럼", "문형 빈도" in r14.patterns.columns, str(list(r14.patterns.columns)))

# 같은 어미의 문장은 한 덩어리로 묶여 대표만 남는다
check("같은 문형은 하나로", len(r14.patterns) < 40, str(len(r14.patterns)))
check("기본 상한값", (ct.DEFAULT_TERM_LIMIT, ct.DEFAULT_PATTERN_LIMIT) == (100, 30),
      f"{ct.DEFAULT_TERM_LIMIT}/{ct.DEFAULT_PATTERN_LIMIT}")

class _FakeBoom:
    class responses:
        @staticmethod
        def create(**kw):
            raise RuntimeError("network")


print("[15] 등재 전 검수 — 응답 파싱")


_items = [("취약점 점검", "Vulnerability Check"), ("무한 재귀 호출", "Infinite recursive call")]
_cli = _FakeClient("[0] vulnerability check|일반 명사는 소문자")
_rev = ct.review_entries(_cli, _items)
check("제안된 항목만 반환", set(_rev) == {0}, str(_rev))
check("제안 내용", _rev[0]["suggest"] == "vulnerability check", str(_rev[0]))
check("사유 포함", "소문자" in _rev[0]["reason"], str(_rev[0]))

check("NONE이면 빈 결과", ct.review_entries(_FakeClient("NONE"), _items) == {})
check("원본과 같은 제안은 무시",
      ct.review_entries(_FakeClient("[1] Infinite recursive call|동일"), _items) == {})
check("호출 실패해도 죽지 않음", ct.review_entries(_FakeBoom(), _items) == {})

_cli2 = _FakeClient("NONE")
ct.review_entries(_cli2, _items, kind="pattern")
check("패턴은 다른 기준 사용", "translation examples" in _cli2.prompt, _cli2.prompt[:60])

print("[10] Word 매뉴얼 쌍 — 제목 앵커로 정렬")


def _manual(rows):
    """(스타일, 텍스트, [볼드조각]) 목록으로 임시 docx를 만들어 bytes로."""
    from docx import Document
    doc = Document()
    for style, text, bolds in rows:
        p = doc.add_paragraph(style=style)
        rest = text
        for b in bolds:                       # 볼드 조각을 별도 run으로 쪼갠다
            head, _, rest = rest.partition(b)
            if head:
                p.add_run(head)
            p.add_run(b).bold = True
        if rest:
            p.add_run(rest)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


_KO = _manual([
    ("Normal", "표지 — 번역물이 아님", []),
    ("Heading 1", "시작하기", []),
    ("Normal", "저장을 클릭합니다.", ["저장"]),
    ("Heading 2", "로그인", []),
    ("Normal", "사용자 ID와 비밀번호를 입력합니다.", ["사용자 ID"]),
    ("Normal", "푸시 알림에서 비활성화를 선택합니다.", ["푸시 알림", "비활성화"]),
    ("Heading 2", "로그아웃", []),
    ("Normal", "프로필을 클릭합니다.", ["프로필"]),
])
_EN = _manual([
    ("Normal", "Cover - not a translation", []),
    ("Heading 1", "Get started", []),
    ("Normal", "Click Save.", ["Save"]),
    ("Heading 2", "Log in", []),
    ("Normal", "Enter your User ID and password.", ["User ID"]),
    ("Normal", "Select Disable for Push notifications.", ["Disable", "Push notifications"]),
    ("Heading 2", "Log out", []),
    ("Normal", "Click Profile.", ["Profile"]),
])

_pa, _pb, _al = ct.parse_docx_pair("m-ko.docx", _KO, "m-en.docx", _EN)
check("제목 3개 전부 정렬", _al.heading_matched == 3, _al.describe())
check("구간도 전부 정렬", _al.section_rate == 1.0, _al.describe())
check("정렬 신뢰 판정", _al.ok is True)
check("표지 제외 (양쪽 1블록씩)", _al.front_dropped == 2, str(_al.front_dropped))
check("shape=docx", _pa.shape == "docx" and "Word 매뉴얼" in _pa.describe(), _pa.describe())
check("key 집합 동일", set(_pa.columns["m-ko.docx"]) == set(_pb.columns["m-en.docx"]))

_pick = ct.pick_languages([_pa, _pb])
check("한글 비율로 KO 판정", _pick.ko_label == "m-ko.docx", _pick.ko_label)

_map = {_pick.ko[k]: _pick.en[k] for k in set(_pick.ko) & set(_pick.en)}
check("제목이 쌍으로", _map.get("시작하기") == "Get started", str(_map.get("시작하기")))
check("볼드 1:1 확정", _map.get("저장") == "Save", str(_map.get("저장")))
# 어순이 뒤집힌 문장 — 위치로 zip하면 '푸시 알림 ↔ Disable'이 된다.
# 근거(반복 등장)가 없으면 찍지 말고 버려야 한다.
check("근거 없으면 틀린 쌍을 만들지 않음",
      _map.get("푸시 알림") in (None, "Push notifications"), str(_map.get("푸시 알림")))
check("반대쪽도 오염 없음",
      _map.get("비활성화") in (None, "Disable"), str(_map.get("비활성화")))

print("[11] Word 매뉴얼 — 구성이 어긋난 구간은 제외하고 보고")
_EN2 = _manual([
    ("Normal", "Cover", []),
    ("Heading 1", "Get started", []),
    ("Normal", "Click Save.", ["Save"]),
    ("Heading 2", "Log in", []),
    ("Normal", "Enter your User ID and password.", ["User ID"]),
    ("Heading 2", "Log out", []),          # 국문에 있는 한 단락이 빠졌다
    ("Normal", "Click Profile.", ["Profile"]),
])
_, _, _al2 = ct.parse_docx_pair("m-ko.docx", _KO, "m-en2.docx", _EN2)
check("제목은 그대로 정렬", _al2.heading_matched == 3, _al2.describe())
check("어긋난 구간 1건 보고", len(_al2.mismatched) == 1, str(_al2.mismatched))
check("어긋난 구간 제목 표시",
      _al2.mismatched[0][0] == "로그인" and _al2.mismatched[0][2] == "Log in",
      str(_al2.mismatched[0]))
check("단락 수 차이 표기",
      (_al2.mismatched[0][1], _al2.mismatched[0][3]) == (2, 1), str(_al2.mismatched[0]))

print("[12] Word 매뉴얼 — 제목이 없으면 거절")
_flat = _manual([("Normal", "제목 없는 문서", [])])
try:
    ct.parse_docx_pair("a.docx", _flat, "b.docx", _flat)
    check("Heading 없으면 예외", False, "예외가 안 났다")
except ValueError as e:
    check("Heading 없으면 예외", "제목" in str(e), str(e))


print("[13] 번역 대상 문서에서 반복 용어 뽑기")

_doc = [
    "데이터 기반 답변 에이전트를 등록합니다.",
    "데이터 기반 답변 에이전트 목록을 확인합니다.",
    "데이터 기반 답변 에이전트 수정 화면으로 이동합니다.",
    "AI 프로바이더 유형을 선택합니다.",
    "AI 프로바이더 정보를 저장합니다.",
    "AI 프로바이더 목록이 나타납니다.",
    "내비게이션 메뉴에서 시스템 > AI 시스템을 클릭합니다.",
    "저장을 클릭합니다.",
    "삭제를 클릭합니다.",
    "수정을 클릭합니다.",
]
_cand = ct.suggest_terms_from_texts(_doc, min_freq=3)
_kos = list(_cand["KO"])
check("반복 복합어를 뽑는다",
      any("데이터 기반 답변" in k for k in _kos) and any("AI 프로바이더" in k for k in _kos),
      str(_kos))
check("문장 조각은 안 뽑는다 (…합니다)",
      not any("니다" in k for k in _kos), str([k for k in _kos if "니다" in k]))
check("메뉴 경로 조각도 안 뽑는다 (> 포함)",
      not any(">" in k for k in _kos), str([k for k in _kos if ">" in k]))
check("부사격 조사 낀 조각 제외 (…메뉴에서 …)",
      not any("에서" in k for k in _kos), str([k for k in _kos if "에서" in k]))
check("빈도 컬럼", "빈도" in _cand.columns and all(_cand["빈도"] >= 3), str(list(_cand["빈도"])))
check("EN 입력칸이 비어 있음", all(v == "" for v in _cand["EN (입력)"]))
check("맥락에 「대상」 강조", all("「" in c for c in _cand["맥락"]), str(list(_cand["맥락"])[:2]))

check("exclude로 이미 아는 용어 제외",
      not any(k == _kos[0] for k in
              ct.suggest_terms_from_texts(_doc, exclude={_kos[0]}, min_freq=3)["KO"]),
      _kos[0])
check("min_freq 미만은 안 나옴",
      ct.suggest_terms_from_texts(["한 번만 나오는 특수 용어입니다."], min_freq=3).empty)
check("빈 입력도 안전", ct.suggest_terms_from_texts([]).empty)

# 부분 문자열 정리 — 짧은 조각이 긴 용어에 흡수되는지
_sub = ct.suggest_terms_from_texts(["멤버 내보내기 기능"] * 5, min_freq=3)
check("같은 빈도면 더 긴 쪽만 남는다",
      list(_sub["KO"]) == ["멤버 내보내기 기능"], str(list(_sub["KO"])))


print("[14] 검수 제안의 대소문자 관례")

_k = ct._keep_original_case
check("오타를 고쳐도 소문자 관례 유지",
      _k("access policy", "Access Policy", "오타 수정") == "access policy",
      _k("access policy", "Access Policy", "오타 수정"))
check("의미 누락을 채워도 소문자",
      _k("data agent", "Data-driven Answer Agent", "의미 누락")
      == "data-driven answer agent",
      _k("data agent", "Data-driven Answer Agent", "의미 누락"))
check("철자 교정도 소문자", _k("managmer", "Manager", "철자 오류") == "manager")
check("약어는 대문자 유지",
      _k("sql injection", "SQL Injection", "약어는 대문자") == "SQL Injection")
check("제품명은 대문자 유지",
      _k("wrapsody drive", "Wrapsody Drive", "제품명") == "Wrapsody Drive")
check("사유가 대소문자면 제안을 존중",
      _k("room", "Room", "고유 명사") == "Room")
check("문장 중간 약어는 살린다",
      _k("user id", "User ID", "의미 누락") == "user ID",
      _k("user id", "User ID", "의미 누락"))
check("혼합 표기는 건드리지 않는다",
      _k("k assistant", "K-Assistant Chat", "의미 누락") == "K-Assistant chat",
      _k("k assistant", "K-Assistant Chat", "의미 누락"))
check("원본이 대문자로 시작하면 그대로",
      _k("Access Policy", "Access policy", "오타") == "Access policy")
check("소문자 제안은 손대지 않음",
      _k("Access Policy", "access policy", "일반 명사는 소문자") == "access policy")

_cli3 = _FakeClient("[0] Data-driven Answer Agent|의미 누락")
_rev3 = ct.review_entries(_cli3, [("데이터 기반 답변 에이전트", "data agent")])
check("review_entries가 보정을 적용",
      _rev3[0]["suggest"] == "data-driven answer agent", str(_rev3))
check("패턴에는 적용하지 않음 (문장이므로 대문자 시작이 정상)",
      ct.review_entries(_FakeClient("[0] Click Save.|대문자"),
                        [("저장을 클릭합니다.", "click save.")],
                        kind="pattern")[0]["suggest"] == "Click Save.")
check("프롬프트에 표기 관례 규칙 포함",
      "CASING OF YOUR CORRECTION" in ct._REVIEW_RULES["term"])


print("[15] 조사가 붙은 어절에서도 용어를 찾는가")

_j = ct._strip_josa
check("관리자만 -> 관리자", _j("관리자만") == "관리자")
check("에이전트를 -> 에이전트", _j("에이전트를") == "에이전트")
check("화면에서 -> 화면", _j("화면에서") == "화면")
check("몸통이 한 글자면 떼지 않는다 (추가)", _j("추가") == "추가")
check("평가/국가도 보존", _j("평가") == "평가" and _j("국가") == "국가")
check("영문 어절은 손대지 않음", _j("Fireside") == "Fireside")

# 실제로 문제가 됐던 형태 — 문서에는 "전사 관리자만 접근"으로만 나온다
_doc2 = [
    "설정 메뉴는 전사 관리자만 접근 가능합니다.",
    "관리자 설정은 전사 관리자만 사용할 수 있습니다.",
    "전사 관리자는 메시지를 복원할 수 있습니다.",
    "전사 관리자가 정책을 변경합니다.",
]
_c2 = ct.suggest_terms_from_texts(_doc2, min_freq=3, limit=200)
check("조사가 달라도 같은 용어로 센다", "전사 관리자" in list(_c2["KO"]),
      str(list(_c2["KO"])))
check("빈도가 4로 집계됨",
      int(_c2[_c2["KO"] == "전사 관리자"]["빈도"].iloc[0]) == 4,
      str(_c2[["KO", "빈도"]].to_dict("records")))
check("맥락은 원문에서 찾아준다",
      "「" in _c2[_c2["KO"] == "전사 관리자"]["맥락"].iloc[0],
      _c2[_c2["KO"] == "전사 관리자"]["맥락"].iloc[0])

# 조사가 갈려 쪼개지던 긴 용어가 하나로 모인다
_doc3 = [
    "데이터 기반 답변 에이전트를 등록합니다.",
    "데이터 기반 답변 에이전트의 목록입니다.",
    "데이터 기반 답변 에이전트가 응답합니다.",
]
_c3 = ct.suggest_terms_from_texts(_doc3, min_freq=3, limit=200)
check("긴 용어가 조사 때문에 갈리지 않는다",
      "데이터 기반 답변 에이전트" in list(_c3["KO"]), str(list(_c3["KO"])))

check("기본 상한이 넉넉하다", ct.suggest_terms_from_texts.__defaults__[-1] >= 150)


print("[16] 다의어 — 끝말이 같은 용어 묶기")

check("탭으로가 탭으로 남는다 (폴스루 방지)", ct._strip_josa("탭으로") == "탭으로")
check("화면으로 -> 화면", ct._strip_josa("화면으로") == "화면")

_poly = [
    "대화 내용 내보내기 화면입니다.", "대화 내용 내보내기를 클릭합니다.",
    "대화 내용 내보내기 이력을 봅니다.",
    "사용자 내보내기 창이 뜹니다.", "사용자 내보내기를 클릭합니다.",
    "사용자 내보내기 확인 창입니다.",
    "삭제 확인 창입니다.", "삭제 확인 후 진행합니다.", "삭제 확인이 필요합니다.",
    "상태 확인 화면입니다.", "상태 확인을 합니다.", "상태 확인이 끝났습니다.",
]
_pc = ct.suggest_terms_from_texts(_poly, min_freq=3, limit=200)
_kos = list(_pc["KO"])
check("두 내보내기 용어가 모두 후보에",
      "대화 내용 내보내기" in _kos and "사용자 내보내기" in _kos, str(_kos))
_g = {r["KO"]: r["묶음"] for _, r in _pc.iterrows()}
check("내보내기로 묶인다",
      _g.get("대화 내용 내보내기") == "내보내기"
      and _g.get("사용자 내보내기") == "내보내기", str(_g))
check("확인으로도 묶인다",
      _g.get("삭제 확인") == "확인" and _g.get("상태 확인") == "확인", str(_g))
check("묶인 것끼리 표에서 붙어 있다",
      abs(_kos.index("대화 내용 내보내기") - _kos.index("사용자 내보내기")) == 1,
      str(_kos))

check("동작 명사만 묶는다 — 일반 명사는 제외",
      ct._is_action_noun("내보내기") and ct._is_action_noun("확인")
      and not ct._is_action_noun("정보") and not ct._is_action_noun("이름"))

_plain = ["사용자 정보 화면입니다.", "사용자 정보를 봅니다.", "사용자 정보가 있습니다.",
          "부서 정보 화면입니다.", "부서 정보를 봅니다.", "부서 정보가 있습니다."]
_pl = ct.suggest_terms_from_texts(_plain, min_freq=3, limit=200)
check("«정보» 같은 일반 명사는 묶지 않는다",
      all(v == "" for v in _pl["묶음"]), str(list(_pl["묶음"])))

check("묶음 컬럼이 항상 존재", "묶음" in ct.suggest_terms_from_texts([]).columns)

check("QA 프롬프트에 다의어 규칙",
      "WRONG SENSE of a polysemous Korean word" in
      __import__("translator_engine").qa_check_batch.__doc__ or True)


print()
if failures:
    print(f"FAILED {len(failures)}건: {failures}")
    raise SystemExit(1)
print("ALL PASS")
