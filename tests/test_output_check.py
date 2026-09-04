"""
산출물 검증 테스트 — 마커 무결성 + 저장 후 대조.

이 스위트가 지키는 것은 "번역이 잘 됐는가"가 아니라 **"잘못된 결과를 코드가
잡아내는가"** 다. 실제 산출물에서 나왔던 실패 모드를 그대로 재현해 두고,
검증기가 그걸 놓치지 않는지 본다.

    python tests/test_output_check.py
"""
from __future__ import annotations

import io
import sys
from pathlib import Path

ROOT = Path(__file__).resolve().parent.parent
sys.path.insert(0, str(ROOT))
sys.path.insert(0, str(ROOT / "tests"))

import output_check as oc
import translator_engine as te

failures = []


def check(name, cond, detail=""):
    print(f"  {'OK  ' if cond else 'FAIL'} {name}" + (f"  {detail}" if not cond else ""))
    if not cond:
        failures.append(name)


def has(findings, needle):
    return any(needle in f.title or needle in f.detail for f in findings)


# ══════════════════════════════════════════════════════════════════
print("[1] 후처리 회귀 — 예전에 본문을 망가뜨리던 것들")

check("v2.1 업데이트가 살아남음",
      te.strip_meta_version_labels("v2.1 update") == "v2.1 update",
      te.strip_meta_version_labels("v2.1 update"))
check("문장 끝 final. 이 잘리지 않음",
      te.strip_meta_version_labels("We ship the final.") == "We ship the final.",
      te.strip_meta_version_labels("We ship the final."))
check("Updated for v2.0 보존",
      te.strip_meta_version_labels("Updated for v2.0") == "Updated for v2.0",
      te.strip_meta_version_labels("Updated for v2.0"))
check("맨 앞 Draft: 라벨은 제거",
      te.strip_meta_version_labels("Draft: Save the file") == "Save the file",
      te.strip_meta_version_labels("Draft: Save the file"))
check("라벨 형식 감지", te.has_meta_version_labels("Draft: A\nRevised: B"))
check("일반 문장은 라벨 아님", not te.has_meta_version_labels("v2.1 update"))
check("관사 자동교정 제거됨", not hasattr(te, "fix_indefinite_articles"))

# ══════════════════════════════════════════════════════════════════
print("[2] 마커 무결성 — 실제 산출물에서 나온 손상 형태")

check("정상 통과",
      te.check_marker_integrity("⟦B⟧규칙⟦/B⟧ 설정", "⟦B⟧Rule⟦/B⟧ settings") == [])
_p = te.check_marker_integrity("⟦HL:yellow⟧AI 공급자 유형⟦/HL⟧",
                               "⟦HL:yellowAI provider type")
check("⟦HL:yellow… 삼킴 감지", _p, str(_p))
check("손상 마커로 보고", any("손상된 마커" in x for x in _p), str(_p))

_p = te.check_marker_integrity("⟦B⟧스키마⟦/B⟧", "⟦/b schema")
check("⟦/b 감지", any("손상된 마커" in x for x in _p), str(_p))

check("닫기 누락 감지",
      te.check_marker_integrity("⟦B⟧A⟦/B⟧ ⟦B⟧B⟦/B⟧", "⟦B⟧A ⟦B⟧B⟦/B⟧") != [])
check("열기/닫기 뒤바뀜 감지",
      any("짝이 어긋" in x
          for x in te.check_marker_integrity("⟦B⟧x⟦/B⟧", "⟦/B⟧x⟦B⟧")))
check("하이퍼링크 소실 감지",
      te.check_marker_integrity("⟦H1⟧링크⟦/H1⟧", "link") != [])
check("줄바꿈 소실 감지",
      te.check_marker_integrity("a⟦LB⟧b", "a b") != [])
check("탭 소실 감지",
      te.check_marker_integrity("a⟦TB⟧b", "a b") != [])
check("하이라이트 색 변경 감지",
      any("색상" in x for x in
          te.check_marker_integrity("⟦HL:yellow⟧x⟦/HL⟧", "⟦HL:green⟧x⟦/HL⟧")))
check("용어 자리표시자 소실 감지",
      te.check_marker_integrity("⟦G0⟧를 클릭", "Click it") != [])
check("없던 자리표시자 생성 감지",
      te.check_marker_integrity("클릭", "Click ⟦G9⟧") != [])
check("마커 순서가 바뀌는 건 허용",
      te.check_marker_integrity("⟦G0⟧ 뒤에 ⟦G1⟧", "⟦G1⟧ after ⟦G0⟧") == [])

# ══════════════════════════════════════════════════════════════════
print("[3] 저장 후 대조 — 망가진 번역본을 만들어 검증기에 통과시킨다")


def _docx(rows, path):
    """rows: (텍스트, 하이라이트색|None, 특수요소|None)"""
    from docx import Document
    from docx.enum.text import WD_COLOR_INDEX
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()
    for text, hl, special in rows:
        p = doc.add_paragraph()
        r = p.add_run(text)
        if hl:
            r.font.highlight_color = getattr(WD_COLOR_INDEX, hl)
        if special == "br":
            r._r.append(OxmlElement("w:br"))
        elif special == "tab":
            r._r.append(OxmlElement("w:tab"))
    doc.save(path)
    return path


TMP = ROOT / "outputs"
TMP.mkdir(exist_ok=True)
_src = _docx([
    ("v2.1 업데이트 내용입니다.", None, None),
    ("노란색 강조된 문장.", "YELLOW", None),
    ("줄바꿈이 있는 문장.", None, "br"),
    ("탭이 있는 문장.", None, "tab"),
    ("K-Assistant 설명 문단.", None, None),
], str(TMP / "_chk_src.docx"))

_bad = _docx([
    ("1 update content.", None, None),          # 버전이 잘림
    ("Highlighted sentence.", None, None),      # 하이라이트 소실
    ("⟦HL:yellowSentence with break.", None, None),   # 마커 노출 + 줄바꿈 소실
    ("Sentence with tab.", None, None),         # 탭 소실
    ("AI Assistant 설명 문단.", None, None),     # 리터럴 변경 + 한국어 잔존
], str(TMP / "_chk_bad.docx"))

_f, _s, _o = oc.verify(_src, _bad, literals=["K-Assistant"])
check("마커 노출 감지", has(_f, "마커가 문서에 그대로"), str(_f))
check("한국어 잔존 감지", has(_f, "한국어가 남아"), str(_f))
check("버전 삭제 감지", has(_f, "버전 표기가 사라"), str(_f))
check("하이라이트 소실 감지", has(_f, "하이라이트가 사라진 문단"), str(_f))
check("하이라이트를 run이 아닌 문단으로 셈", _s.highlight_paras == 1, str(_s.highlight_paras))
check("줄바꿈 소실 감지", has(_f, "줄바꿈 개수"), str(_f))
check("탭 소실 감지", has(_f, "탭 개수"), str(_f))
check("원문 영문 변경 감지", has(_f, "K-Assistant"), str(_f))
check("문단 수는 같으므로 문제 없음", not has(_f, "문단 개수"), str(_f))

_good = _docx([
    ("v2.1 update content.", None, None),
    ("Highlighted sentence.", "YELLOW", None),
    ("Sentence with break.", None, "br"),
    ("Sentence with tab.", None, "tab"),
    ("K-Assistant description paragraph.", None, None),
], str(TMP / "_chk_good.docx"))
_f2, _, _ = oc.verify(_src, _good, literals=["K-Assistant"])
check("정상 번역본은 오류 0건",
      not [x for x in _f2 if x.level == "오류"], str(_f2))

# DNT 로 한국어 제품명 허용
_dnt_out = _docx([
    ("v2.1 update content.", None, None),
    ("Highlighted sentence.", "YELLOW", None),
    ("Sentence with break.", None, "br"),
    ("Sentence with tab.", None, "tab"),
    ("파수 K-Assistant description.", None, None),
], str(TMP / "_chk_dnt.docx"))
check("DNT 없으면 한국어로 잡힘",
      has(oc.verify(_src, _dnt_out)[0], "한국어가 남아"))
check("DNT 지정하면 통과",
      not has(oc.verify(_src, _dnt_out, dnt=["파수"])[0], "한국어가 남아"))

print("[4] 머리글·바닥글까지 본다")
check("본문 외 파트도 수집 대상",
      "word/document.xml" in oc.collect(_src).parts,
      str(oc.collect(_src).parts))

for _p_ in ("_chk_src", "_chk_bad", "_chk_good", "_chk_dnt"):
    (TMP / f"{_p_}.docx").unlink(missing_ok=True)

# ══════════════════════════════════════════════════════════════════
print("[5] 경계 공백 — 삭제하지 말고 마커 밖으로")

_ns = te.normalize_marker_boundary_spaces
check("Feedbacktab 방지",
      _ns("⟦B⟧Agent Feedback ⟦/B⟧tab") == "⟦B⟧Agent Feedback⟦/B⟧ tab",
      _ns("⟦B⟧Agent Feedback ⟦/B⟧tab"))
check("여는 마커 앞으로도 이동",
      _ns("Click⟦B⟧ Delete Agent⟦/B⟧") == "Click ⟦B⟧Delete Agent⟦/B⟧",
      _ns("Click⟦B⟧ Delete Agent⟦/B⟧"))
check("구두점 앞 공백은 되돌림",
      _ns("Click ⟦B⟧Save ⟦/B⟧.") == "Click ⟦B⟧Save⟦/B⟧.",
      _ns("Click ⟦B⟧Save ⟦/B⟧."))
check("줄 머리로 밀린 공백은 버림",
      _ns("⟦B⟧ Leading⟦/B⟧ x") == "⟦B⟧Leading⟦/B⟧ x",
      _ns("⟦B⟧ Leading⟦/B⟧ x"))
check("하이라이트 경계도 처리",
      _ns("x⟦HL:yellow⟧ y ⟦/HL⟧z") == "x ⟦HL:yellow⟧y⟦/HL⟧ z",
      _ns("x⟦HL:yellow⟧ y ⟦/HL⟧z"))
check("마커 없으면 그대로", _ns("plain text") == "plain text")

# ══════════════════════════════════════════════════════════════════
print("[6] 손상 마커 제거 — 추측하지 말고 지운다")

_sb = te.sanitize_broken_markers
check("⟦HL:yellow… 잔재 제거, 본문 AI는 보존",
      _sb("Click ⟦HL:yellowAI provider type") == "Click AI provider type",
      _sb("Click ⟦HL:yellowAI provider type"))
check("⟦/b 제거", _sb("Edit info⟦/b") == "Edit info", _sb("Edit info⟦/b"))
check("깨진 하이퍼링크 제거", "⟦" not in _sb("⟦H12 link"))
check("정상 마커는 보존",
      _sb("⟦B⟧ok⟦/B⟧ ⟦HL:yellow⟧hl⟦/HL⟧") == "⟦B⟧ok⟦/B⟧ ⟦HL:yellow⟧hl⟦/HL⟧")
check("⟦LB⟧ 보존", _sb("a⟦LB⟧b") == "a⟦LB⟧b")
check("마커 없는 텍스트는 무변경", _sb("plain") == "plain")

# ══════════════════════════════════════════════════════════════════
print("[7] 기록 직전 최종 게이트")

_t, _p = te.finalize_markers("⟦HL:yellow⟧AI 공급자 유형⟦/HL⟧",
                             "Click ⟦HL:yellowAI provider type")
check("손상돼도 하이라이트를 되살림", _t.startswith("⟦HL:yellow⟧") and _t.endswith("⟦/HL⟧"), _t)
check("본문에 마커 잔재 없음", "⟦HL:yellowAI" not in _t, _t)
check("남은 문제 없음", _p == [], str(_p))

_t, _p = te.finalize_markers("⟦HL:yellow⟧⟦B⟧부서 추가⟦/B⟧ 창⟦/HL⟧",
                             "The ⟦B⟧Add Department ⟦/B⟧window")
check("공백 이동 + 하이라이트 복원 동시",
      _t == "⟦HL:yellow⟧The ⟦B⟧Add Department⟦/B⟧ window⟦/HL⟧", _t)

_t, _p = te.finalize_markers("⟦B⟧스키마⟦/B⟧ 수정",
                             "Edit the agent schema⟦/b")
check("복구 불가하면 잔재는 지우고 문제로 보고",
      "⟦" not in _t and _p, f"{_t!r} {_p}")

_t, _p = te.finalize_markers("⟦B⟧정상⟦/B⟧", "⟦B⟧Normal⟦/B⟧")
check("정상 입력은 무변경", _t == "⟦B⟧Normal⟦/B⟧" and _p == [], f"{_t!r} {_p}")


# ══════════════════════════════════════════════════════════════════
print("[8] 원문 영문 봉인 — 지시가 아니라 구조로 지킨다")

_sl, _rl = te.seal_literals, te.restore_literals


def _roundtrip(t):
    sealed, m = _sl(t)
    return sealed, m, _rl(sealed, m) == t


_s1, _m1, _ok1 = _roundtrip("K-Assistant는 AI 챗봇입니다. API를 쓰세요.")
check("한글이 붙어도 봉인됨 (K-Assistant는)",
      set(_m1.values()) == {"K-Assistant", "AI", "API"}, str(_m1))
check("왕복 무손실", _ok1, _s1)

_s2, _m2, _ok2 = _roundtrip("⟦HL:yellow⟧v2.1 업데이트⟦/HL⟧")
check("버전 봉인", "v2.1" in _m2.values(), str(_m2))
check("마커 안쪽은 건드리지 않음 — HL이 약어로 잡히지 않음",
      "yellow" not in str(_m2.values()) and _s2.startswith("⟦HL:yellow⟧"), _s2)
check("왕복 무손실", _ok2)

_s3, _m3, _ok3 = _roundtrip("MS-SQL 2012 이상 / JDK v1.7 이상")
check("하이픈 고유명사 봉인", "MS-SQL" in _m3.values(), str(_m3))
check("왕복 무손실", _ok3)

_s4, _m4, _ok4 = _roundtrip("Fireside Room에서 Wrapsody 문서를 엽니다.")
check("일반 대문자 시작 단어는 봉인하지 않음", _m4 == {}, str(_m4))

_s5, _m5, _ok5 = _roundtrip("문서는 https://a.com/b.pdf 와 a@b.com 참고")
check("URL·이메일 봉인", len(_m5) == 2, str(_m5))
check("왕복 무손실", _ok5)

check("봉인할 게 없으면 무변경", _sl("일반 문장입니다.") == ("일반 문장입니다.", {}))

# 봉인 상태에서 대소문자 정규화를 통과시켜도 원문이 지켜지는가
_sealed, _map = _sl("⟦B⟧기본 AI 기능⟦/B⟧ 활성화")
_cased = te._sentence_case_preserving_markers(_sealed.replace("기본", "Enable basic")
                                              .replace("기능", "features")
                                              .replace(" 활성화", ""))
check("문장 대소문자 정규화를 통과해도 AI가 살아남음",
      "AI" in _rl(_cased, _map), _rl(_cased, _map))

print("[9] QA 이후 대소문자 방어")
_ec = te.enforce_literal_casing
check("소문자로 내려간 약어 복구",
      _ec("Enable basic ai features and the api", "AI API 기능")
      == "Enable basic AI features and the API",
      _ec("Enable basic ai features and the api", "AI API 기능"))
check("원문에 없는 단어는 건드리지 않음",
      _ec("nothing to fix here", "AI 기능") == "nothing to fix here")
check("버전 대소문자 복구",
      _ec("V2.1 update", "v2.1 업데이트") == "v2.1 update",
      _ec("V2.1 update", "v2.1 업데이트"))

check("⟦X#⟧도 마커 무결성 검사 대상",
      te.check_marker_integrity("⟦X0⟧ 설명", "description") != [])
check("⟦X#⟧는 손상 마커 제거에서 보존",
      te.sanitize_broken_markers("⟦X0⟧ ok") == "⟦X0⟧ ok")


# ══════════════════════════════════════════════════════════════════
print("[10] 용어별 영문 표기 갈림 검사")

_src2 = _docx([("데이터 기반 답변 에이전트를 등록합니다.", None, None),
               ("데이터 기반 답변 에이전트 목록입니다.", None, None),
               ("데이터 기반 답변 에이전트를 수정합니다.", None, None),
               ("데이터 기반 답변 에이전트를 삭제합니다.", None, None)],
              str(TMP / "_c_src.docx"))
_bad2 = _docx([("Register a data-driven response agent.", None, None),
               ("List of data-based answer agents.", None, None),
               ("Edit the data-driven answering agent.", None, None),
               ("Delete the data-based response agent.", None, None)],
              str(TMP / "_c_bad.docx"))
_good2 = _docx([("Register a data-driven response agent.", None, None),
                ("List of data-driven response agents.", None, None),
                ("Edit the data-driven response agent.", None, None),
                ("Delete the data-driven response agent.", None, None)],
               str(TMP / "_c_good.docx"))

_r_bad = oc.check_term_consistency(oc.collect(_src2), oc.collect(_bad2))
check("표기가 갈리면 잡아낸다", _r_bad, str(_r_bad))
check("커버리지가 낮게 보고됨", _r_bad and _r_bad[0][1] < 0.7, str(_r_bad))
check("경쟁 표기를 함께 보여줌", _r_bad and _r_bad[0][2], str(_r_bad))

_r_good = oc.check_term_consistency(oc.collect(_src2), oc.collect(_good2))
check("일관되면 잡지 않는다", not _r_good, str(_r_good))

check("1단어 조각으로 판단하지 않음 — 2단어 이상만",
      all(len(g.split()) >= 2 for g in oc._en_ngrams("Register a data-driven response agent")),
      str(list(oc._en_ngrams("Register a data-driven response agent"))[:4]))
_short2 = _docx([("Only one paragraph.", None, None)], str(TMP / "_c_short.docx"))
check("문단 수가 어긋나면 판단 보류",
      oc.check_term_consistency(oc.collect(_src2), oc.collect(_short2)) == [])

_f3, _, _ = oc.verify(_src2, _bad2)
check("verify 결과에 경고로 실림",
      any("영문 표기가 갈립니다" in x.title for x in _f3), str(_f3))

for _n in ("_c_src", "_c_bad", "_c_good", "_c_short"):
    (TMP / f"{_n}.docx").unlink(missing_ok=True)

# ══════════════════════════════════════════════════════════════════
print("[11] QA — 제목 포함 · 그룹키 분리 · 글로서리 선별")

_seen = {}


class _FakeQAClient:
    pass


_orig_qa = te.qa_check_batch


def _spy_qa(client, items, style_guide, glossary_pairs, model="gpt-5.2"):
    _seen["items"] = list(items)
    _seen["glossary"] = list(glossary_pairs)
    return {}


try:
    import stub_llm
    stub_llm.install()          # 먼저 스텁을 깔고
    te.qa_check_batch = _spy_qa  # 그 위에 스파이를 얹는다 (순서 중요)
    _hsrc = ROOT / "outputs" / "_qa_src.docx"
    from docx import Document as _D
    _d = _D()
    _p = _d.add_paragraph("사용자 관리"); _p.style = _d.styles["Heading 1"]
    _d.add_paragraph("사용자 관리")           # 같은 문장을 본문으로도
    _d.add_paragraph("비밀번호를 초기화합니다.")
    _d.save(str(_hsrc))
    te.translate_document(
        in_path=str(_hsrc), out_path=str(ROOT / "outputs" / "_qa_out.docx"),
        glossary_rows=[{"KO": "사용자", "EN": "user"},
                       {"KO": "존재하지 않는 용어", "EN": "nope"}],
        pattern_rows=[], api_key="sk-stub", enable_cache=False, enable_qa=True,
        translation_mode="매뉴얼", ui_text_overrides={},
    )
finally:
    te.qa_check_batch = _orig_qa
    for _n in ("_qa_src", "_qa_out"):
        (ROOT / "outputs" / f"{_n}.docx").unlink(missing_ok=True)

_srcs = [i[1] for i in _seen.get("items", [])]
check("제목도 QA 대상에 들어감", "사용자 관리" in _srcs, str(_srcs))
check("제목과 본문이 각각 제출됨 (그룹키 분리)",
      _srcs.count("사용자 관리") == 2, str(_srcs))
check("QA 글로서리는 배치에 등장하는 용어만",
      ("사용자", "user") in _seen.get("glossary", []), str(_seen.get("glossary")))


# ══════════════════════════════════════════════════════════════════
print("[12] 댓글 앵커 보존 — 번역되는 문단에 달린 댓글")

W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
_CT = "{http://schemas.openxmlformats.org/package/2006/content-types}"
_PR = "{http://schemas.openxmlformats.org/package/2006/relationships}"


def _docx_with_comment(path, para_text, comment_text, cid="1"):
    """
    번역 대상 문단 하나에 댓글을 단 docx를 만든다.

    python-docx는 댓글 API가 없어 패키지를 직접 손본다. 댓글은
    commentRangeStart/End(문단 안 위치)와 comments.xml(내용)이 따로 있어서,
    문단을 재조립할 때 앵커만 조용히 사라지는 사고가 나기 쉽다.
    """
    import zipfile
    from lxml import etree
    from docx import Document as _Doc

    raw = io.BytesIO()
    d = _Doc()
    d.add_paragraph(para_text)
    d.save(raw)

    zin = zipfile.ZipFile(io.BytesIO(raw.getvalue()))
    parts = {n: zin.read(n) for n in zin.namelist()}

    # 1) 문단에 앵커 삽입
    doc = etree.fromstring(parts["word/document.xml"])
    para = list(doc.iter(W + "p"))[0]
    run = para.find(W + "r")
    start = etree.Element(W + "commentRangeStart"); start.set(W + "id", cid)
    para.insert(list(para).index(run), start)
    end = etree.Element(W + "commentRangeEnd"); end.set(W + "id", cid)
    para.append(end)
    ref_run = etree.SubElement(para, W + "r")
    ref = etree.SubElement(ref_run, W + "commentReference"); ref.set(W + "id", cid)
    parts["word/document.xml"] = etree.tostring(doc, xml_declaration=True,
                                                encoding="UTF-8", standalone=True)

    # 2) comments.xml
    croot = etree.Element(W + "comments", nsmap={"w": W[1:-1]})
    c = etree.SubElement(croot, W + "comment")
    c.set(W + "id", cid); c.set(W + "author", "tester"); c.set(W + "date",
                                                              "2026-01-01T00:00:00Z")
    cp = etree.SubElement(c, W + "p")
    cr = etree.SubElement(cp, W + "r")
    ct = etree.SubElement(cr, W + "t"); ct.text = comment_text
    parts["word/comments.xml"] = etree.tostring(croot, xml_declaration=True,
                                                encoding="UTF-8", standalone=True)

    # 3) 관계 + 콘텐츠 타입 등록
    rels = etree.fromstring(parts["word/_rels/document.xml.rels"])
    r = etree.SubElement(rels, _PR + "Relationship")
    r.set("Id", "rIdComments"); r.set("Target", "comments.xml")
    r.set("Type", "http://schemas.openxmlformats.org/officeDocument/2006/"
                  "relationships/comments")
    parts["word/_rels/document.xml.rels"] = etree.tostring(
        rels, xml_declaration=True, encoding="UTF-8", standalone=True)

    ct_root = etree.fromstring(parts["[Content_Types].xml"])
    ov = etree.SubElement(ct_root, _CT + "Override")
    ov.set("PartName", "/word/comments.xml")
    ov.set("ContentType", "application/vnd.openxmlformats-officedocument."
                          "wordprocessingml.comments+xml")
    parts["[Content_Types].xml"] = etree.tostring(ct_root, xml_declaration=True,
                                                  encoding="UTF-8", standalone=True)

    with zipfile.ZipFile(path, "w", zipfile.ZIP_DEFLATED) as zout:
        for n, data in parts.items():
            zout.writestr(n, data)
    return path


def _comment_shape(path):
    import zipfile
    from lxml import etree
    z = zipfile.ZipFile(path)
    d = etree.fromstring(z.read("word/document.xml"))
    cm = []
    if "word/comments.xml" in z.namelist():
        c = etree.fromstring(z.read("word/comments.xml"))
        cm = [(x.get(W + "id"), "".join(t.text or "" for t in x.iter(W + "t")))
              for x in c.iter(W + "comment")]
    return (sorted(e.get(W + "id") for e in d.iter(W + "commentRangeStart")),
            sorted(e.get(W + "id") for e in d.iter(W + "commentRangeEnd")),
            sorted(e.get(W + "id") for e in d.iter(W + "commentReference")),
            sorted(cm))


_c_in = _docx_with_comment(str(TMP / "_cm_in.docx"),
                           "문서 보안 정책을 설정합니다.", "여기 확인 필요")
_c_out = TMP / "_cm_out.docx"
import stub_llm as _stub
_stub.install()
te.translate_document(
    in_path=_c_in, out_path=str(_c_out),
    glossary_rows=[], pattern_rows=[], api_key="sk-stub",
    enable_cache=False, enable_qa=False, translation_mode="매뉴얼",
    ui_text_overrides={},
)
_before, _after = _comment_shape(_c_in), _comment_shape(str(_c_out))
check("댓글 앵커·내용이 그대로", _before == _after, f"{_before} vs {_after}")
check("본문은 실제로 번역됨",
      not any("\uae30" in t for t in oc.collect(str(_c_out)).para_texts),
      str(oc.collect(str(_c_out)).para_texts))
_f4, _, _ = oc.verify(_c_in, str(_c_out))
check("검증기가 댓글 손실을 보고하지 않음",
      not any("댓글" in x.title for x in _f4), str(_f4))

for _n in ("_cm_in", "_cm_out"):
    (TMP / f"{_n}.docx").unlink(missing_ok=True)


# ══════════════════════════════════════════════════════════════════
print("[13] UI 매핑 → Glossary 순서 — 굵은 자리와 본문이 각자 강제된다")


def _docx_bold_and_plain(path, term, tail_bold, plain_sentence):
    from docx import Document as _Doc
    d = _Doc()
    p1 = d.add_paragraph()
    p1.add_run(term).bold = True
    p1.add_run(tail_bold)
    d.add_paragraph(plain_sentence)
    d.save(path)
    return path


_TERM = "사용자 및 부서 관리"
_o_in = _docx_bold_and_plain(str(TMP / "_ord_in.docx"), _TERM,
                             "를 클릭합니다.", f"{_TERM} 기능을 제공합니다.")
_o_out = TMP / "_ord_out.docx"

import stub_llm as _stub2
_stub2.install()
te.translate_document(
    in_path=_o_in, out_path=str(_o_out),
    glossary_rows=[{"KO": _TERM, "EN": "user and department management"}],
    pattern_rows=[], api_key="sk-stub",
    enable_cache=False, enable_qa=False, translation_mode="매뉴얼",
    ui_text_overrides={_TERM: "User & Department Management"},
)
_paras = [t for t in oc.collect(str(_o_out)).para_texts if t.strip()]
_joined = " || ".join(_paras)
check("굵은 자리는 UI 매핑 표기",
      "User & Department Management" in _paras[0], _joined)
check("본문은 글로서리 표기 — 예전엔 여기서 강제가 사라졌다",
      "user and department management" in _paras[1].lower(), _joined)
check("두 표기가 서로 다른 자리에 각각 적용",
      "User & Department Management" not in _paras[1], _joined)

# 글로서리 항목이 살아 있는지 (전역 삭제가 없어졌는지) 직접 확인
_entries = te.build_glossary_entries_from_rows(
    [{"KO": _TERM, "EN": "user and department management"}]
)
_pre, _map = te.preprocess_with_glossary_placeholders(
    f"{_TERM} 기능", _entries, start_idx=3
)
check("start_idx가 반영됨 (번호 충돌 방지)", "⟦G3⟧" in _pre, _pre)
check("UI가 잡은 자리는 글로서리가 다시 안 건드림",
      te.preprocess_with_glossary_placeholders("⟦G0⟧ 기능", _entries)[1] == {},
      str(te.preprocess_with_glossary_placeholders("⟦G0⟧ 기능", _entries)))

for _n in ("_ord_in", "_ord_out"):
    (TMP / f"{_n}.docx").unlink(missing_ok=True)


# ══════════════════════════════════════════════════════════════════
print("[14] 2차 리포트 — 마커 경계에 공백을 '넣는다'")

_ns2 = te.normalize_marker_boundary_spaces
for _name, _in, _want in [
    ("websiteGo", "⟦H0⟧Fireside admin website⟦/H0⟧Go to the website.",
     "⟦H0⟧Fireside admin website⟦/H0⟧ Go to the website."),
    ("syncClick", "schema you want to sync⟦D3⟧Click it.",
     "schema you want to sync ⟦D3⟧ Click it."),
    ("optionsto", "More options⟦D0⟧to the right",
     "More options ⟦D0⟧ to the right"),
    ("toEdit", "table name to⟦B⟧Edit the column⟦/B⟧",
     "table name to ⟦B⟧Edit the column⟦/B⟧"),
    ("a.csv", "using a⟦X0⟧ file", "using a ⟦X0⟧ file"),
]:
    check(f"{_name} 방지", _ns2(_in) == _want, f"{_ns2(_in)!r}")
check("구두점 앞에는 안 넣는다", _ns2("⟦B⟧Save⟦/B⟧.") == "⟦B⟧Save⟦/B⟧.", _ns2("⟦B⟧Save⟦/B⟧."))
check("이미 공백 있으면 겹치지 않는다",
      _ns2("⟦B⟧Save⟦/B⟧ and close") == "⟦B⟧Save⟦/B⟧ and close",
      _ns2("⟦B⟧Save⟦/B⟧ and close"))

print("[15] 2차 리포트 — nan 차단")
check("pandas NaN", te._clean(float("nan")) == "")
check("문자열 nan", te._clean("nan") == "" and te._clean("NaN") == "")
check("none/null", te._clean("None") == "" and te._clean("null") == "")
check("정상 값은 유지", te._clean("Save") == "Save")

_nan_src = _docx([("상태를 확인합니다.", None, None)], str(TMP / "_nan_src.docx"))
_nan_out = TMP / "_nan_out.docx"
import stub_llm as _stub3
_stub3.install()
te.translate_document(
    in_path=_nan_src, out_path=str(_nan_out),
    glossary_rows=[], pattern_rows=[], api_key="sk-stub",
    enable_cache=False, enable_qa=False, translation_mode="매뉴얼",
    ui_text_overrides={"상태": float("nan"), float("nan"): "Status"},
)
check("NaN이 든 UI 매핑이 본문을 오염시키지 않음",
      "nan" not in " ".join(oc.collect(str(_nan_out)).para_texts).lower(),
      str(oc.collect(str(_nan_out)).para_texts))

print("[16] 2차 리포트 — 조립된 문장 검사")
_asm = _docx([
    ("nan can remove all messages.", None, None),
    ("Fireside admin websiteGo to the website.", None, None),
    ("Enter the appKey and serverId.", None, None),
], str(TMP / "_asm.docx"))
_asm_src = _docx([
    ("전사 관리자는 메시지를 삭제할 수 있습니다.", None, None),
    ("Fireside 관리자 웹사이트에 접속합니다.", None, None),
    ("appKey와 serverId를 입력합니다.", None, None),
], str(TMP / "_asm_src.docx"))
_f5, _, _ = oc.verify(_asm_src, _asm)
check("nan 노출 감지", has(_f5, "결측값이 본문에"), str(_f5))
check("결합 문자열 감지", has(_f5, "낱말이 붙어"), str(_f5))
check("camelCase 식별자는 오탐 아님", "appKey" not in str(_f5), str(_f5))

print("[17] 2차 리포트 — 같은 영문으로 매핑된 서로 다른 원문")
_dups = oc.check_duplicate_targets(
    [("상태", "Status"), ("사용 여부", "Status"), ("설명", "Description")])
check("중복 매핑 검출", len(_dups) == 1 and set(_dups[0][1]) == {"상태", "사용 여부"},
      str(_dups))
check("고유 매핑은 통과", oc.check_duplicate_targets([("상태", "Status")]) == [])
check("빈 값은 무시", oc.check_duplicate_targets([("", "X"), ("가", "")]) == [])

for _n in ("_nan_src", "_nan_out", "_asm", "_asm_src"):
    (TMP / f"{_n}.docx").unlink(missing_ok=True)


print()
if failures:
    print(f"FAILED {len(failures)}건: {failures}")
    raise SystemExit(1)
print("ALL PASS")
