import copy
import os
import re
import tempfile
import zipfile
import zlib
from collections import Counter, defaultdict
from dataclasses import dataclass
from typing import List, Dict, Tuple, Iterable, Optional, Callable, Any

from docx import Document
from lxml import etree
from openai import OpenAI

import markdown_format

# OOXML namespace constants
_W        = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"
_XML_SPACE = "http://www.w3.org/XML/1998/namespace"


TOTAL_INPUT_TOKENS = 0
TOTAL_CACHED_INPUT_TOKENS = 0
TOTAL_OUTPUT_TOKENS = 0
TOTAL_TOKENS = 0

B_OPEN = "⟦B⟧"
B_CLOSE = "⟦/B⟧"
# soft line break marker — Word `<w:br/>` element.
# **Must NOT start with the letter B** to avoid the LLM confusing it with the
# bold-open ⟦B⟧ marker and splitting it into "⟦B⟧R⟧" style garbage. ⟦LB⟧
# ("line break") is far enough away visually.
BR_MARKER = "⟦LB⟧"
# tab marker — Word `<w:tab/>` element inside a run. Same naming constraint as
# BR_MARKER: must not collide visually with ⟦B⟧.
TAB_MARKER = "⟦TB⟧"

G_PREFIX = "⟦G"
D_PREFIX = "⟦D"
# 원문 영문 리터럴 봉인용. 'L'은 ⟦LB⟧와 눈으로 헷갈리므로 X를 쓴다
# (BR_MARKER가 B를 피한 것과 같은 이유).
X_PREFIX = "⟦X"
SUFFIX = "⟧"

KOREAN_RE = re.compile(r"[가-힣]")
PLACEHOLDER_RE = re.compile(r"⟦G\d+⟧")
LITERAL_PH_RE = re.compile(r"⟦X\d+⟧")
DRAWING_PH_RE = re.compile(r"⟦D\d+⟧")
H_OPEN_RE    = re.compile(r"⟦H(\d+)⟧")
H_CLOSE_RE   = re.compile(r"⟦/H(\d+)⟧")
HL_OPEN_RE   = re.compile(r"⟦HL:([a-zA-Z]+)⟧")
HL_CLOSE     = "⟦/HL⟧"
BR_MARKER_RE = re.compile(re.escape(BR_MARKER))
ALL_MARKER_RE = re.compile(r"(⟦B⟧|⟦/B⟧|⟦I⟧|⟦/I⟧|⟦C\d+⟧|⟦D\d+⟧|⟦H\d+⟧|⟦/H\d+⟧|⟦HL:[a-zA-Z]+⟧|⟦/HL⟧|⟦LB⟧|⟦TB⟧|⟦X\d+⟧)")
# ⟦LB⟧ / ⟦TB⟧ 는 run 내부에서 <w:br/> · <w:tab/> 요소로 되살려야 하므로
# _make_run에서 텍스트를 이 마커 기준으로 쪼갠다.
LAYOUT_MARKER_SPLIT_RE = re.compile(
    rf"({re.escape(BR_MARKER)}|{re.escape(TAB_MARKER)})"
)
MARKER_SPLIT_RE = re.compile(rf"({re.escape(B_OPEN)}|{re.escape(B_CLOSE)}|⟦G\d+⟧)")
# 대소문자 정규화(heading sentence case, UI label 등)를 지날 때 **모든** 마커는
# 원형 그대로 남아야 한다. ⟦B⟧/⟦G n⟧만 보호하면 ⟦HL:yellow⟧의 "HL"이 첫 알파벳으로
# 잡혀 "⟦Hl:yellow⟧ / ⟦/hl⟧"로 훼손되고, 이후 ALL_MARKER_RE 토크나이저가 인식하지
# 못해 마커가 문서에 그대로 찍힌다.
ANY_MARKER_SPLIT_RE = re.compile(
    r"(⟦B⟧|⟦/B⟧|⟦I⟧|⟦/I⟧|⟦C\d+⟧|⟦D\d+⟧|⟦G\d+⟧|⟦H\d+⟧|⟦/H\d+⟧|⟦HL:[a-zA-Z]+⟧|⟦/HL⟧|⟦LB⟧|⟦TB⟧|⟦X\d+⟧)"
)

# 안전장치용 정규식 — LLM이 ⟦LB⟧를 파괴적으로 응답해서 "⟦L⟧B⟧" 같은 잔해나
# 이전 ⟦BR⟧ 잔해가 남을 때 감지·복구. `_make_run` 직전에 실행.
_LB_REMNANT_RE = re.compile(r"⟦L\s*⟧\s*B\s*⟧|⟦B\s*R\s*⟧|⟦BR⟧|⟦B\s*⟧R\s*⟧|R\s*⟧(?=\s*[A-Z가-힣])", re.IGNORECASE)

UI_LOWER_WORDS = {
    "name",
    "names",
    "list",
    "lists",
    "detail",
    "details",
    "setting",
    "settings",
    "information",
    "field",
    "fields",
    "filter",
    "filters",
    "menu",
    "menus",
    "tab",
    "tabs",
    "status",
    "type",
    "types",
    "history",
    "option",
    "options",
    "message",
    "messages",
    "group",
    "groups",
    "owner",
    "owners",
    "user",
    "users",
    "rule",
    "rules",
    "policy",
    "policies",
    "log",
    "logs",
    "guideline",
    "guidelines",
    "pattern",
    "patterns",
    "tag",
    "tags",
    "value",
    "values",
    "result",
    "results",
    "data",
    "items",
    "item",
    "dialog",
    "window",
    "button",
    "buttons",
    "criteria",
    "class",
    "level",
}


@dataclass(frozen=True)
class GlossaryEntry:
    ko: str
    en: str
    dnt: bool
    case_sensitive: bool
    product: str
    note: str


def reset_token_counters():
    global TOTAL_INPUT_TOKENS, TOTAL_CACHED_INPUT_TOKENS, TOTAL_OUTPUT_TOKENS, TOTAL_TOKENS
    TOTAL_INPUT_TOKENS = 0
    TOTAL_CACHED_INPUT_TOKENS = 0
    TOTAL_OUTPUT_TOKENS = 0
    TOTAL_TOKENS = 0


# 폭 없는(zero-width) 문자들. 모델 응답에 이따금 섞여 들어오는데, 렌더링에
# 보이지 않아 사람 리뷰로는 절대 걸러지지 않고 검색·복사·diff에서만 문제가
# 된다. 한국어→영어 번역에서는 이 문자들이 의미를 갖는 경우가 없으므로 제거해도
# 안전하다. (U+00A0 non-breaking space는 원문 의도일 수 있어 건드리지 않는다.)
_ZERO_WIDTH_RE = re.compile(
    "["
    "﻿"   # ZERO WIDTH NO-BREAK SPACE (BOM) — 실제로 관측된 케이스
    "​"   # ZERO WIDTH SPACE
    "‌"   # ZERO WIDTH NON-JOINER
    "‍"   # ZERO WIDTH JOINER
    "⁠"   # WORD JOINER
    "­"   # SOFT HYPHEN
    "]"
)


def strip_zero_width(text: str) -> str:
    return _ZERO_WIDTH_RE.sub("", text)


def contains_korean(text: str) -> bool:
    # Strip drawing placeholders before checking — they are not translatable text
    clean = DRAWING_PH_RE.sub("", text) if text else text
    return bool(clean and KOREAN_RE.search(clean))


def make_marker(prefix: str, i: int) -> str:
    return f"{prefix}{i}{SUFFIX}"


def _to_bool(v: Any) -> bool:
    if v is None:
        return False
    s = str(v).strip().lower()
    return s in {"true", "y", "yes", "1"}


# 유효한 값처럼 보이지만 실은 결측값이 문자열로 새어 나온 것들.
# pandas의 NaN은 **참(truthy)** 이라 `v or ""` 를 그냥 통과한다.
_NULLISH = {"nan", "none", "null", "<na>", "nat"}


def _clean(v: Any) -> str:
    if v is None:
        return ""
    try:
        if v != v:                      # NaN은 자기 자신과 같지 않다
            return ""
    except Exception:
        pass
    s = str(v).strip()
    if s.lower() in _NULLISH:
        return ""
    return s


def _cap_first_alpha(s: str) -> str:
    """Capitalise the first alphabetic character of the *visible text*.

    마커(⟦HL:yellow⟧, ⟦B⟧ …) 내부 글자는 건너뛴다 — 그렇지 않으면 마커로 시작하는
    문단에서 마커의 첫 글자만 손대고 정작 본문 첫 글자는 소문자로 남는다.
    """
    parts = ANY_MARKER_SPLIT_RE.split(s)
    for pi, part in enumerate(parts):
        if not part or ANY_MARKER_SPLIT_RE.fullmatch(part):
            continue
        for i, ch in enumerate(part):
            if ch.isalpha():
                parts[pi] = part[:i] + ch.upper() + part[i + 1:]
                return "".join(parts)
    return s


def build_glossary_entries_from_rows(rows: List[dict]) -> List[GlossaryEntry]:
    entries: List[GlossaryEntry] = []

    for r in rows:
        ko = _clean(r.get("KO"))
        en = _clean(r.get("EN"))
        if not ko or not en:
            continue

        entries.append(
            GlossaryEntry(
                ko=ko,
                en=en,
                dnt=_to_bool(r.get("DNT")),
                case_sensitive=_to_bool(r.get("Case-sensitive")),
                product=_clean(r.get("Product")),
                note=_clean(r.get("Note")),
            )
        )

    deduped: List[GlossaryEntry] = []
    seen = set()
    for e in entries:
        key = (e.ko, e.en, e.product, e.note, e.dnt, e.case_sensitive)
        if key not in seen:
            seen.add(key)
            deduped.append(e)

    deduped.sort(key=lambda e: len(e.ko), reverse=True)
    return deduped


def build_pattern_pairs_from_rows(rows: List[dict]) -> List[Tuple[str, str]]:
    pairs: List[Tuple[str, str]] = []
    seen = set()

    for r in rows:
        ko = _clean(r.get("KO"))
        en = _clean(r.get("EN"))
        if not ko or not en:
            continue

        key = (ko, en)
        if key not in seen:
            seen.add(key)
            pairs.append((ko, en))

    return pairs


def preprocess_with_glossary_placeholders(
    text: str,
    entries: List[GlossaryEntry],
    start_idx: int = 0,
) -> Tuple[str, Dict[str, GlossaryEntry]]:
    """
    Replace glossary KO terms in *text* with ⟦G0⟧ … placeholders.

    Uses a regex that normalises internal whitespace (\s+) so that terms
    like '민감 정보' still match even when the paragraph runs were merged
    with a different number of spaces (or a zero-width space, thin space, etc.).
    Longer KO terms are tried first (entries are pre-sorted by length desc)
    to avoid a short term swallowing part of a longer one.

    start_idx는 앞선 전처리(UI 매핑)가 이미 ⟦G0⟧…를 쓴 경우 번호가 겹치지
    않게 이어 붙이기 위한 것이다.
    """
    out = text
    mapping: Dict[str, GlossaryEntry] = {}
    idx = start_idx

    for entry in entries:
        if not entry.ko:
            continue

        # Build a pattern that matches the KO term with flexible internal whitespace.
        # Each space in the glossary key becomes \s+ so ' 민감 정보 ' matches
        # '민감정보', '민감  정보', etc.
        escaped_parts = [re.escape(part) for part in entry.ko.split()]
        pattern = r"\s*".join(escaped_parts) if escaped_parts else re.escape(entry.ko)

        # Quick pre-check: at least the first word must be present
        if escaped_parts and not re.search(escaped_parts[0], out):
            continue

        if re.search(pattern, out):
            ph = make_marker(G_PREFIX, idx)
            idx += 1
            out = re.sub(pattern, ph, out)
            mapping[ph] = entry

    return out, mapping


def preprocess_ui_overrides_in_bold(
    text: str,
    ui_overrides: Dict[str, str],
    start_idx: int = 0,
) -> Tuple[str, Dict[str, GlossaryEntry], int]:
    """
    Replace UI text-mapping KO terms with ⟦G#⟧ placeholders **only within
    ⟦B⟧…⟦/B⟧ (bold) segments**.

    UI text mappings are entered by the user to force specific EN wording
    for on-screen labels — which are always rendered as bold in the source
    docx. Applying the same replacement to non-bold occurrences (regular
    prose that happens to contain the same Korean word) produces awkward
    output like "Change the file Retention policy or Delete them", so the
    substitution is confined to bold spans.

    `start_idx` lets callers continue an existing G-index sequence (e.g.
    after a preceding call to `preprocess_with_glossary_placeholders`) so
    the resulting placeholders don't collide.

    Returns (out, mapping, next_idx).
    """
    if not ui_overrides:
        return text, {}, start_idx

    mapping: Dict[str, GlossaryEntry] = {}
    idx = start_idx
    # 긴 term부터 시도 — 짧은 term이 긴 term 안에서 매치되어 부분 치환되는 것 방지
    sorted_terms = sorted(
        [(ko, en) for ko, en in ui_overrides.items() if ko and en],
        key=lambda kv: len(kv[0]),
        reverse=True,
    )

    def _replace_inside_bold(match):
        nonlocal idx
        inner = match.group(1)
        for ko, en in sorted_terms:
            escaped_parts = [re.escape(part) for part in ko.split()]
            pattern = r"\s*".join(escaped_parts) if escaped_parts else re.escape(ko)
            if re.search(pattern, inner):
                ph = make_marker(G_PREFIX, idx)
                idx += 1
                inner = re.sub(pattern, ph, inner)
                mapping[ph] = GlossaryEntry(
                    ko=ko.strip(),
                    en=en.strip(),
                    dnt=False,
                    case_sensitive=True,
                    product="",
                    note="UI text override (bold only)",
                )
        return f"{B_OPEN}{inner}{B_CLOSE}"

    out = re.sub(rf"{re.escape(B_OPEN)}(.+?){re.escape(B_CLOSE)}", _replace_inside_bold, text, flags=re.DOTALL)
    return out, mapping, idx


def _is_at_sentence_or_bold_start(text_before: str) -> bool:
    """
    Return True if the position immediately after text_before is the start of
    a sentence or the start of a bold segment — i.e. the restored term should
    be capitalized.

    Rules:
    - Nothing before → start of text → capitalize.
    - Last non-space character is sentence-ending punctuation (.  !  ?) → capitalize.
    - text_before ends with the bold-open marker ⟦B⟧ (after stripping trailing
      spaces) → we are the first word inside a bold segment → capitalize.
    """
    s = text_before.rstrip()
    if not s:
        return True
    if s[-1] in ".!?":
        return True
    if s.endswith(B_OPEN):
        return True
    return False


# ──────────────────────────────────────────────────────────────────────
# 원문 영문 리터럴 봉인
#
# 프롬프트에 "기존 영문은 바꾸지 마라"라고 써도 모델은 지키지 않는다.
# 실제 산출물에서 K-Assistant가 AI Assistant로 바뀌었고, "Enable basic AI
# features"가 "…basic ai features"가 됐다. 뒤엣것은 모델이 아니라 우리
# 후처리(_sentence_case_preserving_markers)가 저지른 것이다 — 첫 글자 뒤의
# 알파벳을 전부 소문자로 내리기 때문이다.
#
# 지시로 막을 수 없으면 **구조로** 막는다. 번역 전에 자리표시자로 바꿔두면
# 모델도 후처리도 건드릴 수 없다. 복원은 대소문자 정규화까지 모두 끝난 뒤.
# ──────────────────────────────────────────────────────────────────────

# 경계에 \b를 쓰면 안 된다. 파이썬 정규식의 \w는 한글을 포함하므로
# "K-Assistant는" 에서 t와 는 사이에 경계가 생기지 않아 매칭이 통째로 실패한다.
# 실제로 리포트가 문제 삼은 K-Assistant가 바로 이 형태였다.
# 그래서 "영문·숫자가 아니면 경계"로 직접 정의한다.
_LB = r"(?<![A-Za-z0-9])"
_RB = r"(?![A-Za-z0-9])"
_LITERAL_RE = re.compile(
    r"https?://[^\s⟦⟧]+"                                     # URL
    r"|[\w.+-]+@[\w-]+\.[\w.-]+"                             # 이메일
    rf"|{_LB}v\d+(?:\.\d+)+{_RB}"                            # v2.1, v1.7.1
    rf"|{_LB}[A-Za-z][A-Za-z0-9+#_-]*\.(?:docx?|xlsx?|pptx?|pdf|txt|csv"
    rf"|json|xml|zip|png|jpe?g|gif|svg|mdx?){_RB}"           # 파일명
    rf"|{_LB}[A-Z][A-Za-z]*-[A-Z][A-Za-z0-9]+{_RB}"           # K-Assistant, MS-SQL
    rf"|{_LB}(?:SaaS|PaaS|IaaS|IoT|iOS|macOS){_RB}"           # 혼합 대소문자 관용 표기
    rf"|{_LB}[A-Z][A-Z0-9]{{1,7}}{_RB}"                       # AI, API, SQL, JDK…
)


def seal_literals(text: str) -> Tuple[str, Dict[str, str]]:
    """
    번역·후처리가 건드리면 안 되는 원문 영문을 ⟦X#⟧로 봉인한다.

    마커 안쪽은 절대 건드리지 않는다 — ⟦HL:yellow⟧의 "HL"이 대문자 두 글자라
    약어 규칙에 걸리기 때문이다. 그래서 마커 기준으로 쪼갠 뒤 사이 구간만 본다.
    """
    if not text:
        return text, {}
    mapping: Dict[str, str] = {}
    parts = _split_preserving_markers(text)
    out = []
    for part in parts:
        if not part or _is_marker_token(part):
            out.append(part)
            continue

        def _seal(m):
            ph = make_marker(X_PREFIX, len(mapping))
            mapping[ph] = m.group(0)
            return ph

        out.append(_LITERAL_RE.sub(_seal, part))
    return "".join(out), mapping


def restore_literals(text: str, mapping: Dict[str, str]) -> str:
    """봉인한 영문을 원래 표기 그대로 되돌린다."""
    if not mapping or not text:
        return text
    # 자리표시자가 통째로 사라진 경우(모델이 삼킴)는 여기서 되살릴 수 없다.
    # check_marker_integrity가 ⟦X#⟧ 소실로 잡아준다.
    for ph, original in mapping.items():
        text = text.replace(ph, original)
    return text


def enforce_literal_casing(text: str, source_text: str) -> str:
    """
    QA(Pass 2)가 흐트러뜨린 원문 영문의 대소문자를 되돌린다.

    Pass 2는 봉인 바깥에서 돌기 때문에 여기서 한 번 더 지켜야 한다.
    치환이 아니라 **대소문자만** 맞춘다 — 단어 자체가 바뀐 경우는 QA의
    판단일 수 있으므로 건드리지 않는다.
    """
    if not text or not source_text:
        return text
    for lit in dict.fromkeys(_LITERAL_RE.findall(source_text)):
        if len(lit) < 2:
            continue
        text = re.sub(rf"{_LB}{re.escape(lit)}{_RB}", lit, text,
                      flags=re.IGNORECASE)
    return text


def _lower_phrase(en: str) -> str:
    """
    문장 중간에 놓일 용어를 소문자로. **첫 글자만이 아니라 구 전체**를 본다.

    첫 글자만 내리면 "Agent Feedback"이 "agent Feedback"이 되어 어중간하다.
    다만 약어(SQL, ID, AI)와 혼합 표기(K-Assistant, SaaS)는 그대로 둔다 —
    이들은 문장 어디에 오든 표기가 고정이다.
    """
    if not en:
        return en
    words = []
    for w in en.split(" "):
        core = w.strip(".,;:()[]")
        if core.isupper() or any(c.isupper() for c in core[1:]):
            words.append(w)                       # 약어·혼합 표기는 유지
        else:
            words.append(w[:1].lower() + w[1:])
    return " ".join(words)


def restore_glossary_placeholders(
    text: str,
    mapping: Dict[str, GlossaryEntry],
) -> str:
    """
    Restore each glossary placeholder with position-aware capitalisation.

    - DNT or Case-sensitive entries: always inserted exactly as stored in EN.
    - Normal entries (lowercase EN in glossary):
        • Capitalise the first letter when the placeholder sits at the very
          start of the text, immediately after sentence-ending punctuation,
          or immediately after the bold-open marker ⟦B⟧.
        • Lower-case otherwise.

    This means glossary EN values should always be stored in lowercase.
    The function handles capitalisation automatically based on context.
    """
    out = text

    for ph, entry in mapping.items():
        # Fixed casing: DNT or explicitly case-sensitive terms are never touched.
        if entry.dnt or entry.case_sensitive:
            out = out.replace(ph, entry.en)
            continue

        # Variable casing: process each occurrence independently so we can
        # inspect what precedes that specific occurrence.
        result = ""
        remaining = out
        while True:
            idx = remaining.find(ph)
            if idx == -1:
                result += remaining
                break

            # Everything accumulated so far + text up to this placeholder
            # gives us the "before" context for position detection.
            before = result + remaining[:idx]
            at_start = _is_at_sentence_or_bold_start(before)

            en = entry.en
            if en:
                # 어느 위치든 구 전체를 먼저 소문자 관례로 맞춘 뒤, 문장
                # 처음이면 첫 글자만 올린다. 그래야 같은 용어가 한 문서에서
                # "Agent Feedback"과 "agent feedback"으로 갈리지 않는다.
                en = _lower_phrase(en)
                if at_start:
                    en = en[0].upper() + en[1:]

            result += remaining[:idx] + en
            remaining = remaining[idx + len(ph):]

        out = result

    out = re.sub(r"\s+([.,;:!?])", r"\1", out)
    out = re.sub(r"[ \t]{2,}", " ", out)
    return out.strip()


def enforce_case_sensitive_glossary(
    text: str,
    glossary_entries: List[GlossaryEntry],
) -> str:
    """
    Restore the exact casing of DNT and case-sensitive glossary EN values.

    Some downstream steps — heading sentence-case ([_sentence_case_preserving_markers]),
    QA revisions, generic capitalisation helpers — can drift the casing of
    product names like 'Wrapsody eCo' to 'Wrapsody eco'. This function does a
    case-insensitive sweep and rewrites every occurrence with the stored EN
    value, so the user's case_sensitive=True flag actually holds end-to-end.
    Word boundaries are used to avoid partial matches inside other words.
    """
    for entry in glossary_entries:
        if not (entry.dnt or entry.case_sensitive):
            continue
        if not entry.en:
            continue
        pattern = re.compile(rf"\b{re.escape(entry.en)}\b", re.IGNORECASE)
        text = pattern.sub(entry.en, text)
    return text


def _lxml_all_runs(p_elem):
    """Return every w:r in the paragraph in document order,
    including those nested inside w:hyperlink."""
    return p_elem.findall(f".//{{{_W}}}r")


def _run_is_comment_ref(r_elem) -> bool:
    """True for the w:r that holds w:commentReference (the comment bubble icon).
    This run must NEVER be deleted or rebuilt."""
    return r_elem.find(f"{{{_W}}}commentReference") is not None


def _run_highlight_val(r_elem) -> Optional[str]:
    """Return the highlight colour string (e.g. 'yellow') or None."""
    rPr = r_elem.find(f"{{{_W}}}rPr")
    if rPr is None:
        return None
    hl = rPr.find(f"{{{_W}}}highlight")
    return hl.get(f"{{{_W}}}val") if hl is not None else None


def _lxml_run_is_bold(r_elem) -> bool:
    """True when the run carries an active w:b element."""
    rPr = r_elem.find(f"{{{_W}}}rPr")
    if rPr is None:
        return False
    b = rPr.find(f"{{{_W}}}b")
    if b is None:
        return False
    val = b.get(f"{{{_W}}}val")
    return val not in ("false", "0", "False")


def _lxml_text_elem(r_elem):
    """Return the w:t child element of a run, or None."""
    return r_elem.find(f"{{{_W}}}t")


def _run_has_writable_content(r_elem) -> bool:
    """
    True when the run holds content the rewrite step regenerates itself —
    text (<w:t>), a soft line break (<w:br/>), or a tab (<w:tab/>).

    Runs holding *only* a break or a tab used to be invisible to the rebuild
    logic (it looked for <w:t> alone), so they survived while the same break
    was also re-emitted from the ⟦LB⟧ marker — silently doubling every blank
    line in the output.
    """
    return any(
        r_elem.find(f"{{{_W}}}{tag}") is not None
        for tag in ("t", "br", "tab")
    )


def _run_is_non_text(r_elem) -> bool:
    """True when the run contains a drawing, symbol, or embedded object
    (i.e. an icon) rather than translatable text."""
    return any(
        r_elem.find(f"{{{_W}}}{tag}") is not None
        for tag in ("drawing", "sym", "object", "pict")
    )


def _parse_marked_segments(marked: str) -> List[Tuple]:
    """Split a marked string into typed segments.

    Returns list of (type, text) where type is one of:
      {'bold': bool, 'hl': str|None}  — plain text with formatting
      'drawing'                         — drawing placeholder
      ('h', int)                        — text inside hyperlink <int>
    """
    segments: List[Tuple] = []
    tokens = ALL_MARKER_RE.split(marked)
    bold   = False
    in_hl:    Optional[str] = None
    buf: List[str] = []

    in_hlink: Optional[int] = None

    def _flush() -> None:
        if not buf:
            return
        text = "".join(buf)
        if in_hlink is not None:
            segments.append((("h", in_hlink), text))
        else:
            segments.append(({"bold": bold, "hl": in_hl}, text))
        buf.clear()

    for tok in tokens:
        if tok == B_OPEN:
            _flush(); bold = True
        elif tok == B_CLOSE:
            _flush(); bold = False
        elif DRAWING_PH_RE.fullmatch(tok):
            _flush(); segments.append(("drawing", tok))
        elif H_OPEN_RE.fullmatch(tok):
            _flush(); in_hlink = int(H_OPEN_RE.match(tok).group(1))
        elif H_CLOSE_RE.fullmatch(tok):
            _flush(); in_hlink = None
        elif HL_OPEN_RE.fullmatch(tok):
            _flush(); in_hl = HL_OPEN_RE.match(tok).group(1)
        elif tok == HL_CLOSE:
            _flush(); in_hl = None
        elif tok:
            buf.append(tok)

    _flush()
    return [(b, t) for b, t in segments if t]


def paragraph_to_marked_text(paragraph) -> Tuple[str, Dict, Dict, int]:
    """
    Extract paragraph text with:
    - ⟦B⟧…⟦/B⟧  bold markers
    - ⟦D0⟧        inline drawing placeholders
    - ⟦H0⟧…⟦/H0⟧ hyperlink span markers (so the LLM can keep link text together)
    - ⟦LB⟧        soft line break (<w:br/>)
    - ⟦TB⟧        tab (<w:tab/>)

    Walks *direct children* of w:p so hyperlink nodes are handled as a unit
    rather than mixing their inner runs with regular paragraph runs.

    Returns
    -------
    marked_text   : str
    drawing_map   : {placeholder_str: lxml_run_element}
    hyperlink_map : {index: {'elem': w:hyperlink_element, 'runs': [run_elements]}}
    trailing_lb   : how many soft line breaks the source had at the very end
                    (stripped from marked_text, re-applied verbatim on write-back
                    so the author's intentional blank line survives)
    """
    p_elem = paragraph._p
    parts: List[str] = []
    in_bold = False
    in_hl: Optional[str] = None
    d_idx = 0
    h_idx = 0
    drawing_map: Dict[str, Any] = {}
    hyperlink_map: Dict[int, Any] = {}

    for child in p_elem:
        ctag = child.tag.split("}")[1] if "}" in child.tag else child.tag

        if ctag == "r":
            if _run_is_comment_ref(child):
                continue
            if _run_is_non_text(child):
                if in_bold: parts.append(B_CLOSE); in_bold = False
                if in_hl:   parts.append(HL_CLOSE); in_hl = None
                ph = f"{D_PREFIX}{d_idx}{SUFFIX}"
                parts.append(ph); drawing_map[ph] = child; d_idx += 1
                continue
            # 한 run 안의 자식 요소들을 문서 순서대로 훑는다 — <w:t> 외에
            # <w:br/> (soft line break) 와 <w:tab/> 도 만나면 ⟦LB⟧ / ⟦TB⟧
            # 마커로 편입시켜, 원문의 줄바꿈·탭 위치가 번역·재쓰기 단계까지
            # 살아남게 한다.
            text_parts: List[str] = []
            for sub in child:
                sub_tag = sub.tag.split("}")[1] if "}" in sub.tag else sub.tag
                if sub_tag == "t":
                    text_parts.append(sub.text or "")
                elif sub_tag == "br":
                    text_parts.append(BR_MARKER)
                elif sub_tag == "tab":
                    text_parts.append(TAB_MARKER)
            if not text_parts:
                continue
            text = "".join(text_parts)
            if not text:
                continue
            is_bold = _lxml_run_is_bold(child)
            hl_val  = _run_highlight_val(child)
            if in_hl is not None and hl_val != in_hl:
                if in_bold:
                    parts.append(B_CLOSE); in_bold = False
                parts.append(HL_CLOSE); in_hl = None
            if hl_val is not None and in_hl is None:
                parts.append(f"⟦HL:{hl_val}⟧"); in_hl = hl_val
            if is_bold and not in_bold:
                parts.append(B_OPEN); in_bold = True
            elif not is_bold and in_bold:
                parts.append(B_CLOSE); in_bold = False
            parts.append(text)

        elif ctag == "hyperlink":
            if in_bold: parts.append(B_CLOSE); in_bold = False
            if in_hl:   parts.append(HL_CLOSE); in_hl = None
            hl_runs = child.findall(f"{{{_W}}}r")
            hl_text = "".join(
                r.find(f"{{{_W}}}t").text or ""
                for r in hl_runs
                if r.find(f"{{{_W}}}t") is not None
            )
            if hl_text:
                parts.append(f"⟦H{h_idx}⟧"); parts.append(hl_text); parts.append(f"⟦/H{h_idx}⟧")
                hyperlink_map[h_idx] = {"elem": child, "runs": hl_runs}
                h_idx += 1

    if in_bold: parts.append(B_CLOSE)
    if in_hl:   parts.append(HL_CLOSE)

    marked = "".join(parts)
    # 문단 끝의 ⟦LB⟧는 번역 대상 텍스트에서 떼어낸다 — LLM에게 넘기면 위치가
    # 흐트러지고, QA가 없던 걸 덧붙이기도 하기 때문. 다만 **버리지는 않고**
    # 개수를 세어 두었다가 write-back 시 그대로 복원한다. 작성자가 의도적으로
    # 넣은 빈 줄(= 단락 사이 시각적 간격)이 사라지지 않게 하기 위함.
    m_tail = re.search(rf"((?:{re.escape(BR_MARKER)}\s*)+)$", marked)
    trailing_lb = 0
    if m_tail:
        trailing_lb = m_tail.group(1).count(BR_MARKER)
        marked = marked[: m_tail.start()]

    return marked, drawing_map, hyperlink_map, trailing_lb


_BOLD_SEGMENT_RE = re.compile(r"⟦B⟧(.+?)⟦/B⟧", re.DOTALL)
_INNER_MARKER_RE = re.compile(r"⟦[A-Za-z/]+(?::[A-Za-z]+)?\d*⟧")


# 1x1 transparent PNG — sanitize_docx 폴백 placeholder.
# (70 bytes; PIL로 생성하고 verify 통과한 valid PNG)
_FALLBACK_PNG = bytes.fromhex(
    "89504E470D0A1A0A0000000D49484452000000010000000108060000001F15C489"
    "0000000D49444154789C6360606060000000050001A5F645400000000049454E44"
    "AE426082"
)

_IMAGE_EXTS = (".png", ".jpg", ".jpeg", ".gif", ".bmp", ".tif", ".tiff", ".svg", ".webp", ".emf", ".wmf")


def sanitize_docx(in_path: str) -> str:
    """
    Return a path to a docx whose zip entries all pass CRC-32 verification.

    .docx is a zip container; pipelines like Wrapsody decrypt occasionally
    leave one embedded image with a CRC mismatch, and python-docx then refuses
    to open the whole file. Strategy:

    1. Probe every entry. If all read cleanly → return the original path.
    2. Otherwise rebuild the zip:
       - readable entries → copied as-is
       - broken **image** entries → replaced with a valid placeholder PNG
         (preferring another valid image from the same docx for visual
         consistency, falling back to a 1x1 transparent PNG). The filename
         and zip path stay the same so the rels/content-types references
         that point to it keep working.
       - broken **non-image** entries → re-raised, because dropping or
         emptying word/document.xml or a rels file would silently corrupt
         the document.
    """
    # ── 1차 probe: 어떤 entry가 깨졌는지 + valid image bytes 한 개 확보 ─
    bad_entries: list[str] = []
    valid_image_bytes: Optional[bytes] = None
    try:
        with zipfile.ZipFile(in_path, "r") as z:
            for zinfo in z.infolist():
                try:
                    data = z.read(zinfo.filename)
                    lname = zinfo.filename.lower()
                    if valid_image_bytes is None and lname.endswith(_IMAGE_EXTS):
                        valid_image_bytes = data
                except (zipfile.BadZipFile, zlib.error, RuntimeError):
                    bad_entries.append(zinfo.filename)
    except (zipfile.BadZipFile, zlib.error, RuntimeError):
        # zip 자체가 깨졌으면 손쓸 도리 없음 — 원본 그대로 넘김
        return in_path

    if not bad_entries:
        return in_path

    placeholder = valid_image_bytes or _FALLBACK_PNG

    # 깨진 entry 중 이미지가 아닌 것이 하나라도 있으면 안전하게 복구할 수 없음
    non_image_bad = [n for n in bad_entries if not n.lower().endswith(_IMAGE_EXTS)]
    if non_image_bad:
        raise RuntimeError(
            "docx 내부 파일이 손상되었습니다 (이미지 외 part): "
            + ", ".join(non_image_bad[:3])
            + " — Wrapsody에서 decrypt를 다시 시도해 주세요."
        )

    # ── 새 zip으로 rebuild ─────────────────────────────────────────
    fd, dst = tempfile.mkstemp(suffix=".docx", prefix="repaired_")
    os.close(fd)
    bad_set = set(bad_entries)
    with zipfile.ZipFile(in_path, "r") as zin, \
         zipfile.ZipFile(dst, "w", zipfile.ZIP_DEFLATED) as zout:
        for zinfo in zin.infolist():
            if zinfo.filename in bad_set:
                # 깨진 이미지 → placeholder bytes로 교체. 이름/경로/확장자 유지.
                zout.writestr(zinfo.filename, placeholder)
            else:
                data = zin.read(zinfo.filename)
                zout.writestr(zinfo, data)
    return dst


def _make_context_excerpt(paragraph_text: str, target: str, around: int = 30) -> str:
    """
    Carve a short window of the source paragraph around `target` and wrap the
    target with 「…」 for visual emphasis. Streamlit's data_editor cells render
    as plain text, so we use a marker character pair rather than markdown.
    """
    idx = paragraph_text.find(target)
    if idx < 0:
        return f"「{target}」"
    start = max(0, idx - around)
    end = min(len(paragraph_text), idx + len(target) + around)
    excerpt = paragraph_text[start:end].strip()
    excerpt = excerpt.replace(target, f"「{target}」", 1)
    prefix = "…" if start > 0 else ""
    suffix = "…" if end < len(paragraph_text) else ""
    return f"{prefix}{excerpt}{suffix}"


def extract_bold_texts_with_context(in_path: str) -> List[Tuple[str, str]]:
    """
    Like [extract_bold_texts], but for each unique bold KO segment also returns
    the sentence-level context in which it first appears, with the target
    surrounded by 「…」 so the user can disambiguate "관리" (메뉴 라벨인지 직책인지)
    while filling out the mapping table.

    Returns a list of (ko, context_excerpt) tuples, preserving first-occurrence order.
    """
    doc = Document(sanitize_docx(in_path))
    seen: set = set()
    result: List[Tuple[str, str]] = []
    for p in iter_all_paragraphs(doc):
        if is_heading_paragraph(p):
            continue
        marked, _, _, _ = paragraph_to_marked_text(p)
        # 컨텍스트는 plain text 기준 — 모든 마커 제거 후 단락 텍스트
        plain_para = re.sub(r"⟦[^⟧]+⟧", "", marked)
        for match in _BOLD_SEGMENT_RE.finditer(marked):
            text = _INNER_MARKER_RE.sub("", match.group(1)).strip()
            if not text or not contains_korean(text):
                continue
            if text in seen:
                continue
            seen.add(text)
            context = _make_context_excerpt(plain_para, text)
            result.append((text, context))
    return result


def extract_bold_texts(in_path: str) -> List[str]:
    """
    Pull every bold-formatted text segment containing Korean from a docx file.

    Used by the UI text-mapping screen so the user can pre-translate the
    sentences/words that appear inside ⟦B⟧…⟦/B⟧ in the document. Returned
    list is de-duplicated and preserves the order each segment first appears.

    Skips:
    - Heading / Title paragraphs (Heading 1/2/…, Title): these are document
      structure labels, not UI strings, so the user writes their own English
      version separately and they shouldn't pollute the mapping table.
    - Nested markers (drawings, hyperlinks, highlights) inside a bold span:
      stripped — only the human-readable text is returned.
    """
    doc = Document(sanitize_docx(in_path))
    seen: set = set()
    result: List[str] = []
    for p in iter_all_paragraphs(doc):
        if is_heading_paragraph(p):
            continue  # heading/title은 UI 텍스트 매핑 대상이 아님
        marked, _, _, _ = paragraph_to_marked_text(p)
        for match in _BOLD_SEGMENT_RE.finditer(marked):
            text = _INNER_MARKER_RE.sub("", match.group(1)).strip()
            if not text or not contains_korean(text):
                continue
            if text in seen:
                continue
            seen.add(text)
            result.append(text)
    return result


# Per-run styling tags that must NOT be inherited by rebuilt runs —
# each segment carries its own bold/italic/etc. from the translated markers.
_FORMATTING_TAGS = frozenset(
    # "highlight" is intentionally excluded so that highlighted text keeps
    # its background colour after run rebuilding.
    {"b", "bCs", "i", "iCs", "u", "strike", "dstrike", "rStyle"}
)


def _make_rPr_template(rPr_src):
    """Clone an rPr element keeping only structural properties (font, size,
    colour, spacing…) and stripping per-character styling (bold, italic,
    underline, rStyle) so the template can be safely reused across segments
    with different formatting."""
    if rPr_src is None:
        return None
    tmpl = copy.deepcopy(rPr_src)
    for tag in _FORMATTING_TAGS:
        el = tmpl.find(f"{{{_W}}}{tag}")
        if el is not None:
            tmpl.remove(el)
    hl = tmpl.find(f"{{{_W}}}highlight")
    if hl is not None:
        tmpl.remove(hl)
    return tmpl


def _repair_lb_remnants(text: str) -> str:
    """
    Salvage line-break marker damage the LLM may have caused.

    If the LLM (rare but observed) writes ⟦BR⟧ as literal, or splits ⟦LB⟧
    into "⟦L⟧B⟧" / "⟦B⟧R⟧" / trailing "R⟧", we normalise those artefacts
    back to a clean ⟦LB⟧ so the downstream <w:br/> substitution still fires
    at the right place instead of leaving garbage in the final docx.
    """
    if not text or "⟦" not in text and "R⟧" not in text and "B⟧" not in text:
        return text
    # Ordered replacements — most specific first.
    text = re.sub(r"⟦BR⟧", BR_MARKER, text)              # ⟦BR⟧ → ⟦LB⟧
    text = re.sub(r"⟦L\s*⟧\s*B\s*⟧", BR_MARKER, text)   # ⟦L⟧B⟧ → ⟦LB⟧
    text = re.sub(r"⟦B\s*⟧\s*R\s*⟧", BR_MARKER, text)   # ⟦B⟧R⟧ → ⟦LB⟧
    text = re.sub(r"⟦B\s*R\s*⟧", BR_MARKER, text)        # ⟦BR⟧ variant → ⟦LB⟧
    # 탭 마커도 같은 방식으로 훼손될 수 있다 (⟦T⟧B⟧ / ⟦T B⟧)
    text = re.sub(r"⟦T\s*⟧\s*B\s*⟧", TAB_MARKER, text)
    text = re.sub(r"⟦T\s*B\s*⟧", TAB_MARKER, text)
    # 마지막 안전장치: 남아있는 "R⟧" (앞에 ⟦B⟧가 이미 처리된 잔해)를 제거
    text = re.sub(r"(?<![⟦A-Za-z0-9])R\s*⟧\s*", BR_MARKER, text)
    return text


def _make_run(rPr_template, text: str, props):
    """Create a fresh w:r with a cloned rPr template, bold flag, and text.

    If `text` contains any ⟦LB⟧ markers they are converted to actual
    <w:br/> soft-line-break elements interleaved with <w:t> chunks — this
    preserves the original document's paragraph-internal line breaks
    across translation. Also repairs any damaged BR/LB marker remnants
    that leaked in from LLM responses.
    """
    text = _repair_lb_remnants(text)
    is_bold = props.get("bold", False) if isinstance(props, dict) else bool(props)
    hl_val  = props.get("hl",   None)  if isinstance(props, dict) else None
    r = etree.Element(f"{{{_W}}}r")
    if rPr_template is not None:
        rPr_new = copy.deepcopy(rPr_template)
        if is_bold:
            etree.SubElement(rPr_new, f"{{{_W}}}b")
        if hl_val:
            hl_el = etree.SubElement(rPr_new, f"{{{_W}}}highlight")
            hl_el.set(f"{{{_W}}}val", hl_val)
        r.append(rPr_new)
    elif is_bold or hl_val:
        rPr_new = etree.SubElement(r, f"{{{_W}}}rPr")
        if is_bold:
            etree.SubElement(rPr_new, f"{{{_W}}}b")
        if hl_val:
            hl_el = etree.SubElement(rPr_new, f"{{{_W}}}highlight")
            hl_el.set(f"{{{_W}}}val", hl_val)

    # ⟦LB⟧ / ⟦TB⟧ 마커가 있으면 <w:t> + <w:br/> + <w:tab/> + <w:t> ... 순으로
    # 여러 요소를 문서 순서 그대로 삽입한다.
    if BR_MARKER in text or TAB_MARKER in text:
        for piece in LAYOUT_MARKER_SPLIT_RE.split(text):
            if piece == BR_MARKER:
                etree.SubElement(r, f"{{{_W}}}br")
            elif piece == TAB_MARKER:
                etree.SubElement(r, f"{{{_W}}}tab")
            elif piece:
                t_elem = etree.SubElement(r, f"{{{_W}}}t")
                t_elem.text = piece
                if piece[0] == " " or piece[-1] == " ":
                    t_elem.set(f"{{{_XML_SPACE}}}space", "preserve")
    else:
        t_elem = etree.SubElement(r, f"{{{_W}}}t")
        t_elem.text = text
        if text and (text[0] == " " or text[-1] == " "):
            t_elem.set(f"{{{_XML_SPACE}}}space", "preserve")
    return r


def _set_run_text(r_elem, text: str, is_bold: bool) -> None:
    """Set w:t text and sync w:b bold on a single run element."""
    t_elem = _lxml_text_elem(r_elem)
    if t_elem is None:
        return
    t_elem.text = text
    if text and (text[0] == " " or text[-1] == " "):
        t_elem.set(f"{{{_XML_SPACE}}}space", "preserve")
    else:
        t_elem.attrib.pop(f"{{{_XML_SPACE}}}space", None)
    rPr = r_elem.find(f"{{{_W}}}rPr")
    if rPr is None:
        if is_bold:
            rPr = etree.SubElement(r_elem, f"{{{_W}}}rPr")
            etree.SubElement(rPr, f"{{{_W}}}b")
            r_elem.insert(0, rPr)
    else:
        b_elem = rPr.find(f"{{{_W}}}b")
        if is_bold and b_elem is None:
            etree.SubElement(rPr, f"{{{_W}}}b")
        elif not is_bold and b_elem is not None:
            rPr.remove(b_elem)


def _write_paragraph_inplace(p_elem, translated_marked: str,
                             drawing_map: Dict, hyperlink_map: Dict,
                             trailing_lb: int = 0) -> None:
    """
    Write translated text back into the paragraph XML — three-path strategy.

    PATH 1 — Drawing paragraph
        Groups runs around drawing elements; first run of each group gets the
        combined translated text for that segment; drawing runs stay untouched.

    PATH 2 — Hyperlink paragraph
        ⟦H0⟧…⟦/H0⟧ text → written into the hyperlink's first inner run.
        Plain text → direct w:r children rebuilt from a clean rPr template,
        inserted at positions that preserve the before/after-hyperlink order.

    PATH 3 — Plain paragraph
        rPr template cloned from first run (stripping per-segment styling).
        All direct w:r removed; rebuilt one per (bold, text) segment.
    """
    # 끝부분 ⟦LB⟧는 번역/QA가 임의로 붙였을 수 있으므로 일단 전부 걷어낸 뒤,
    # **원문이 실제로 갖고 있던 개수(trailing_lb)만큼만** 다시 붙인다.
    # → 원문의 의도적인 빈 줄(단락 사이 간격)은 살리고, LLM이 지어낸 여분의
    #   빈 줄은 들어오지 못하게 하는 양방향 방어.
    translated_marked = re.sub(
        rf"({re.escape(BR_MARKER)}\s*)+$", "", translated_marked or ""
    )
    if trailing_lb > 0:
        translated_marked += BR_MARKER * trailing_lb

    all_runs = _lxml_all_runs(p_elem)
    if not all_runs:
        return

    has_drawing   = any(_run_is_non_text(r) for r in all_runs)
    has_hyperlink = bool(hyperlink_map)

    # ── Unified write-back (drawing / hyperlink / highlight / comment-safe) ───
    #
    # All paragraph types use the same strategy:
    # 1. Build run groups separated by drawing runs (direct children only)
    # 2. Each group also tracks which hyperlink nodes belong to it
    # 3. Translated segments for each group are split at hyperlink boundaries
    #    and the resulting text chunks are interleaved with the hyperlink nodes
    # 4. Drawing runs stay in their XML positions; comment-ref runs are never removed

    segs = _parse_marked_segments(translated_marked)

    def _get_rPr_template():
        for child in p_elem:
            ctag = child.tag.split("}")[1] if "}" in child.tag else child.tag
            if ctag == "r" and not _run_is_comment_ref(child) and not _run_is_non_text(child):
                t = _lxml_text_elem(child)
                if t is not None:
                    rPr = child.find(f"{{{_W}}}rPr")
                    if rPr is not None:
                        return _make_rPr_template(rPr)
        return None

    rPr_tmpl = _get_rPr_template()

    def _rebuild_group(group_runs: List, slot_segs: List[Tuple], group_hls: List) -> None:
        if group_runs:
            insert_pos = list(p_elem).index(group_runs[0])
        else:
            insert_pos = len(list(p_elem))

        for r in group_runs:
            p_elem.remove(r)

        for hi, hl_info in hyperlink_map.items():
            if hl_info["elem"] in group_hls:
                hl_runs = hl_info["runs"]
                hl_text = "".join(
                    t for b, t in slot_segs if isinstance(b, tuple) and b == ("h", hi)
                )
                if hl_runs:
                    _set_run_text(hl_runs[0], hl_text, False)
                    for r in hl_runs[1:]:
                        _set_run_text(r, "", False)

        chunks: List[List[Tuple]] = []
        hl_after_chunk: List = []
        cur_chunk: List[Tuple] = []
        for b, t in slot_segs:
            if isinstance(b, tuple) and b[0] == "h":
                hl_elem = hyperlink_map.get(b[1], {}).get("elem")
                chunks.append(cur_chunk)
                hl_after_chunk.append(hl_elem if hl_elem in group_hls else None)
                cur_chunk = []
            else:
                cur_chunk.append((b, t))
        chunks.append(cur_chunk)

        offset = 0
        for ci, chunk in enumerate(chunks):
            for props, text in chunk:
                p_elem.insert(insert_pos + offset, _make_run(rPr_tmpl, text, props))
                offset += 1
            if ci < len(hl_after_chunk) and hl_after_chunk[ci] is not None:
                hl_elem = hl_after_chunk[ci]
                cur_pos = list(p_elem).index(hl_elem)
                target  = insert_pos + offset
                if cur_pos != target:
                    p_elem.remove(hl_elem)
                    p_elem.insert(target if cur_pos > target else target - 1, hl_elem)
                offset += 1

    # Build run groups (separated by drawing runs) from direct children
    direct_children = list(p_elem)
    run_groups: List[List] = []
    hl_groups:  List[List] = []
    cur_runs: List = []
    cur_hls:  List = []
    for child in direct_children:
        ctag = child.tag.split("}")[1] if "}" in child.tag else child.tag
        if ctag == "r" and _run_is_non_text(child):
            run_groups.append(cur_runs); hl_groups.append(cur_hls)
            cur_runs = []; cur_hls = []
        elif ctag == "r" and not _run_is_comment_ref(child) and _run_has_writable_content(child):
            cur_runs.append(child)
        elif ctag == "hyperlink" and any(info["elem"] is child for info in hyperlink_map.values()):
            cur_hls.append(child)
    run_groups.append(cur_runs); hl_groups.append(cur_hls)

    # Split translated segments at drawing markers → one slot per run group
    text_slots: List[List[Tuple]] = []
    cur_slot: List[Tuple] = []
    for seg in segs:
        if seg[0] == "drawing":
            text_slots.append(cur_slot); cur_slot = []
        else:
            cur_slot.append(seg)
    text_slots.append(cur_slot)

    for gi in range(len(run_groups)):
        _rebuild_group(run_groups[gi], text_slots[gi] if gi < len(text_slots) else [], hl_groups[gi])


# 마커처럼 생겼지만 우리가 아는 형태가 아닌 것 — ⟦HL:yellow (닫힘 없음),
# ⟦/b (소문자), ⟧만 남은 조각 등. 이런 게 Word에 그대로 찍혀 왔다.
# 색 이름은 Word가 쓰는 값으로 한정한다. [A-Za-z]+ 로 두면 탐욕적으로 매칭돼
# "⟦HL:yellowAI provider" 에서 뒤따르는 본문 "AI"까지 먹어버린다.
_HL_COLOR_ALT = (
    "darkBlue|darkCyan|darkGray|darkGrey|darkGreen|darkMagenta|darkRed|darkYellow|"
    "lightGray|lightGrey|black|blue|cyan|green|magenta|red|white|yellow|none"
)
_MARKER_DEBRIS_RE = re.compile(
    rf"⟦/?HL(?::(?:{_HL_COLOR_ALT}))?⟧?|⟦/?[A-Za-z]{{1,2}}\d*⟧?",
    re.IGNORECASE,
)


def sanitize_broken_markers(text: str) -> str:
    """
    정상 마커는 그대로 두고, 손상된 마커 잔재만 지운다.

    복구(repair_*)는 어디까지나 추측이라 형태가 조금만 어긋나도 못 잡는다.
    실제 산출물에는 "⟦HL:yellowAI provider type" 이나 "…schema⟦/b" 가 그대로
    남았다. 마지막 방어선은 **추측하지 않는 것** — 못 알아보는 마커는 텍스트로
    출력하느니 지운다. 서식은 잃어도 문서에 쓰레기 문자가 박히지는 않는다.
    """
    if not text or ("⟦" not in text and "⟧" not in text):
        return text
    parts = re.split(f"({_KNOWN_MARKER_RE.pattern})", text)
    out = []
    for i, part in enumerate(parts):
        if i % 2 == 1:                      # 정상 마커 — 손대지 않는다
            out.append(part)
            continue
        part = _MARKER_DEBRIS_RE.sub("", part)
        out.append(part.replace("⟦", "").replace("⟧", ""))
    return "".join(out)


def finalize_markers(src_marked: str, translated: str) -> Tuple[str, List[str]]:
    """
    Word에 쓰기 **직전** 마지막 관문. (정리된 텍스트, 남은 문제 목록)

    지금까지 마커 복구는 번역 직후에만 돌았다. 그런데 그 뒤로도 대소문자
    정규화·문장부호 보정 같은 후처리가 줄줄이 이어지고, 거기서 깨진 마커는
    아무도 다시 보지 않은 채 문서에 기록됐다. 하이라이트가 사라진 8개 문단과
    본문에 노출된 마커 2곳이 모두 이 구간에서 생겼다.

    그래서 순서를 이렇게 잡는다.
      1. 복구할 수 있는 건 복구하고
      2. 경계 공백을 정리하고
      3. 그래도 못 알아보는 잔재는 지우고
      4. 하이라이트가 통째로 날아갔으면 문단 전체에 다시 씌우고
      5. 마지막으로 검사해 남은 문제를 보고한다
    """
    text = repair_bold_markers(translated)
    text = repair_hl_markers(text)
    text = normalize_marker_boundary_spaces(text)
    text = sanitize_broken_markers(text)
    text = apply_highlight_fallback(text, src_marked)
    return text, check_marker_integrity(src_marked, text)


def strip_bold_markers(text: str) -> str:
    """Remove ⟦B⟧/⟦/B⟧ bold markers and ⟦HL:colour⟧/⟦/HL⟧ highlight markers."""
    text = text.replace(B_OPEN, "").replace(B_CLOSE, "")
    text = HL_OPEN_RE.sub("", text)
    text = text.replace(HL_CLOSE, "")
    return text


def set_paragraph_text_preserve_style(paragraph, text: str) -> None:
    text = strip_bold_markers(text)
    if paragraph.runs:
        paragraph.runs[0].text = text
        for r in paragraph.runs[1:]:
            r.text = ""
    else:
        paragraph.add_run(text)


def repair_bold_markers(text: str) -> str:
    if not text:
        return text

    # 1. ⟦Brule → ⟦B⟧rule
    text = re.sub(r"⟦B([A-Za-z])", r"⟦B⟧\1", text)

    # 2. ⟦/Brule → ⟦/B⟧rule
    text = re.sub(r"⟦/B([A-Za-z])", r"⟦/B⟧\1", text)

    # 3. ⟦B rule → ⟦B⟧rule
    text = re.sub(r"⟦B\s+", "⟦B⟧", text)

    # 4. ⟦/B rule → ⟦/B⟧rule
    text = re.sub(r"⟦/B\s+", "⟦/B⟧", text)

    # 5. 잘못된 닫힘 제거 (⟧/B⟧ → ⟦/B⟧)
    text = re.sub(r"⟧/B⟧", "⟦/B⟧", text)

    # 6. 개수 맞추기
    opens = text.count("⟦B⟧")
    closes = text.count("⟦/B⟧")

    if opens > closes:
        text += "⟦/B⟧"
    elif closes > opens:
        for _ in range(closes - opens):
            text = text.replace("⟦/B⟧", "", 1)

    return text


# 인라인 구간을 여는/닫는 마커 — 경계 공백을 밖으로 밀어낼 대상.
_SPAN_OPEN_RE = re.compile(r"(⟦B⟧|⟦I⟧|⟦HL:[a-zA-Z]+⟧|⟦H\d+⟧)[ \t]+")
_SPAN_CLOSE_RE = re.compile(r"[ \t]+(⟦/B⟧|⟦/I⟧|⟦/HL⟧|⟦/H\d+⟧)")
# 줄이 시작하는 자리 — 여기에 밀려온 공백은 버린다.
_LINE_HEAD_RE = re.compile(r"(^|⟦LB⟧|⟦TB⟧)[ \t]+")
_LINE_TAIL_RE = re.compile(r"[ \t]+($|⟦LB⟧|⟦TB⟧)")

# 경계에 **없는** 공백을 새로 넣어야 하는 자리.
#
# 공백을 옮기는 것만으로는 부족하다. 원문이 한국어면 조사가 마커에 딱 붙어
# 있어(⟦/H0⟧에 접속합니다) 옮길 공백 자체가 없다. 번역되면 그 자리가
# 영단어끼리 맞붙어 "websiteGo", "syncClick", "optionsto", "toEdit",
# "a.csv"가 된다. 영문에서 낱말이 붙어 있으면 그건 오타지 서식이 아니다.
_GLUE_AFTER_RE = re.compile(
    r"(⟦/B⟧|⟦/I⟧|⟦/HL⟧|⟦/H\d+⟧|⟦D\d+⟧|⟦C\d+⟧|⟦G\d+⟧|⟦X\d+⟧)(?=[A-Za-z0-9])"
)
_GLUE_BEFORE_RE = re.compile(
    r"(?<=[A-Za-z0-9])(⟦B⟧|⟦I⟧|⟦HL:[a-zA-Z]+⟧|⟦H\d+⟧|⟦D\d+⟧|⟦C\d+⟧|⟦G\d+⟧|⟦X\d+⟧)"
)


def normalize_marker_boundary_spaces(text: str) -> str:
    """
    마커 경계에 낀 공백을 **삭제하지 않고 마커 바깥으로 옮긴다.**

    예전 구현은 그냥 지웠다. 그래서 원문이
        ⟦B⟧Agent Feedback ⟦/B⟧탭
    처럼 공백을 볼드 안쪽에 두고 있으면 번역 결과가
        Agent Feedbacktab
    으로 붙어버렸다. 실제 산출물의 "Delete Agentwindow", "Agent Feedbacktab"이
    전부 이것이다. 원문 Fireside v2.5 기준으로 볼드 닫기 직후에 글자가 붙는
    자리가 212곳이라 영향 범위가 작지 않다.

    공백은 의미를 가진 문자다. 없앨 게 아니라 구간 밖에 두면 된다.
        "…에서⟦B⟧ 사용자⟦/B⟧를"  ->  "…에서 ⟦B⟧사용자⟦/B⟧ 를"  (뒤에서 조사 정리)
    """
    text = _SPAN_OPEN_RE.sub(lambda m: " " + m.group(1), text)
    text = _SPAN_CLOSE_RE.sub(lambda m: m.group(1) + " ", text)
    # 구두점 앞으로 밀려난 공백은 되돌린다 ("Rule ." -> "Rule.")
    text = re.sub(r"[ \t]+([.,;:!?)\]}…])", r"\1", text)
    # 붙어버린 낱말 사이에는 공백을 새로 넣는다
    text = _GLUE_AFTER_RE.sub(r"\1 ", text)
    text = _GLUE_BEFORE_RE.sub(r" \1", text)
    # 줄 머리/꼬리로 밀려난 공백은 버린다
    text = _LINE_HEAD_RE.sub(r"\1", text)
    text = _LINE_TAIL_RE.sub(r"\1", text)
    text = re.sub(r"  +", " ", text)
    return text


# ──────────────────────────────────────────────────────────────────────
# 마커 무결성 검사
#
# 지금까지는 손상된 마커를 repair_*() 로 "고쳐서" 통과시켰다. 그런데 실제
# 산출물에는 ⟦HL:yellowAI provider type 이나 ⟦/b 처럼 복구 규칙이 잡지 못하는
# 형태가 그대로 Word에 기록됐다. 복구는 어디까지나 추측이고, 추측이 틀리면
# 조용히 망가진 문서가 나간다.
#
# 그래서 판단 기준을 바꾼다 — **고칠 수 있으면 고치고, 검사에 실패하면 그
# 후보 번역을 거부한다.** 무엇을 거부할지 판단하려면 먼저 "무엇이 잘못됐는가"
# 를 말할 수 있어야 하므로, 이 함수는 문제 목록을 사람이 읽는 문장으로
# 돌려준다. 빈 리스트면 무결성 통과.
# ──────────────────────────────────────────────────────────────────────

# ID를 가진 마커 — 원문과 **같은 것이 같은 개수만큼** 있어야 한다.
_ID_MARKER_KINDS = (
    ("용어 자리표시자 ⟦G#⟧", re.compile(r"⟦G\d+⟧")),
    ("원문 영문 봉인 ⟦X#⟧", re.compile(r"⟦X\d+⟧")),
    ("이미지 ⟦D#⟧", re.compile(r"⟦D\d+⟧")),
    ("주석 ⟦C#⟧", re.compile(r"⟦C\d+⟧")),
    ("하이퍼링크 열기 ⟦H#⟧", re.compile(r"⟦H\d+⟧")),
    ("하이퍼링크 닫기 ⟦/H#⟧", re.compile(r"⟦/H\d+⟧")),
)
# 개수만 맞으면 되는 마커
_COUNT_MARKER_KINDS = (
    ("줄바꿈 ⟦LB⟧", BR_MARKER),
    ("탭 ⟦TB⟧", TAB_MARKER),
)
# 열기/닫기가 짝을 이뤄야 하는 마커
_PAIR_MARKER_KINDS = (
    ("굵게", B_OPEN, B_CLOSE),
    ("기울임", "⟦I⟧", "⟦/I⟧"),
)
# 우리가 아는 정상 마커 전부. ⟦HL:…⟧을 ⟦H\d+⟧보다 앞에 둘 필요는 없다
# (H 뒤에 숫자를 요구하므로 겹치지 않는다) — 그래도 읽기 좋게 앞에 둔다.
_KNOWN_MARKER_RE = re.compile(
    r"⟦HL:[a-zA-Z]+⟧|⟦/HL⟧|⟦/?B⟧|⟦/?I⟧|⟦C\d+⟧|⟦D\d+⟧|⟦G\d+⟧|⟦X\d+⟧|"
    r"⟦/?H\d+⟧|⟦LB⟧|⟦TB⟧"
)


def _pairing_ok(text: str, open_tok: str, close_tok: str) -> bool:
    """열기/닫기가 순서대로 짝을 이루는가 (중첩 허용)."""
    depth = 0
    for m in re.finditer(
        rf"{re.escape(open_tok)}|{re.escape(close_tok)}", text
    ):
        depth += 1 if m.group(0) == open_tok else -1
        if depth < 0:
            return False
    return depth == 0


def check_marker_integrity(src: str, out: str) -> List[str]:
    """
    번역 후보(out)가 원문 marked text(src)의 마커 구조를 지켰는지 검사한다.

    반환값은 사람이 읽는 문제 설명 목록. 빈 리스트면 통과.
    """
    problems: List[str] = []
    if not src:
        return problems
    out = out or ""

    # 1. ID 마커는 멀티셋이 같아야 한다 (순서는 언어마다 달라질 수 있으므로 무시)
    for label, rx in _ID_MARKER_KINDS:
        a, b = Counter(rx.findall(src)), Counter(rx.findall(out))
        if a == b:
            continue
        lost = sorted((a - b).elements())
        extra = sorted((b - a).elements())
        bits = []
        if lost:
            bits.append(f"사라짐 {' '.join(lost[:5])}")
        if extra:
            bits.append(f"없던 것 생김 {' '.join(extra[:5])}")
        problems.append(f"{label}: " + " · ".join(bits))

    # 2. 개수만 맞으면 되는 마커
    for label, tok in _COUNT_MARKER_KINDS:
        na, nb = src.count(tok), out.count(tok)
        if na != nb:
            problems.append(f"{label}: {na}개 → {nb}개")

    # 3. 열기/닫기 짝
    for label, op, cl in _PAIR_MARKER_KINDS:
        na, nb = src.count(op), out.count(op)
        if na != nb or src.count(cl) != out.count(cl):
            problems.append(
                f"{label} 마커 개수 불일치: "
                f"열기 {na}→{nb} · 닫기 {src.count(cl)}→{out.count(cl)}"
            )
        elif not _pairing_ok(out, op, cl):
            problems.append(f"{label} 마커 짝이 어긋남 (닫기가 열기보다 먼저이거나 남음)")

    # 4. 하이라이트 — 색상까지 같아야 한다
    a, b = Counter(HL_OPEN_RE.findall(src)), Counter(HL_OPEN_RE.findall(out))
    if a != b:
        lost = sorted((a - b).elements())
        extra = sorted((b - a).elements())
        bits = []
        if lost:
            bits.append(f"사라진 색 {' '.join(lost)}")
        if extra:
            bits.append(f"생긴 색 {' '.join(extra)}")
        problems.append("하이라이트 색상 불일치: " + " · ".join(bits))
    elif src.count(HL_CLOSE) != out.count(HL_CLOSE):
        problems.append(
            f"하이라이트 닫기 개수: {src.count(HL_CLOSE)}개 → {out.count(HL_CLOSE)}개"
        )
    elif not _pairing_ok(HL_OPEN_RE.sub("⟦HL⟧", out), "⟦HL⟧", HL_CLOSE):
        problems.append("하이라이트 마커 짝이 어긋남")

    # 5. 하이퍼링크는 ID별로 열기·닫기가 하나씩
    for m in set(H_OPEN_RE.findall(out)):
        if out.count(f"⟦H{m}⟧") != out.count(f"⟦/H{m}⟧"):
            problems.append(f"하이퍼링크 ⟦H{m}⟧ 열기/닫기 개수가 다름")

    # 6. 정상 마커를 모두 지운 뒤에도 ⟦ ⟧ 가 남으면 손상된 마커다.
    #    ⟦HL:yellowAI provider… / ⟦/b 같은 형태가 여기서 잡힌다.
    residue = _KNOWN_MARKER_RE.sub("", out)
    if "⟦" in residue or "⟧" in residue:
        broken = re.findall(r"⟦[^⟦⟧]{0,30}|[^⟦⟧]{0,10}⟧", residue)
        sample = " / ".join(x.strip() for x in broken[:3] if x.strip())
        problems.append(f"손상된 마커가 남아 있음: {sample}")

    return problems

def repair_hl_markers(text: str) -> str:
    """
    Fix common LLM garbling of ⟦HL:colour⟧ … ⟦/HL⟧ highlight markers.

    - ⟦HL: yellow⟧  → ⟦HL:yellow⟧  (space after colon)
    - ⟦Hl:Yellow⟧ / ⟦/hl⟧ → ⟦HL:yellow⟧ / ⟦/HL⟧  (case garbling — LLM 응답이나
      sentence-case 후처리로 마커 안 글자가 뒤집힌 경우. 그대로 두면 토크나이저가
      마커로 인식하지 못해 "⟦Hl:yellow⟧" 문자열이 docx에 그대로 찍힌다.)
    - Unmatched open → append ⟦/HL⟧ at end
    - Unmatched close → remove excess ⟦/HL⟧
    """
    if not text:
        return text
    low = text.lower()
    if "⟦hl" not in low and "⟦/hl" not in low:
        return text
    # Normalise case/space garbling back to the canonical form
    text = re.sub(
        r"⟦\s*[Hh][Ll]\s*:\s*([a-zA-Z]+)\s*⟧",
        lambda m: f"⟦HL:{m.group(1).lower()}⟧",
        text,
    )
    text = re.sub(r"⟦\s*/\s*[Hh][Ll]\s*⟧", HL_CLOSE, text)
    opens  = len(HL_OPEN_RE.findall(text))
    closes = text.count(HL_CLOSE)
    if opens > closes:
        text = text + HL_CLOSE * (opens - closes)
    elif closes > opens:
        for _ in range(closes - opens):
            idx = text.rfind(HL_CLOSE)
            if idx >= 0:
                text = text[:idx] + text[idx + len(HL_CLOSE):]
    return text


def apply_highlight_fallback(translated: str, source_marked: str) -> str:
    """
    Safety net: if the source had ⟦HL:colour⟧ markers but the LLM dropped
    them from the translation, wrap the entire translated text in the dominant
    highlight colour so highlight is never silently lost.
    """
    src_colours = HL_OPEN_RE.findall(source_marked)
    if not src_colours:
        return translated                   # source had no highlight
    if HL_OPEN_RE.search(translated):
        return translated                   # LLM preserved markers — good
    # Fallback: whole-paragraph highlight with the dominant colour
    from collections import Counter
    colour = Counter(src_colours).most_common(1)[0][0]
    return f"⟦HL:{colour}⟧{translated}⟦/HL⟧"


def _split_preserving_markers(text: str) -> List[str]:
    return ANY_MARKER_SPLIT_RE.split(text)


def _is_marker_token(part: str) -> bool:
    return bool(ANY_MARKER_SPLIT_RE.fullmatch(part) or PLACEHOLDER_RE.fullmatch(part))


def _sentence_case_preserving_markers(text: str) -> str:
    parts = _split_preserving_markers(text)
    out = []
    first_alpha_done = False

    for part in parts:
        if not part:
            out.append(part)
            continue

        if _is_marker_token(part):
            out.append(part)
            continue

        chars = list(part)

        if not first_alpha_done:
            for i, ch in enumerate(chars):
                if ch.isalpha():
                    chars[i] = ch.upper()
                    for j in range(i + 1, len(chars)):
                        if chars[j].isalpha():
                            chars[j] = chars[j].lower()
                    first_alpha_done = True
                    break
            out.append("".join(chars))
        else:
            lowered = []
            for ch in chars:
                lowered.append(ch.lower() if ch.isalpha() else ch)
            out.append("".join(lowered))

    return "".join(out)


def normalize_ui_label_text(text: str) -> str:
    parts = _split_preserving_markers(text)
    out = []
    first_alpha_word_seen = False
    word_re = re.compile(r"[A-Za-z][A-Za-z0-9/-]*")

    for part in parts:
        if not part:
            out.append(part)
            continue

        if _is_marker_token(part):
            out.append(part)
            continue

        def repl(match):
            nonlocal first_alpha_word_seen
            word = match.group(0)
            lower = word.lower()

            if not first_alpha_word_seen:
                first_alpha_word_seen = True
                return word

            if lower in UI_LOWER_WORDS:
                return lower

            return word

        out.append(word_re.sub(repl, part))

    return "".join(out)


def normalize_heading_text(text: str) -> str:
    if not text:
        return text

    s = text.strip()
    s = re.sub(r"[.。]+$", "", s)
    s = _sentence_case_preserving_markers(s)
    s = re.sub(r"[ \t]{2,}", " ", s).strip()
    s = re.sub(r"[.。]+$", "", s)
    return s


def normalize_ui_in_bold_segments(text: str) -> str:
    """
    Normalise each bold segment as a UI label: first alpha word kept as-is
    (already cased correctly by restore step), subsequent words lowercased
    if they are in UI_LOWER_WORDS, and the whole segment has its first alpha
    character capitalised.

    NOTE: Each bold segment is normalised independently, so a term restored
    at the start of a bold segment is already capitalised by
    restore_glossary_placeholders; this function only normalises the
    surrounding plain words inside the segment.
    """
    def repl(match):
        inner = match.group(1)
        inner = normalize_ui_label_text(inner)
        inner = _cap_first_alpha(inner)
        return B_OPEN + inner + B_CLOSE

    return re.sub(
        re.escape(B_OPEN) + r"(.*?)" + re.escape(B_CLOSE),
        repl,
        text,
        flags=re.DOTALL,
    )


# fix_indefinite_articles()는 제거했다.
#
# 철자만 보고 a/an을 바꾸는 규칙은 구조적으로 안전하지 않다. 관사는 뒤따르는
# 단어의 **발음**을 따르는데 정규식은 철자밖에 못 본다.
#     a user      -> an user      (틀림)
#     a URL       -> an URL       (틀림, "유알엘"로 읽는다)
#     a university-> an university(틀림)
# 게다가 역방향 규칙은 치환문이 대문자 "A"로 고정돼 있어 문장 중간의
#     an hour -> A hour
# 처럼 관사와 대소문자를 동시에 망가뜨렸다.
# 관사 교정은 정규식이 아니라 QA 프롬프트에서 다룬다.


def capitalize_bullet_lines(text: str) -> str:
    # ⟦LB⟧로 표현된 soft line break도 하나의 "줄"로 취급해야 한다. 그렇지
    # 않으면 "1. …⟦LB⟧2. …" 형태의 문단에서 두 번째 항목 이후가 대문자화되지
    # 않는다.
    if BR_MARKER in text:
        return BR_MARKER.join(
            capitalize_bullet_lines(seg) for seg in text.split(BR_MARKER)
        )

    lines = text.splitlines()
    out_lines = []
    num_prefix_re = re.compile(r"^\s*\d+[\.\)]\s+")
    bullet_prefixes = ("- ", "• ", "∙ ", "* ")

    for line in lines:
        if not line.strip():
            out_lines.append(line)
            continue

        stripped = line.lstrip()
        indent = line[: len(line) - len(stripped)]

        handled = False
        for bp in bullet_prefixes:
            if stripped.startswith(bp):
                rest = stripped[len(bp) :]
                rest = _cap_first_alpha(rest)
                out_lines.append(indent + bp + rest)
                handled = True
                break
        if handled:
            continue

        m = num_prefix_re.match(stripped)
        if m:
            pre = stripped[: m.end()]
            rest = stripped[m.end() :]
            rest = _cap_first_alpha(rest)
            out_lines.append(indent + pre + rest)
            continue

        out_lines.append(line)

    return "\n".join(out_lines)


def restore_sentence_period(translated: str, source: str) -> str:
    """
    If the source sentence ended with a period and the translation does not
    end with any terminal punctuation, append a period.

    The LLM occasionally drops the trailing period when the sentence ends with
    a glossary placeholder (e.g. 'Enter ⟦B⟧basic information⟦/B⟧' → no '.').

    When both sides carry the same number of ⟦LB⟧ soft line breaks the check is
    made per line, so each line of a multi-line paragraph keeps (or loses) its
    period independently instead of only the last one being considered.
    """
    if BR_MARKER in source and line_breaks_match(translated, source):
        return BR_MARKER.join(
            restore_sentence_period(t, s)
            for t, s in zip(translated.split(BR_MARKER), source.split(BR_MARKER))
        )

    src_stripped = source.rstrip()
    if not src_stripped.endswith("."):
        return translated          # source had no period — nothing to restore

    tr = translated.rstrip()
    if tr and tr[-1] not in ".!?:;":
        translated = tr + "."
    return translated


_LIST_LINE_RE = re.compile(r"^\s*([-•∙*]|\d+[.)]|[a-zA-Z][.)])\s+")


def normalize_paragraph_breaks(s: str, source_marked: str = "") -> str:
    """
    Reconcile the line structure of an LLM response with the source paragraph.

    Two regimes, decided by whether the SOURCE paragraph had soft line breaks:

    * source has ⟦LB⟧ — the author deliberately broke this paragraph into
      several lines (Shift+Enter). The model very often answers with real
      newline characters instead of the ⟦LB⟧ token, so we *promote* every
      newline to ⟦LB⟧ rather than destroying it. Collapsing here was the main
      reason multi-line notes came back as one run-on line.

    * source has no ⟦LB⟧ — a translated Word paragraph should stay one
      continuous line, so stray newlines are collapsed into spaces. Line
      breaks are still preserved when the paragraph is actually a bulleted /
      numbered list (each line begins with a bullet or number marker).
    """
    if not s:
        return s
    s = s.replace("\r\n", "\n").replace("\r", "\n")
    s = s.strip("\n")

    if source_marked and BR_MARKER in source_marked:
        # 개행 → ⟦LB⟧ 승격. 마커 주변 공백은 정리하되 줄 자체는 유지.
        s = re.sub(r"[ \t]*\n+[ \t]*", BR_MARKER, s)
        s = re.sub(rf"\s*{re.escape(BR_MARKER)}\s*", BR_MARKER, s)
        return s

    if "\n" not in s:
        return s

    lines = s.split("\n")
    non_empty = [ln for ln in lines if ln.strip()]
    # 2번째 라인부터 다수가 list marker로 시작하면 list로 간주 → 원본 유지
    if len(non_empty) >= 2 and sum(bool(_LIST_LINE_RE.match(ln)) for ln in non_empty[1:]) >= 1:
        s = re.sub(r"\n{3,}", "\n\n", s)
        return s

    # 그 외엔 단순 산문 — 모든 line break를 space로 흡수
    return re.sub(r"\s*\n+\s*", " ", s).strip()


def line_breaks_match(translated: str, source_marked: str) -> bool:
    """True when the translation carries exactly as many ⟦LB⟧ as the source."""
    return (translated or "").count(BR_MARKER) == (source_marked or "").count(BR_MARKER)


def enforce_line_breaks(translated: str, source_marked: str) -> str:
    """
    Last-resort reconciliation of the ⟦LB⟧ count against the source.

    `translate_marked_paragraph` already guarantees the count structurally by
    translating line-by-line, so this only has to catch damage introduced
    afterwards (marker repair, glossary restore, QA rewrite). It never invents
    breaks out of thin air — it only recovers ones the model expressed in a
    different form (real newlines, damaged markers) and trims surplus ones.
    """
    want = (source_marked or "").count(BR_MARKER)
    text = _repair_lb_remnants(translated or "")

    if want == 0:
        # 원문에 줄바꿈이 없었다면 번역이 만들어낸 ⟦LB⟧는 전부 여분이다.
        if BR_MARKER in text:
            text = re.sub(rf"\s*{re.escape(BR_MARKER)}\s*", " ", text)
            text = re.sub(r" {2,}", " ", text).strip()
        return text

    # 모델이 ⟦LB⟧ 대신 실제 개행으로 답한 경우 승격
    if "\n" in text:
        text = re.sub(r"[ \t]*\n+[ \t]*", BR_MARKER, text)

    have = text.count(BR_MARKER)
    if have > want:
        # 여분은 뒤에서부터 공백으로 흡수 (앞쪽 줄 구조가 더 신뢰도 높음)
        parts = text.split(BR_MARKER)
        head = parts[: want + 1]
        tail = parts[want + 1:]
        head[-1] = " ".join([head[-1]] + tail)
        text = BR_MARKER.join(head)
        text = re.sub(r" {2,}", " ", text)

    return text


def iter_all_paragraphs(doc: Document) -> Iterable:
    for p in doc.paragraphs:
        yield p
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    yield p


def is_heading_paragraph(p) -> bool:
    """
    Recognise heading / title paragraphs across Word locales, custom styles,
    and manually-outlined paragraphs.

    Detection order:
    1. Style name / style_id (English "Heading 1"/"Title", Korean "제목 1"/"표제")
       and its ancestor style chain — so a user-defined style that inherits
       from Heading 1 is caught.
    2. The paragraph's `w:pPr/w:outlineLvl` value. Any explicit outline level
       (0-8) means the author marked this paragraph as a heading via
       Word's Outline setting, even if the applied style isn't a heading style.
    """
    style = getattr(p, "style", None)
    seen_ids: set = set()
    while style is not None:
        sid = (getattr(style, "style_id", None) or "")
        name = (getattr(style, "name", None) or "")
        sid_l = sid.lower()
        name_l = name.lower()

        if sid_l.startswith("heading") or sid_l == "title":
            return True
        if name_l.startswith("heading") or name_l in ("title",):
            return True
        if name.startswith("제목") or name in ("표제",):
            return True

        if sid in seen_ids:
            break
        seen_ids.add(sid)
        style = getattr(style, "base_style", None)

    # outline level 확인 — 스타일과 무관하게 Word "Outline" 설정된 경우
    try:
        p_elem = getattr(p, "_p", None)
        if p_elem is not None:
            outline_nodes = p_elem.findall(f".//{{{_W}}}pPr/{{{_W}}}outlineLvl")
            if outline_nodes:
                val = outline_nodes[0].get(f"{{{_W}}}val")
                # 0-8은 heading level. 9는 body text.
                if val is not None and val.isdigit() and int(val) < 9:
                    return True
    except Exception:
        pass

    return False


def paragraph_has_hyperlink(paragraph) -> bool:
    return paragraph._p.xpath(".//w:hyperlink") != []


def _write_paragraph(p, translated_marked: str,
                     drawing_map: Optional[Dict] = None,
                     hyperlink_map: Optional[Dict] = None,
                     trailing_lb: int = 0) -> None:
    """Write translated text into the paragraph, preserving all XML structure."""
    if is_heading_paragraph(p):
        translated_marked = translated_marked.rstrip()
        if translated_marked.endswith("."):
            translated_marked = translated_marked[:-1].rstrip()

    _write_paragraph_inplace(
        p._p, translated_marked, drawing_map or {}, hyperlink_map or {}, trailing_lb
    )


def normalize_for_scoring(text: str) -> str:
    text = text.replace(B_OPEN, " ").replace(B_CLOSE, " ")
    text = re.sub(r"⟦G\d+⟧", " ", text)
    text = text.replace("~", " ")
    return text


def normalize_colon_label_line(text: str) -> str:
    """
    Normalize 'Label:' patterns so the label reads correctly as a UI element name.

    Handles two structures:
    1. Plain label before colon:  'rule Name:'         -> 'Rule name:'
    2. Bold-close + plain word:   '⟦B⟧Rule⟦/B⟧ Name:' -> '⟦B⟧Rule⟦/B⟧ name:'
       (the plain word after ⟦/B⟧ is a continuation of the bold label and must be
       lowercased when it is a UI_LOWER_WORD)

    The two passes run in order; Pass 2 skips any position already handled by Pass 1
    to prevent re-capitalising words that were just lowercased.
    """
    # Pass 1 — lowercase UI_LOWER_WORDS that sit between ⟦/B⟧ and ':'
    def _lower_after_bold(m):
        return m.group(1) + (m.group(2).lower() if m.group(2).lower() in UI_LOWER_WORDS else m.group(2)) + m.group(3)

    pass1 = re.sub(
        rf"({re.escape(B_CLOSE)}\s+)([A-Za-z][A-Za-z0-9]*)(\s*:)",
        _lower_after_bold,
        text,
    )

    # Pass 2 — normalize plain (non-bold-wrapped) labels before ':'
    # Skip positions that are immediately preceded by ⟦/B⟧ (already handled by Pass 1)
    def repl(m):
        before = pass1[: m.start()]
        if re.search(rf"{re.escape(B_CLOSE)}\s*$", before):
            return m.group(0)   # already handled — leave untouched
        label_norm = normalize_ui_label_text(m.group(1).strip())
        label_norm = _cap_first_alpha(label_norm)
        return f"{label_norm}:"

    return re.sub(r"\b([A-Za-z][A-Za-z0-9 ]{1,50}):", repl, pass1)


def looks_like_heading_text(text: str) -> bool:
    """Heuristic: does this look like a heading rather than a body sentence?

    Korean source sentences are often short (≤ 3 words) even when they are
    step instructions ("기본 정보를 입력합니다."), so we suppress heading
    detection when the text contains Korean characters.  Only apply the
    short-text heuristic to already-translated (English) text.
    """
    s = strip_bold_markers(text).strip()

    if not s:
        return False
    if "\n" in s:
        return False
    if len(s) > 50:
        return False
    if s.endswith(":"):
        return False
    # Korean text: never treat as heading via this heuristic —
    # Korean step instructions are naturally short (≤ 3 words).
    if KOREAN_RE.search(s):
        return False

    words = s.split()
    if len(words) <= 3:
        return True

    return False


def tokenize_koreanish(text: str) -> List[str]:
    norm = normalize_for_scoring(text)
    return re.findall(r"[가-힣A-Za-z0-9]+", norm)


def select_relevant_patterns(
    source_text: str,
    patterns: List[Tuple[str, str]],
    max_pattern: int = 3,
) -> List[Tuple[str, str]]:
    source_tokens = set(tokenize_koreanish(source_text))

    def score_patterns(ko: str) -> int:
        pattern_tokens = set(tokenize_koreanish(ko))
        return len(source_tokens & pattern_tokens)

    scored_patterns = []
    for ko, en in patterns:
        score = score_patterns(ko)
        if score > 0:
            scored_patterns.append((score, ko, en))

    scored_patterns.sort(key=lambda x: (-x[0], -len(x[1])))
    return [(ko, en) for score, ko, en in scored_patterns[:max_pattern]]


def _track_usage(resp) -> None:
    """Fold one Responses API call's usage into the module-level counters."""
    global TOTAL_INPUT_TOKENS, TOTAL_CACHED_INPUT_TOKENS, TOTAL_OUTPUT_TOKENS, TOTAL_TOKENS

    usage = getattr(resp, "usage", None)
    if not usage:
        return
    input_tokens = getattr(usage, "input_tokens", 0) or 0
    output_tokens = getattr(usage, "output_tokens", 0) or 0
    total_tokens = getattr(usage, "total_tokens", 0) or 0

    input_details = getattr(usage, "input_tokens_details", None)
    cached_tokens = getattr(input_details, "cached_tokens", 0) or 0

    TOTAL_INPUT_TOKENS += input_tokens
    TOTAL_CACHED_INPUT_TOKENS += cached_tokens
    TOTAL_OUTPUT_TOKENS += output_tokens
    TOTAL_TOKENS += total_tokens


def translate_paragraph_with_patterns(
    client: OpenAI,
    source_text: str,
    pattern_examples: List[Tuple[str, str]],
    model: str = "gpt-5.2",
    translation_mode: str = "Manual",
    style_reference: str = "",
    line_count: int = 0,
) -> str:
    """
    Translate one marked paragraph (or one line of it).

    `line_count` > 0 switches the prompt into *line-aligned* mode: the input is
    a `[N] …` numbered list of the source's soft-line-break-separated lines and
    the model must echo exactly that many `[N]` blocks back. See
    `translate_marked_paragraph`.
    """
    line_mode = line_count > 0

    pattern_block = (
        "\n".join([f"- {ko} -> {en}" for ko, en in pattern_examples]) or "(none)"
    )

    if translation_mode in {"UI", "UI 텍스트"}:
        style_rules = """
- Prefer short, direct UI-style English.
- Keep labels and instructions concise.
- Avoid unnecessary words.
- Use product-style wording similar to Microsoft or Google UI.
"""
    else:
        style_rules = """
- Prefer natural, clear manual/documentation English.
- Use complete sentences where appropriate.
- Keep the tone professional and concise.
- Use enterprise software documentation style.
"""

    style_ref_block = (
        f"\nExisting English in this document (match tone, vocabulary, sentence structure):\n{style_reference}\n"
        if style_reference else ""
    )

    # ── 줄 구조 규칙 ────────────────────────────────────────────────────
    # 예전에는 "⟦LB⟧를 유지하라"와 "출력은 반드시 한 줄이어야 한다"가 같은
    # 프롬프트에 동시에 들어가 있어 모델이 줄바꿈을 지우는 쪽으로 기울었다.
    # 이제 원문의 줄 구조에 따라 둘 중 하나만 제시한다.
    if line_mode:
        layout_rules = f"""
- The input is a numbered list of {line_count} source line(s) in the form `[0] …`,
  `[1] …`. These are the lines of ONE Word paragraph, separated by soft line
  breaks in the original document.
- Translate each line INDEPENDENTLY and output exactly {line_count} block(s),
  each starting with its own `[N]` header on a new line, in ascending order,
  reusing the SAME numbers. Never merge two source lines into one block and
  never split one source line across two blocks.
- Each `[N]` block must be a single line of text. Do NOT emit ⟦LB⟧ inside a block.
- Do not output the `[N]` headers as part of the translated text itself.
"""
    else:
        layout_rules = """
- Output must be a SINGLE line. Do NOT insert line breaks, and do NOT emit any
  ⟦LB⟧ marker — the source paragraph has none.
"""

    tab_rule = (
        "- ⟦TB⟧ = a tab character in the source. Keep it, unsplit, at the same\n"
        "  position relative to the surrounding text.\n"
        if TAB_MARKER in source_text else ""
    )

    prompt = f"""
Translate Korean to natural, professional English. Produce a draft, then revise it so a native English speaker would not flag awkwardness, grammar errors, or literal-translation tells.

Rules:
- Preserve markers EXACTLY: ⟦G#⟧, ⟦B⟧, ⟦/B⟧, ⟦I⟧, ⟦/I⟧, ⟦C#⟧, ⟦D#⟧, ⟦H#⟧/⟦/H#⟧, ⟦HL:colour⟧/⟦/HL⟧, ⟦TB⟧.
{tab_rule}- ⟦G#⟧ placeholders are FIXED glossary terms. Output them BYTE-FOR-BYTE unchanged.
  NEVER translate, paraphrase, expand, or substitute a ⟦G#⟧ placeholder with any word.
- ⟦C#⟧ = a sealed literal (inline code, an escaped character, a URL). Output it
  BYTE-FOR-BYTE unchanged and keep it at the same position in the sentence.
- ⟦I⟧…⟦/I⟧ = italic span — translate the text inside, keep the pair around it.
- ⟦D#⟧ = inline icon/image — keep it where it naturally fits in the sentence.
- ⟦H#⟧…⟦/H#⟧ = hyperlink span — translate the text inside, keep the markers around it.
- IMPORTANT — you MAY MOVE ⟦H#⟧…⟦/H#⟧ and ⟦D#⟧ to wherever English word order
  needs them. Korean puts the link first ("⟦H0⟧the admin site⟦/H0⟧에 접속합니다")
  but English needs it last ("Go to ⟦H0⟧the admin site⟦/H0⟧."). Rewrite the
  sentence naturally and carry the span with it. Do NOT leave the span stranded
  at the front and then repeat its meaning in the rest of the sentence — that
  produces "⟦H0⟧the admin site⟦/H0⟧Go to the site." which is wrong.
- Put a space around a marker whenever English needs one between words. Korean
  particles attach directly to the marker; English words do not.
- ⟦HL:colour⟧…⟦/HL⟧ = highlighted text — translate the inside and keep the markers
  around the same semantic content in the translation.
- If the source contains existing English words or phrases, leave them EXACTLY as-is.
  Do NOT rephrase, reword, or "improve" any English that is already present.
- Use the pattern examples only as reference guidance.
- Do not copy irrelevant examples.
- Avoid repetition and awkward literal wording.
- Grammar: maintain subject-verb agreement and correct singular/plural agreement.
  Korean does not always mark plurality, but English does. Prefer "all users participating",
  "individual users", "authorized users" over "all user", "individual user", "authorized user".
- Subjects: avoid bare generic singular subjects without articles. "Can user prevent..." is
  ungrammatical — use "Can users prevent..." or supply a concrete subject like
  "Can administrators prevent users from...".
- Naturalize literal renderings into idiomatic English. For example:
    "신개념 협업 플랫폼" → "next-generation collaboration platform" (NOT "new-concept platform")
    "활용 가이드"        → "user guide" or "usage guide" (NOT "utilization guide")
  After your initial draft, re-read each sentence and replace any phrasing that reads as
  Korean-influenced literal English with the way a native speaker would actually write it.
- Spell English words correctly. NEVER transliterate Korean phonetic spellings.
  "프로그램" is "program" (never "logram", "pulogeurem", etc.). Treat any non-word
  English output as a serious error and self-correct.
- Articles: use "a/an/the" appropriately. Korean has no articles, so do not omit them.
- Do not force title case.
- For headings, concise phrase-style English is preferred.
- NEVER merge markers with words. "⟦B⟧rule" is correct; "⟦Brule" is invalid.
- Keep markers as separate tokens.
- Output ONLY the translated text. No explanation, no extra lines.
- ABSOLUTELY DO NOT emit meta labels ("Draft:", "Revised:", "Original:",
  "Before:", "After:", "Version 1:", "v2:", "Option A:", "Improved:",
  "Correction:") or any before/after comparison. Give ONE clean final
  translation, no alternatives, no explanation.
{layout_rules}{style_rules}{style_ref_block}
Reference pattern examples:
{pattern_block}

Text to translate:
{source_text}
""".strip()

    resp = client.responses.create(
        model=model,
        input=prompt,
        reasoning={"effort": "low"},
        text={"verbosity": "low"},
    )

    _track_usage(resp)

    return resp.output_text.strip()


_NUMBERED_LINE_RE = re.compile(r"^[ \t]*\[(\d+)\][ \t]?", re.MULTILINE)


def _parse_numbered_lines(text: str, expected: int) -> Optional[List[str]]:
    """
    Parse a `[0] … [1] …` response into an ordered list of `expected` strings.

    Returns None when the response does not cover exactly indices 0..expected-1,
    so the caller can fall back to per-line translation rather than silently
    losing or duplicating a line.
    """
    if not text:
        return None
    matches = list(_NUMBERED_LINE_RE.finditer(text))
    if not matches:
        return None

    bodies: Dict[int, str] = {}
    for i, m in enumerate(matches):
        idx = int(m.group(1))
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[m.end():end].strip()
        # 한 줄 안에서 모델이 개행을 넣었다면 공백으로 흡수 — 줄 경계는
        # 어디까지나 [N] 헤더가 정의한다.
        body = re.sub(r"\s*\n+\s*", " ", body).strip()
        if idx in bodies:
            return None          # 중복 인덱스 → 신뢰 불가
        bodies[idx] = body

    if sorted(bodies) != list(range(expected)):
        return None
    if any(not b for b in bodies.values()):
        return None
    return [bodies[i] for i in range(expected)]


def translate_marked_paragraph(
    client: OpenAI,
    source_text: str,
    pattern_examples: List[Tuple[str, str]],
    model: str = "gpt-5.2",
    translation_mode: str = "Manual",
    style_reference: str = "",
) -> str:
    """
    Translate a marked paragraph, keeping its soft-line-break structure intact.

    Paragraphs without ⟦LB⟧ go straight through the normal single-shot path.
    Paragraphs *with* ⟦LB⟧ are split into lines, translated as a numbered list,
    and rejoined — so the number of line breaks is guaranteed by construction
    instead of depending on the model honouring a marker. If the numbered
    response cannot be parsed reliably we translate line by line, which is
    slower but still structurally exact.
    """
    if BR_MARKER not in source_text:
        return translate_paragraph_with_patterns(
            client=client,
            source_text=source_text,
            pattern_examples=pattern_examples,
            model=model,
            translation_mode=translation_mode,
            style_reference=style_reference,
        )

    lines = source_text.split(BR_MARKER)
    # 빈 줄(연속된 ⟦LB⟧)은 번역 대상이 아니므로 자리만 지킨다.
    todo = [i for i, ln in enumerate(lines) if ln.strip()]
    if not todo:
        return source_text

    numbered = "\n".join(f"[{k}] {lines[i]}" for k, i in enumerate(todo))
    try:
        raw = translate_paragraph_with_patterns(
            client=client,
            source_text=numbered,
            pattern_examples=pattern_examples,
            model=model,
            translation_mode=translation_mode,
            style_reference=style_reference,
            line_count=len(todo),
        )
        parsed = _parse_numbered_lines(raw, len(todo))
    except Exception:
        parsed = None

    if parsed is None:
        parsed = [
            translate_paragraph_with_patterns(
                client=client,
                source_text=lines[i],
                pattern_examples=pattern_examples,
                model=model,
                translation_mode=translation_mode,
                style_reference=style_reference,
            ).strip()
            for i in todo
        ]

    out = list(lines)
    for k, i in enumerate(todo):
        out[i] = parsed[k]
    return BR_MARKER.join(out)


def translate_remaining_korean(
    client: OpenAI,
    text: str,
    model: str = "gpt-5.2",
) -> str:
    if not contains_korean(text):
        return text

    prompt = f"""
Translate any remaining Korean in the text into natural English.

Rules:
- Preserve markers EXACTLY: ⟦B⟧, ⟦/B⟧, ⟦G#⟧, ⟦LB⟧, ⟦TB⟧.
- ⟦LB⟧ (soft line break) and ⟦TB⟧ (tab) must be kept in place and in the same
  number as the input. Never split, add, or drop one.
- ⟦G#⟧ placeholders are FIXED terms — output them UNCHANGED. Do NOT translate them.
- Do NOT alter, rephrase, or improve any English that is already present.
- Only translate Korean words/phrases; leave everything else exactly as-is.
- Do not force title case.
- Output ONLY the translated text. No explanation, no extra lines.

Text:
{text}
""".strip()

    resp = client.responses.create(
        model=model,
        input=prompt,
        reasoning={"effort": "low"},
        text={"verbosity": "low"},
    )

    return resp.output_text.strip()


def extract_doc_style_guide(
    client: OpenAI,
    english_samples: List[str],
    model: str = "gpt-5.2",
) -> str:
    """
    Pre-pass: compress existing English content into a compact style guide.

    Called once per document before any paragraph translation. The returned
    string is reused as the `style_reference` for every per-paragraph LLM call,
    so it must be short (~250 words) — both to save tokens and to maximise
    prompt-cache hits across batches.
    """
    global TOTAL_INPUT_TOKENS, TOTAL_CACHED_INPUT_TOKENS, TOTAL_OUTPUT_TOKENS, TOTAL_TOKENS

    if not english_samples:
        return ""

    samples_text = "\n".join(english_samples)

    prompt = f"""Extract a CONCISE style guide from this product document's existing English content. The guide will be reused to keep newly translated Korean paragraphs consistent with the rest of the document.

Cover only what is distinctive (skip generic English advice):
- Specific terminology preferences this product uses (e.g. "Save" vs "Apply", "page" vs "screen")
- Sentence patterns for instructions, headings, and UI labels
- Formality level / tone (UI text vs manual prose)
- Capitalization conventions (sentence case, title case, lowercase UI words)
- Any product-specific phrasing the document repeats

Output as a compact bullet list. No introduction, no examples — bullets only. Maximum 250 words.

Existing English content:
{samples_text}
""".strip()

    resp = client.responses.create(
        model=model,
        input=prompt,
        reasoning={"effort": "low"},
        text={"verbosity": "low"},
    )

    usage = getattr(resp, "usage", None)
    if usage:
        input_tokens = getattr(usage, "input_tokens", 0) or 0
        output_tokens = getattr(usage, "output_tokens", 0) or 0
        total_tokens = getattr(usage, "total_tokens", 0) or 0
        input_details = getattr(usage, "input_tokens_details", None)
        cached_tokens = getattr(input_details, "cached_tokens", 0) or 0
        TOTAL_INPUT_TOKENS += input_tokens
        TOTAL_CACHED_INPUT_TOKENS += cached_tokens
        TOTAL_OUTPUT_TOKENS += output_tokens
        TOTAL_TOKENS += total_tokens

    return resp.output_text.strip()


_QA_HEADER_RE = re.compile(r"\[(\d+)\]")

# 메타 라벨 정규식들 — LLM이 무단으로 "Draft: X Revised: Y" 형태를 반환하는
# 경우를 감지·복구하기 위함. Draft/Revised 뿐 아니라 흔한 대체 표현도 커버.
# 이전 구현은 이 라벨들을 **문장 어디서든** 찾아 그 앞을 통째로 버렸다.
# 그래서 본문을 조용히 잘라먹었다.
#     "v2.1 update"          -> "1 update"     (v2. 가 라벨로 잡힘)
#     "...decide the final." -> ""             (final. 이 라벨로 잡힘)
# 라벨은 원래 "Draft: … Revised: …" 처럼 **줄 맨 앞에 콜론과 함께** 오는
# 형태다. 그 형태만 보고, 본문 중간은 건드리지 않는다. 버전 표기(v1/v2)와
# 마침표는 라벨 판정에서 뺀다 — 정상 텍스트와 구별되지 않기 때문.
_META_LABEL_WORDS = (
    r"Draft|Revised|Original|Before|After|Correction|Corrected|"
    r"Improved|Final(?:\s*version)?|Candidate\s*[12]|Option\s*[AB]"
)
_META_PREFIX_RE = re.compile(
    rf"^\s*(?:{_META_LABEL_WORDS})\s*:\s*", re.IGNORECASE
)
# 응답 전체가 라벨 비교 형식인지 판정할 때만 쓴다 (자르지 않고 거부용).
_META_LABEL_ANY_RE = re.compile(
    rf"(?:^|[\n\r])\s*(?:{_META_LABEL_WORDS})\s*:", re.IGNORECASE
)


def strip_meta_version_labels(text: str) -> str:
    """
    맨 앞의 "Draft:" / "Revised:" 같은 메타 라벨 **하나만** 떼어낸다.

    프롬프트가 금지해도 LLM이 가끔 버전 비교 형식으로 답하는 것에 대한
    최소한의 안전망이다. 본문을 자르지 않는 것이 핵심 — 잘라내야 할 만큼
    응답이 망가졌다면 그건 여기서 손볼 게 아니라 **응답 자체를 거부**해야
    한다(has_meta_version_labels 참고).
    """
    if not text:
        return text
    return _META_PREFIX_RE.sub("", text, count=1).strip()


def has_meta_version_labels(text: str) -> bool:
    """응답이 "Draft: … Revised: …" 형식인가 — 거부·재시도 판정용."""
    return bool(text) and bool(_META_LABEL_ANY_RE.search(text))
# 메타 라벨 감지 — 이런 label이 들어있으면 LLM이 프롬프트를 무시하고 다양한
# 버전 비교를 응답에 담은 경우. 안전을 위해 그 revision을 통째로 버림.
_QA_META_LABEL_RE = re.compile(
    r"\b(Draft|Revised|Original|Before|After|Correction|Suggestion|Alternative|"
    r"Option\s+[A-Z]|Version\s*\d|v\d\s*[:.]|Candidate\s*\d)\s*[:.]",
    re.IGNORECASE,
)


def parse_qa_response(text: str) -> Dict[int, str]:
    """
    Parse '[N] revised text' blocks from the QA response.

    Multi-line revisions are supported: each [N] header captures everything up
    to the next [N] header (or end of text). Returns {} when the LLM signals
    no changes (output is empty or 'NONE').

    Safety net: if any [N] block contains meta labels (Draft/Revised/Before/
    After/Option A/…) it means the LLM ignored the prompt and returned
    version-comparison text instead of a clean revision. We skip that block
    so the pass-1 translation is kept as-is.
    """
    revisions: Dict[int, str] = {}
    if not text:
        return revisions
    matches = list(_QA_HEADER_RE.finditer(text))
    if not matches:
        return revisions
    for i, m in enumerate(matches):
        idx = int(m.group(1))
        start = m.end()
        end = matches[i + 1].start() if i + 1 < len(matches) else len(text)
        body = text[start:end].strip()
        if not body or body.upper() == "NONE":
            continue
        if _QA_META_LABEL_RE.search(body):
            # LLM이 여러 버전 비교 형식으로 응답 — pass-1 유지 (revision 무시)
            continue
        revisions[idx] = body
    return revisions


def qa_check_batch(
    client: OpenAI,
    items: List[Tuple[int, str, str]],
    style_guide: str,
    glossary_pairs: List[Tuple[str, str]],
    model: str = "gpt-5.2",
) -> Dict[int, str]:
    """
    Pass-2 consistency check on a batch of translated paragraphs.

    Items are (local_index, korean_source, english_translation). The local
    index is what the LLM echoes back in `[N]` headers, so callers should map
    it to their own paragraph identifiers.

    Returns {local_index: revised_translation} containing ONLY items the LLM
    chose to revise. Items not in the result should keep their pass-1
    translation unchanged.
    """
    global TOTAL_INPUT_TOKENS, TOTAL_CACHED_INPUT_TOKENS, TOTAL_OUTPUT_TOKENS, TOTAL_TOKENS

    if not items:
        return {}

    # 앞에서 80개를 자르던 것을 **이 배치에 실제로 등장하는 용어**만 고르는
    # 방식으로 바꿨다. 글로서리가 수백 개일 때 앞 80개를 자르면 정작 지금
    # 검사할 문단에 나오는 용어가 통째로 빠진다 — 일관성 검사인데 근거를
    # 안 주는 셈이었다. 등장하는 것만 주면 개수가 줄고 정확도는 오른다.
    _batch_src = " ".join(src for _, src, _ in items)
    _relevant = [(ko, en) for ko, en in glossary_pairs if ko and ko in _batch_src]
    if not _relevant:
        _relevant = glossary_pairs[:40]        # 하나도 안 걸리면 최소한만
    glossary_block = (
        "\n".join(f"- {ko} -> {en}" for ko, en in _relevant[:120]) or "(none)"
    )
    pairs_block = "\n\n".join(
        f"[{idx}]\nSOURCE: {src}\nTRANSLATION: {tr}" for idx, src, tr in items
    )

    prompt = f"""You are a translation QA reviewer. Each numbered pair below contains a Korean SOURCE and its English TRANSLATION produced in a first pass.

Revise a translation ONLY when it has a clear problem:
- Terminology inconsistent with the document's existing English style
- Sentence pattern that does not match the document's tone
- Awkward, unnatural, or overly literal English (Korean-influenced phrasing)
  e.g. "new-concept platform" → "next-generation platform", "utilization guide" → "user guide"
- Subject-verb or singular/plural agreement error
  e.g. "all user" → "all users", "individual user" → "individual users", "authorized user" → "authorized users"
- Awkward bare generic singular subject
  e.g. "Can user prevent users" → "Can users prevent" or "Can administrators prevent users"
- Missing article where English requires one (Korean has no articles, so the draft may omit them)
- Misspelling, typo, or non-word — especially Korean phonetic transliteration
  e.g. "logram" → "program", any nonsense English output is wrong
- WRONG SENSE of a polysemous Korean word. Several Korean words map to different
  English verbs depending on what they act on. Check the object, not the word:
    · 내보내기 — "export" for data/files/messages, but "remove"/"kick out" for a
      person leaving a group ("Room에서 사용자 내보내기" = remove a user from a Room)
    · 삭제 — "delete" for records, "remove" for a list entry
    · 저장 — "save" for a document, "store" for data at rest
    · 확인 — "check"/"view" for inspecting, "confirm"/"OK" for a dialog button
  A translation that picked the wrong sense is a factual error, not a style
  preference — revise it. e.g. Korean "내보낸 이력" (export history) rendered as
  "History of Removing" is WRONG.
- Missing or duplicated marker, or marker merged with a word
- Glossary term not preserved (see glossary list below) — including the exact casing of case-sensitive terms

Do NOT revise translations that are already acceptable. Do NOT make stylistic preference changes that are not grounded in the style guide. Do NOT shorten or expand for taste. When in doubt, leave it alone.

Strict rules:
- Preserve markers EXACTLY: ⟦B⟧, ⟦/B⟧, ⟦HL:colour⟧, ⟦/HL⟧, ⟦H#⟧, ⟦/H#⟧, ⟦D#⟧, ⟦LB⟧, ⟦TB⟧.
- ⟦LB⟧ is a soft line break that exists in the SOURCE. Your revision must contain
  EXACTLY as many ⟦LB⟧ markers as the TRANSLATION you were given, in the same
  positions. Never remove one to make the text flow as a single line, never add
  one, and never write a real newline in its place. A revision whose ⟦LB⟧ count
  differs from the input will be discarded.
- Glossary translations listed below are mandatory — keep those exact English words with their exact casing.
- Do NOT translate or alter any English already present (other than the specific problems above).

Output format:
- For each item that needs revision: a header line `[N]` followed by the revised translation. Nothing else for that item.
- If nothing needs revision, output exactly: NONE
- Do NOT add explanations, comments, or numbering of any other kind.

ABSOLUTELY FORBIDDEN in your output — if you emit ANY of these your revision will be discarded:
- Meta labels of ANY kind: "Draft:", "Revised:", "Original:", "Before:", "After:",
  "Correction:", "Suggestion:", "Alternative:", "Option A/B", "Version 1/2", "v1/v2".
- Any before/after comparison, side-by-side text, or multiple candidate versions.
- Any preamble like "Here is the revised text" or trailing commentary.
- Real newline characters within a single item's revised text. If the item's
  translation is split across several lines, those splits are represented by
  ⟦LB⟧ markers — keep the markers and stay on one physical line.

Emit ONE clean, final revised text per [N] header. Nothing else.

Style guide:
{style_guide or "(no style guide available)"}

Glossary (mandatory translations):
{glossary_block}

Pairs to review:
{pairs_block}
""".strip()

    resp = client.responses.create(
        model=model,
        input=prompt,
        reasoning={"effort": "low"},
        text={"verbosity": "low"},
    )

    usage = getattr(resp, "usage", None)
    if usage:
        input_tokens = getattr(usage, "input_tokens", 0) or 0
        output_tokens = getattr(usage, "output_tokens", 0) or 0
        total_tokens = getattr(usage, "total_tokens", 0) or 0
        input_details = getattr(usage, "input_tokens_details", None)
        cached_tokens = getattr(input_details, "cached_tokens", 0) or 0
        TOTAL_INPUT_TOKENS += input_tokens
        TOTAL_CACHED_INPUT_TOKENS += cached_tokens
        TOTAL_OUTPUT_TOKENS += output_tokens
        TOTAL_TOKENS += total_tokens

    return parse_qa_response(resp.output_text)


# ─────────────────────────────────────────────────────────────────────────
# 문서 포맷 어댑터
#
# 엔진 본체(글로서리 치환, 마커 복구, QA, 후처리)는 전부 순수 문자열 처리라
# 포맷과 무관하다. 포맷에 묶이는 건 "번역 단위를 어떻게 꺼내고 어떻게 다시
# 써넣느냐" 두 끝단뿐이므로 그것만 어댑터로 분리한다. 중간 언어는 기존
# 마크드 텍스트(⟦B⟧ ⟦G#⟧ ⟦H#⟧ …)를 그대로 쓴다.
# ─────────────────────────────────────────────────────────────────────────

MARKDOWN_EXTENSIONS = (".md", ".markdown", ".mdx")
SUPPORTED_EXTENSIONS = (".docx",) + MARKDOWN_EXTENSIONS


@dataclass
class TransUnit:
    """번역 단위 하나. `ref`는 어댑터가 write에서 쓰는 포맷별 부속물."""
    src: str                 # 마크드 텍스트 — LLM이 보는 유일한 것
    is_heading: bool         # 포맷 고유 판정 (docx: 스타일, md: `#`)
    ref: Any


class DocxAdapter:
    """Word — 기존 동작 그대로. 문단이 곧 번역 단위."""

    def __init__(self, in_path: str):
        self.in_path = in_path
        self.doc = None

    def load(self) -> None:
        self.doc = Document(sanitize_docx(self.in_path))

    def english_samples(self, max_chars: int = 3000) -> List[str]:
        samples: List[str] = []
        for p in iter_all_paragraphs(self.doc):
            text = p.text.strip()
            if text and not contains_korean(text) and len(text) > 15:
                samples.append(text)
            if sum(len(s) for s in samples) >= max_chars:
                break
        return samples

    def units(self) -> List[TransUnit]:
        out: List[TransUnit] = []
        for p in iter_all_paragraphs(self.doc):
            marked, drawing_map, hyperlink_map, trailing_lb = paragraph_to_marked_text(p)
            out.append(TransUnit(
                src=marked,
                is_heading=is_heading_paragraph(p),
                ref=(p, drawing_map, hyperlink_map, trailing_lb),
            ))
        return out

    def write(self, unit: TransUnit, translated: str) -> None:
        p, drawing_map, hyperlink_map, trailing_lb = unit.ref
        _write_paragraph(p, translated, drawing_map, hyperlink_map, trailing_lb)

    def save(self, out_path: str) -> None:
        self.doc.save(out_path)


class MarkdownAdapter:
    """
    Markdown / MDX.

    쓰기를 모아뒀다가 save에서 한 번에 반영한다 — 원문 오프셋 구간을 치환하는
    방식이라 뒤에서부터 적용해야 앞쪽 좌표가 밀리지 않기 때문.
    """

    def __init__(self, in_path: str, keep_heading_anchor: bool = True):
        self.in_path = in_path
        self.keep_heading_anchor = keep_heading_anchor
        self.text = ""
        self.encoding = "utf-8"
        self.newline = "\n"
        self.bom = False
        self.md_units: List[markdown_format.MdUnit] = []
        self._pending: List[Tuple[markdown_format.MdUnit, str]] = []

    def load(self) -> None:
        self.text, self.encoding, self.newline, self.bom = markdown_format.read_text(self.in_path)
        self.md_units = markdown_format.parse_markdown(self.text)
        self._pending = []

    def english_samples(self, max_chars: int = 3000) -> List[str]:
        samples: List[str] = []
        for u in self.md_units:
            text = re.sub(r"⟦[^⟧]*⟧", "", u.src).strip()
            if text and not contains_korean(text) and len(text) > 15:
                samples.append(text)
            if sum(len(s) for s in samples) >= max_chars:
                break
        return samples

    def units(self) -> List[TransUnit]:
        return [TransUnit(src=u.src, is_heading=u.is_heading, ref=u) for u in self.md_units]

    def write(self, unit: TransUnit, translated: str) -> None:
        self._pending.append((unit.ref, translated))

    def save(self, out_path: str) -> None:
        text = markdown_format.apply_translations(
            self.text, self._pending, keep_anchor=self.keep_heading_anchor
        )
        markdown_format.write_text(out_path, text, self.encoding, self.newline, self.bom)


def make_adapter(in_path: str):
    ext = os.path.splitext(in_path)[1].lower()
    if ext in MARKDOWN_EXTENSIONS:
        return MarkdownAdapter(in_path)
    return DocxAdapter(in_path)


def extract_korean_paragraphs(in_path: str) -> List[str]:
    """
    번역 대상이 되는 한국어 문단만 모아 돌려준다.

    글로서리 후보를 뽑으려면 굵은 라벨뿐 아니라 **본문 전체**가 필요하다.
    같은 용어가 문서 곳곳에서 몇 번 반복되는지가 곧 용어성의 근거이기
    때문이다. 마커는 지우고 순수 텍스트만 남긴다 — 빈도만 세면 되므로.
    """
    ext = os.path.splitext(in_path)[1].lower()
    out: List[str] = []
    if ext in MARKDOWN_EXTENSIONS:
        text, _, _, _ = markdown_format.read_text(in_path)
        for u in markdown_format.parse_markdown(text):
            if contains_korean(u.src):
                out.append(u.src)
    else:
        doc = Document(in_path)
        for p in iter_all_paragraphs(doc):
            if contains_korean(p.text):
                out.append(p.text)
    return [strip_zero_width(ALL_MARKER_RE.sub(" ", t)).strip() for t in out]

def extract_bold_terms(in_path: str) -> List[Tuple[str, str]]:
    """
    UI 텍스트 매핑 화면용 — 굵게 표시된 한국어 조각과 그 문맥.

    Word 매뉴얼과 Docusaurus 문서 모두 **굵게 = 화면 라벨** 관습을 쓰므로
    포맷이 달라도 같은 기능이 그대로 성립한다.
    """
    ext = os.path.splitext(in_path)[1].lower()
    if ext not in MARKDOWN_EXTENSIONS:
        return extract_bold_texts_with_context(in_path)

    text, _, _, _ = markdown_format.read_text(in_path)
    seen: set = set()
    result: List[Tuple[str, str]] = []
    for u in markdown_format.parse_markdown(text):
        if u.is_heading:
            continue
        plain = re.sub(r"⟦[^⟧]+⟧", "", u.src)
        for match in _BOLD_SEGMENT_RE.finditer(u.src):
            term = _INNER_MARKER_RE.sub("", match.group(1)).strip()
            if not term or not contains_korean(term) or term in seen:
                continue
            seen.add(term)
            result.append((term, _make_context_excerpt(plain, term)))
    return result


def verify_saved_output(in_path: str, out_path: str,
                        glossary_rows: Optional[List[dict]] = None) -> List[dict]:
    """
    저장이 끝난 문서를 원본과 대조한다. 실패해도 번역을 막지는 않는다.

    반환: [{"level": "오류"|"경고", "title": …, "detail": …}, …]

    번역 파이프라인은 문단을 하나씩 다루므로, 링크·아이콘을 XML에 되꽂은
    **뒤에** 생기는 문장 오류를 구조적으로 볼 수 없다. 그래서 결과물을
    바깥에서 다시 읽는 층을 둔다. output_check는 읽기 전용이라 여기서
    실패해도 산출물에는 영향이 없다.
    """
    if not (out_path.lower().endswith(".docx")
            and in_path.lower().endswith(".docx")):
        return []                     # 마크다운 경로는 대상이 아니다
    try:
        import output_check
    except Exception:
        return []

    rows = glossary_rows or []
    dnt = [_clean(r.get("KO")) for r in rows if _to_bool(r.get("DNT"))]
    literals = [_clean(r.get("EN")) for r in rows
                if _to_bool(r.get("DNT")) or _to_bool(r.get("Case-sensitive"))]
    try:
        findings, _, _ = output_check.verify(
            in_path, out_path,
            dnt=[d for d in dnt if d],
            literals=[l for l in literals if l],
        )
        out = [{"level": f.level, "title": f.title, "detail": f.detail}
               for f in findings]
        # 서로 다른 국문이 같은 영문으로 등재돼 있으면 문서에 "Status, and
        # Status"처럼 구분 불가능한 문장이 생긴다. 데이터 문제라 산출물만
        # 봐서는 원인을 알 수 없으므로 여기서 함께 짚는다.
        for en, kos in output_check.check_duplicate_targets(
            [(_clean(r.get("KO")), _clean(r.get("EN"))) for r in rows]
        ):
            out.append({
                "level": "경고",
                "title": f"서로 다른 용어가 «{en}» 하나로 매핑돼 있습니다",
                "detail": " · ".join(kos)
                          + " — 문서에서 구분되지 않습니다. 글로서리를 나누세요.",
            })
        return out
    except Exception as e:
        return [{"level": "경고", "title": "저장 후 검증을 마치지 못했습니다",
                 "detail": str(e)}]


def translate_document(
    in_path: str,
    out_path: str,
    glossary_rows: List[dict],
    pattern_rows: List[dict],
    api_key: str,
    enable_cache: bool = True,
    enable_qa: bool = True,
    qa_batch_size: int = 8,
    model: str = "gpt-5.2",
    translation_mode: str = "Manual",
    progress_callback: Optional[Callable[[int, int], None]] = None,
    ui_text_overrides: Optional[Dict[str, str]] = None,
) -> Dict[str, int]:
    reset_token_counters()

    glossary_entries = build_glossary_entries_from_rows(glossary_rows)

    # UI 텍스트 매핑은 Glossary와 달리 **⟦B⟧…⟦/B⟧ (bold) 영역 안에만** 적용한다.
    # (Glossary는 본문 어디든 치환. 하지만 UI 라벨은 화면에서 bold이므로
    #  본문에 같은 KO가 나와도 자연 번역을 유지해야 한다.)
    # → glossary_entries에 합치지 않고 별도 dict으로 유지. 각 단락 preprocess
    #   시 preprocess_ui_overrides_in_bold로 별도 치환.
    ui_overrides_clean: Dict[str, str] = {}
    if ui_text_overrides:
        for ko, en in ui_text_overrides.items():
            # _clean()을 쓴다 — 호출자가 pandas DataFrame에서 값을 꺼내며
            # 빈 셀을 str(nan)="nan"으로 만들어 넘기는 사고가 실제로 있었다.
            # "nan"이 치환어로 등록되면 본문에 그대로 찍힌다.
            ko_s, en_s = _clean(ko), _clean(en)
            if ko_s and en_s:
                ui_overrides_clean[ko_s] = en_s
        # 예전에는 UI 매핑과 같은 KO를 glossary_entries에서 **통째로 지웠다**.
        # 그러면 굵게 표시되지 않은 본문에서는 그 용어의 강제가 사라져,
        # 같은 "사용자 및 부서 관리"가 User and Group management /
        # User and department management로 갈렸다.
        # 지우지 않고 **순서**로 푼다 — UI 매핑을 먼저 적용하면 굵은 구간은
        # 이미 자리표시자로 바뀌어 있어 글로서리가 건드릴 수 없고, 나머지
        # 본문에는 글로서리가 그대로 걸린다.

    patterns = build_pattern_pairs_from_rows(pattern_rows)
    # QA prompt엔 UI 매핑도 참고용으로 함께 (일관성 유지 위해)
    glossary_pairs_for_qa = [(e.ko, e.en) for e in glossary_entries] + list(ui_overrides_clean.items())

    client = OpenAI(api_key=api_key)
    adapter = make_adapter(in_path)
    adapter.load()
    cache: Dict[str, str] = {}

    # ── 문서 내 기존 영문 단락 수집 ────────────────────────────────────
    # Pre-pass(extract_doc_style_guide)에서 한 번만 사용되므로 캡을 넉넉히
    # 잡아도 호출당 비용은 고정. v2.0→v2.1처럼 영문이 풍부한 문서에서
    # 가이드 품질을 끌어올리기 위함.
    english_samples = adapter.english_samples()

    # ── Pre-pass: doc-specific style guide 추출 (LLM 1회 호출) ─────────
    style_guide = ""
    if english_samples:
        try:
            style_guide = extract_doc_style_guide(client, english_samples, model=model)
        except Exception:
            # 가이드 추출 실패는 치명적이지 않음 — 가이드 없이 진행
            style_guide = ""

    # 번역 대상 유닛 — 한국어가 없는 단락/블록은 여기서 걸러진다.
    # 마크다운에서는 인라인 봉인이 이 판정보다 먼저 끝나 있어야 한다
    # (`![img](/img/분석실행.png)` 처럼 경로에만 한글이 있는 경우 때문).
    units: List[TransUnit] = [u for u in adapter.units() if contains_korean(u.src)]

    # heading 여부는 여기서 한 번만 확정한다 — 포맷 고유 판정(스타일/`#`)에
    # 텍스트 휴리스틱을 OR로 얹은 값을 Pass 1·2가 공유한다.
    for u in units:
        u.is_heading = u.is_heading or looks_like_heading_text(u.src)

    total_paras = len(units)

    if total_paras == 0:
        adapter.save(out_path)
        return {
            "input_tokens": TOTAL_INPUT_TOKENS,
            "cached_tokens": TOTAL_CACHED_INPUT_TOKENS,
            "output_tokens": TOTAL_OUTPUT_TOKENS,
            "total_tokens": TOTAL_TOKENS,
            "paragraphs_translated": 0,
        }

    # ── 진행률 총량 사전 계산 ─────────────────────────────────────────
    # Pass 2(QA) 대상은 "src 첫 등장이고 heading이 아닌" 단락. Pass 1과
    # 동일한 필터링 규칙을 미리 돌려서 Pass 1/2 합산 진행률을 확정한다.
    qa_estimated = 0
    if enable_qa:
        seen_for_estimate = set()
        for u in units:
            key = (u.src, bool(u.is_heading))
            if key in seen_for_estimate:
                continue
            seen_for_estimate.add(key)
            qa_estimated += 1
    total_work = total_paras + qa_estimated

    # ── Pass 1: 번역 (쓰기는 미루고 메모리에 누적) ────────────────────
    pass1_results: List[Dict] = []

    for idx, unit in enumerate(units):
        src = unit.src
        # heading 여부는 위에서 확정됨 — UI 매핑 적용 여부와 case-sensitive
        # 재복원 여부를 이 값으로 가른다.
        is_heading = unit.is_heading

        if enable_cache and src in cache:
            translated = cache[src]
        else:
            # 1) UI 텍스트 매핑을 **먼저** — 오직 굵은 구간, non-heading에서만.
            #    Heading은 무조건 sentence case로 정규화될 예정이므로 UI 매핑을
            #    적용해봐야 결국 case가 바뀌어 어색해질 수 있다.
            ui_map: Dict[str, GlossaryEntry] = {}
            next_idx = 0
            ui_pre = src
            if not is_heading:
                ui_pre, ui_map, next_idx = preprocess_ui_overrides_in_bold(
                    src, ui_overrides_clean, start_idx=0
                )
            # 2) 남은 텍스트에 Glossary 치환 (본문 전체).
            #    UI가 잡은 자리는 이미 ⟦G#⟧이므로 여기서 다시 걸리지 않는다.
            gl_pre, gl_map = preprocess_with_glossary_placeholders(
                ui_pre, glossary_entries, start_idx=next_idx
            )
            if ui_map:
                gl_map = {**ui_map, **gl_map}

            # 원문 영문 봉인 — 여기부터 복원 전까지 AI/API/v2.1/K-Assistant는
            # ⟦X#⟧로 가려져 있어 모델도, 우리 후처리도 건드릴 수 없다.
            gl_pre, lit_map = seal_literals(gl_pre)

            selected_pattern_examples = select_relevant_patterns(gl_pre, patterns)

            translated = translate_marked_paragraph(
                client=client,
                source_text=gl_pre,
                pattern_examples=selected_pattern_examples,
                model=model,
                translation_mode=translation_mode,
                style_reference=style_guide,
            )

            # 마커가 망가진 채 후처리로 넘어가면 그 뒤로는 복구가 추측이
            # 된다. 값이 싼 단계에서 한 번 더 물어보는 편이 낫다.
            if check_marker_integrity(gl_pre, translated):
                retry = translate_marked_paragraph(
                    client=client,
                    source_text=gl_pre,
                    pattern_examples=selected_pattern_examples,
                    model=model,
                    translation_mode=translation_mode,
                    style_reference=style_guide,
                )
                if retry.strip() and len(check_marker_integrity(gl_pre, retry)) \
                        < len(check_marker_integrity(gl_pre, translated)):
                    translated = retry

            # 0-a) 폭 없는 문자 제거 — **모든 후처리보다 먼저** 해야 한다.
            #      모델이 문장 끝에 U+FEFF를 붙여 보내면 restore_sentence_period가
            #      "마침표가 없다"고 오판해 마침표를 하나 더 붙이는 식으로,
            #      보이지도 않는 문자가 텍스트 판정을 줄줄이 어긋나게 만든다.
            translated = strip_zero_width(translated).strip()

            # 0) 줄 구조 정리 — 모델이 ⟦LB⟧ 대신 개행으로 답했거나 여분을
            #    붙였을 때 원문 기준으로 맞춘다. 이후 후처리들이 ⟦LB⟧를 "줄"로
            #    인식하므로 가장 먼저 수행해야 한다.
            translated = normalize_paragraph_breaks(translated, gl_pre)
            translated = enforce_line_breaks(translated, gl_pre)

            # 1) marker 복구
            translated = repair_bold_markers(translated)
            translated = repair_hl_markers(translated)

            # 1-b) bold 경계 공백 제거
            translated = normalize_marker_boundary_spaces(translated)

            # 1-c) HL 마커 누락 폴백
            translated = apply_highlight_fallback(translated, src)

            # 2) glossary 복원 (위치 기반 대소문자)
            translated = restore_glossary_placeholders(translated, gl_map or {})

            # 3) colon label normalize
            translated = normalize_colon_label_line(translated)

            # 4) 남은 한국어 fallback 번역
            if contains_korean(translated):
                translated = translate_remaining_korean(client, translated, model=model)
                translated = strip_zero_width(translated)
                translated = enforce_line_breaks(translated, gl_pre)
                translated = repair_bold_markers(translated)
                translated = repair_hl_markers(translated)
                translated = normalize_marker_boundary_spaces(translated)
                translated = apply_highlight_fallback(translated, src)
                translated = restore_glossary_placeholders(translated, gl_map or {})
                translated = normalize_colon_label_line(translated)

            # 5) heading / 일반 문단 후처리 (is_heading은 위에서 계산됨)
            if is_heading:
                translated = normalize_heading_text(translated)
                translated = normalize_ui_label_text(translated)
                translated = _cap_first_alpha(translated)
                translated = translated.rstrip()
                if translated.endswith("."):
                    translated = translated[:-1]
            else:
                translated = normalize_ui_in_bold_segments(translated)
                translated = _cap_first_alpha(translated)

            # 6) 마지막 품질 보정
            translated = capitalize_bullet_lines(translated)
            translated = restore_sentence_period(translated, src)
            translated = normalize_paragraph_breaks(translated, src)
            # pass-1이 "Draft: X ... Revised: Y" 형태를 뱉는 경우 방어
            translated = strip_meta_version_labels(translated)
            # 후처리 단계에서 줄 구조가 흐트러졌으면 원문 기준으로 최종 정렬
            translated = enforce_line_breaks(translated, src)
            # 6-b) 후처리(대소문자 정규화 등)가 마커를 훼손했을 경우 최종 복구.
            #      깨진 마커는 write-back 토크나이저가 인식하지 못해 문서에
            #      그대로 출력되므로 XML 기록 직전에 한 번 더 정규화한다.
            translated = repair_hl_markers(translated)

            # 7) case-sensitive/DNT glossary 용어의 정확한 대소문자 복원.
            #    Heading에도 적용한다. 원칙은 여전히 "heading은 sentence case"
            #    지만, 사용자가 글로서리에 **명시적으로 case-sensitive로 등록한**
            #    용어는 그 표기가 곧 의도다. URL·API 같은 약어를 엔진이 알아서
            #    대문자화하지는 않고(그건 사용자 몫), 등록된 것만 존중한다.
            translated = enforce_case_sensitive_glossary(translated, glossary_entries)

            # 8) 봉인 해제. 반드시 대소문자 정규화가 **모두 끝난 뒤**여야 한다.
            #    먼저 풀면 heading sentence-case가 "AI"를 "ai"로 내려버린다.
            translated = restore_literals(translated, lit_map)

            if enable_cache:
                cache[src] = translated

        pass1_results.append({
            "unit": unit,
            "src": src,
            "translated": translated,
        })

        if progress_callback:
            progress_callback(idx + 1, total_work)

    # ── Pass 2: batch QA (일관성 검사) ───────────────────────────────
    # src별 그룹핑으로 동일 원문은 한 번만 QA → 결과를 모든 사본에 적용.
    if enable_qa:
        # 그룹키에 heading 여부를 넣는다. 예전엔 src만 썼는데, 같은 문장이
        # 제목으로도 본문으로도 쓰이면 한 그룹이 되고 대표(group[0])가 제목이면
        # **본문까지 통째로 QA에서 빠졌다**. 제목과 본문은 대소문자 규칙부터
        # 다르므로 애초에 같은 리비전을 공유해서도 안 된다.
        src_groups: Dict[Tuple[str, bool], List[Dict]] = defaultdict(list)
        for r in pass1_results:
            src_groups[(r["src"], bool(r["unit"].is_heading))].append(r)

        # 제목도 QA한다. 예전엔 건너뛰었는데, 그 탓에 "Reset an user password",
        # "Save ai provider information" 같은 제목 오류가 그대로 남았다.
        # 제목이야말로 눈에 가장 먼저 띄는 자리다.
        qa_items: List[Dict] = []
        for (src, _is_heading), group in src_groups.items():
            qa_items.append({
                "src": src,
                "translated": group[0]["translated"],
                "group": group,
                "is_heading": _is_heading,
            })

        qa_done = 0
        for batch_start in range(0, len(qa_items), qa_batch_size):
            batch = qa_items[batch_start:batch_start + qa_batch_size]
            batch_input = [
                (i, item["src"], item["translated"])
                for i, item in enumerate(batch)
            ]

            try:
                revisions = qa_check_batch(
                    client=client,
                    items=batch_input,
                    style_guide=style_guide,
                    glossary_pairs=glossary_pairs_for_qa,
                    model=model,
                )
            except Exception:
                # QA 실패 시 batch 전체 skip — Pass 1 결과 유지
                revisions = {}

            for i, item in enumerate(batch):
                if i not in revisions:
                    continue
                # Pass 1과 같은 이유로 후처리 전에 폭 없는 문자를 먼저 제거한다.
                revised = strip_zero_width(revisions[i]).strip()
                if not revised:
                    continue  # 안전장치: 빈 응답으로 덮어쓰지 않음
                # Pass 2 응답에도 동일한 marker 후처리를 적용
                revised = repair_bold_markers(revised)
                revised = repair_hl_markers(revised)
                revised = normalize_marker_boundary_spaces(revised)
                revised = apply_highlight_fallback(revised, item["src"])
                # QA 응답에서도 line break/meta label 정리 (pass-1과 동일 수준으로)
                revised = normalize_paragraph_breaks(revised, item["src"])
                revised = strip_meta_version_labels(revised)
                revised = enforce_line_breaks(revised, item["src"])
                # 줄 구조를 되살릴 수 없는 리비전은 채택하지 않는다 — QA는
                # 표현 다듬기가 목적이므로, 줄바꿈을 잃는 대가로 받아들일
                # 만한 개선은 없다. Pass 1 결과를 그대로 유지한다.
                if not line_breaks_match(revised, item["src"]):
                    continue
                # 제목은 sentence case 규칙을 다시 씌운다 — QA가 본문 기준으로
                # 고쳐 놓으면 제목만 표기 규칙이 어긋난다.
                if item.get("is_heading"):
                    revised = normalize_heading_text(revised)
                    revised = normalize_ui_label_text(revised)
                    revised = _cap_first_alpha(revised).rstrip()
                    if revised.endswith("."):
                        revised = revised[:-1]
                # QA가 case-sensitive 용어를 흐트러뜨리는 경우도 복원
                revised = enforce_case_sensitive_glossary(revised, glossary_entries)
                # Pass 2는 봉인 바깥에서 돌므로 원문 영문 대소문자를 여기서 지킨다
                revised = enforce_literal_casing(revised, item["src"])
                for r in item["group"]:
                    r["translated"] = revised
                    if enable_cache:
                        cache[item["src"]] = revised

            qa_done += len(batch)
            if progress_callback:
                progress_callback(total_paras + qa_done, total_work)

    # ── Final: 모든 결과를 한 번에 쓰기 ───────────────────────────────
    # 여기가 문서로 나가는 유일한 통로다. 마커 검사를 다른 데 두면 이후
    # 후처리가 또 망가뜨릴 수 있으므로, 관문은 반드시 이 자리여야 한다.
    marker_repaired = marker_failed = 0
    for r in pass1_results:
        text, problems = finalize_markers(r["unit"].src, r["translated"])
        if problems:
            marker_failed += 1
        elif text != r["translated"]:
            marker_repaired += 1
        adapter.write(r["unit"], text)

    adapter.save(out_path)

    # ── 저장 후 자체 검증 ────────────────────────────────────────────
    # 여기까지의 검사는 전부 **문단 단위**라, 링크와 아이콘을 Word에 다시
    # 조립한 뒤 만들어지는 문장은 아무도 보지 않았다. "websiteAccess it."
    # 같은 오류가 그 틈에서 나온다. 저장된 문서를 다시 읽어 문장을 본다.
    verification = verify_saved_output(in_path, out_path, glossary_rows)

    return {
        "input_tokens": TOTAL_INPUT_TOKENS,
        "cached_tokens": TOTAL_CACHED_INPUT_TOKENS,
        "output_tokens": TOTAL_OUTPUT_TOKENS,
        "total_tokens": TOTAL_TOKENS,
        "paragraphs_translated": total_paras,
        "marker_repaired": marker_repaired,
        "marker_failed": marker_failed,
        "verification": verification,
    }