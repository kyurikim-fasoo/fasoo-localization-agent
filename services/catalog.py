"""
i18n 카탈로그 → 글로서리 / UI 라벨 후보 추출.

전제:
    제품의 다국어 메시지 카탈로그는 key로 정렬돼 있다. 즉 "어느 한국어가 어느
    영어에 대응하는가"라는 정렬 문제가 애초에 없다. 그래서 이 모듈은 LLM을
    쓰지 않는다 — join·집계·분류만으로 후보가 나온다. 판단이 필요한 지점
    (1:N 충돌 대표형 선택, 용어성 판정)은 사람이 화면에서 고른다.

지원 포맷:
    고객사마다 파일 생김새가 다르므로 아래를 자동 감지한다.
      A. 배열 + 언어 필드   [{"key":"a.b", "en":"Save"}]
      B. 배열 + 양쪽 언어   [{"key":"a.b", "en":"Save", "ko":"저장"}]
      C. 평탄 객체          {"a.b": "저장"}          (i18next/vue-i18n)
      D. 중첩 객체          {"a": {"b": "저장"}}     → 경로를 key로 평탄화
      E. 파일 1개 / 2개     위 어느 것이든

    F. Word 매뉴얼 쌍   같은 버전의 국문본/영문본 .docx 두 개
                       → 제목(Heading) 순서를 앵커로 블록을 맞춰 key를 만든다.
                         parse_docx_pair() 참고.

    어느 컬럼이 한국어인지는 파일명이 아니라 **한글 비율**로 판정한다.
    파일명에 ko/en 표시가 없어도 동작하게 하기 위함.
"""
from __future__ import annotations

import io
import re
from collections import Counter, defaultdict
from dataclasses import dataclass, field
from difflib import SequenceMatcher
from typing import Any, Dict, List, Optional, Tuple

import pandas as pd

_KOREAN_RE = re.compile(r"[가-힣]")
# key로 쓰일 법한 필드명 (우선순위 순)
_KEY_FIELD_HINTS = ("key", "id", "msgid", "code", "name", "property")
# 값 필드가 아님이 분명한 것 — 메타데이터
_META_FIELDS = ("comment", "description", "context", "note", "updated", "created")

LABEL_MAX_KO_CHARS = 25
LABEL_MAX_EN_WORDS = 5

# `{0}` `%s` `{{name}}` — 값이 끼워지는 메시지 템플릿. 용어가 아니라 패턴이다.
_PLACEHOLDER_RE = re.compile(r"\{\d+\}|\{\{|%[sd]\b")
# 한국어 종결어미 — 마침표가 없어도 문장인 경우를 잡는다.
#   강한 표지는 위치를 가리지 않고 찾는다. 실제 카탈로그에는
#   "시간을 입력하세요. (단위: 초)" 처럼 **문장이 끝난 뒤 괄호가 붙는** 항목이
#   흔해서, 끝만 봐서는 놓친다. UI 라벨에 '습니다/하세요'가 들어갈 일은 없다.
_SENT_STRONG_RE = re.compile(r"(습니다|합니다|입니다|하세요|하십시오|십시오)")
#   약한 표지는 다른 단어 안에 우연히 들어갈 수 있어 끝에서만 본다.
_SENT_TAIL_RE = re.compile(r"(한다|된다|세요|어요|아요|겠다|하라|해라)$")
# 조사로 끝나면 용어가 아니라 문장 조각이다 ("…취약점 검출부터").
_JOSA_TAIL_RE = re.compile(r"(부터|까지|에서|으로|에게|보다|처럼|만큼|이며|하며)$")


def _ko_core(ko: str) -> str:
    """끝의 문장부호·괄호·따옴표를 털어낸 몸통 — 종결어미 판정용."""
    return ko.strip().rstrip(").]:;,\"'”’ ").strip()


# ──────────────────────────────────────────────────────────────────────
# 파싱
# ──────────────────────────────────────────────────────────────────────

@dataclass
class ParsedFile:
    """파일 하나를 '컬럼 여러 개'로 정규화한 결과."""
    name: str
    shape: str                              # array | flat | nested
    key_field: str
    columns: Dict[str, Dict[str, str]]      # 컬럼명 -> {key: text}
    n_keys: int
    duplicate_keys: int = 0

    def describe(self) -> str:
        cols = ", ".join(self.columns)
        shape_ko = {"array": "배열형", "flat": "평탄 객체",
                    "nested": "중첩 객체", "docx": "Word 매뉴얼"}[self.shape]
        dup = f" · 중복키 {self.duplicate_keys:,}" if self.duplicate_keys else ""
        return f"{shape_ko} · key={self.key_field} · 컬럼[{cols}] · {self.n_keys:,}개{dup}"


def _flatten(obj: Any, prefix: str = "") -> Dict[str, str]:
    """중첩 객체를 'a.b.c' 경로 key로 평탄화. 문자열 값만 남긴다."""
    out: Dict[str, str] = {}
    if isinstance(obj, dict):
        for k, v in obj.items():
            out.update(_flatten(v, f"{prefix}.{k}" if prefix else str(k)))
    elif isinstance(obj, list):
        for i, v in enumerate(obj):
            out.update(_flatten(v, f"{prefix}[{i}]"))
    elif isinstance(obj, str):
        out[prefix] = obj
    return out


def _pick_key_field(rows: List[dict]) -> Optional[str]:
    """배열형에서 key 역할을 하는 필드 찾기."""
    fields = [f for f in rows[0] if isinstance(rows[0].get(f), (str, int))]
    for hint in _KEY_FIELD_HINTS:
        for f in fields:
            if f.lower() == hint:
                return f
    # 힌트가 없으면 '값이 거의 다 고유하고 한글이 없는' 필드를 고른다
    best, best_score = None, 0.0
    for f in fields:
        vals = [str(r.get(f, "")) for r in rows if r.get(f)]
        if not vals:
            continue
        uniq = len(set(vals)) / len(vals)
        head = vals[:500]
        ko = sum(1 for v in head if _KOREAN_RE.search(v)) / len(head)
        score = uniq * (1 - ko)
        if score > best_score:
            best, best_score = f, score
    return best if best_score > 0.8 else None


def parse_json(name: str, data: Any) -> ParsedFile:
    """JSON 한 덩어리를 ParsedFile로. 형태를 인식하지 못하면 ValueError."""
    if isinstance(data, list):
        rows = [r for r in data if isinstance(r, dict)]
        if not rows:
            raise ValueError("배열이지만 객체가 없습니다.")
        key_field = _pick_key_field(rows)
        if key_field is None:
            raise ValueError("key 역할을 하는 필드를 찾지 못했습니다.")
        value_fields = [
            f for f in rows[0]
            if f != key_field
            and isinstance(rows[0].get(f), str)
            and f.lower() not in _META_FIELDS
        ]
        if not value_fields:
            raise ValueError("텍스트 값 필드가 없습니다.")
        columns: Dict[str, Dict[str, str]] = {f: {} for f in value_fields}
        seen: Counter = Counter()
        for r in rows:
            k = str(r.get(key_field, "")).strip()
            if not k:
                continue
            seen[k] += 1
            for f in value_fields:
                v = r.get(f)
                if isinstance(v, str) and v.strip():
                    columns[f][k] = v
        dupes = sum(1 for c in seen.values() if c > 1)
        return ParsedFile(name, "array", key_field, columns, len(seen), dupes)

    if isinstance(data, dict):
        flat = _flatten(data)
        if not flat:
            raise ValueError("문자열 값을 찾지 못했습니다.")
        # 실제로 계층이 있었는지로 판단한다. 평탄 객체도 key에 점이 들어가므로
        # ("a.save") 결과 key만 봐서는 구분되지 않는다.
        nested = any(isinstance(v, (dict, list)) for v in data.values())
        return ParsedFile(
            name, "nested" if nested else "flat", "(경로)", {name: flat}, len(flat)
        )

    raise ValueError("최상위가 배열도 객체도 아닙니다.")


# ──────────────────────────────────────────────────────────────────────
# Word 매뉴얼 쌍 (국문본 + 영문본)
#
# JSON 카탈로그와 달리 매뉴얼에는 key가 없다. 그래서 **구조**로 key를 만든다.
# 같은 버전의 국/영문 매뉴얼은 같은 템플릿에서 나오므로 제목(Heading) 순서가
# 그대로 보존된다 — 이걸 앵커로 삼아 구간을 자르고, 구간 안에서는 블록을
# 순서대로 맞춘다. 문장 단위로 끊지 않으므로 "국문 1문장 = 영문 2문장"은
# 문제가 되지 않는다. 단락 경계는 번역돼도 거의 안 바뀌기 때문이다.
#
# 볼드 런은 화면 라벨이라 가장 값진 후보지만, **위치로 짝지으면 안 된다**.
# 한 단락에 볼드가 여럿이면 한/영 어순이 뒤집히기 때문이다.
#     KO  「푸시 알림」에서 「비활성화」를 선택할 경우
#     EN  If you select 「Disable」 for 「Push notifications」
# 그대로 zip하면 정반대 쌍이 나온다. 그래서 양쪽 1개씩인 단락만 직접
# 확정하고, 나머지는 문서 전체의 공기(co-occurrence) 빈도로 매칭한다.
# ──────────────────────────────────────────────────────────────────────

# Word는 한 단어도 여러 run으로 쪼갠다 ("Fireside Admin" + "istrator").
# 그래서 인접한 굵은 런은 반드시 먼저 합쳐야 한다.
_LABEL_TRIM = " :\u00b7-.,\u201c\u201d\u2018\u2019\"\'"
# 공기 매칭 신뢰 임계값 (Dice 계수). 낮추면 수확은 늘지만 오매칭이 섞인다.
DOCX_COOC_MIN = 0.5
# 1등이 2등보다 이만큼은 앞서야 채택한다. 다중 볼드 단락이 문서에 한 번만
# 나오면 후보들이 동점이 되는데, 그때 찍으면 정반대 쌍이 등재된다.
# 글로서리를 오염시키느니 버린다.
DOCX_COOC_MARGIN = 1.2


@dataclass
class DocxAlignment:
    """두 매뉴얼을 맞춘 결과. 화면에 "얼마나 믿을 만한가"를 보여주기 위한 통계."""
    name_a: str
    name_b: str
    headings_a: int = 0
    headings_b: int = 0
    heading_matched: int = 0
    sections: int = 0
    sections_matched: int = 0
    mismatched: List[Tuple[str, int, str, int]] = field(default_factory=list)
    front_dropped: int = 0
    n_heading: int = 0
    n_block: int = 0
    n_bold_direct: int = 0
    n_bold_cooc: int = 0

    @property
    def heading_total(self) -> int:
        return max(self.headings_a, self.headings_b)

    @property
    def heading_rate(self) -> float:
        return self.heading_matched / self.heading_total if self.heading_total else 0.0

    @property
    def section_rate(self) -> float:
        return self.sections_matched / self.sections if self.sections else 0.0

    @property
    def n_label(self) -> int:
        return self.n_bold_direct + self.n_bold_cooc

    @property
    def ok(self) -> bool:
        """이 정렬을 믿고 후보를 뽑아도 되는가."""
        return self.heading_rate >= 0.8 and self.section_rate >= 0.7

    def describe(self) -> str:
        return (f"제목 {self.heading_matched}/{self.heading_total} · "
                f"구간 {self.sections_matched}/{self.sections} · "
                f"라벨 {self.n_label}쌍")


def _docx_label(t: str) -> str:
    return re.sub(r"\s+", " ", t).strip(_LABEL_TRIM).strip()


def _para_lines(p) -> List[Tuple[str, List[str]]]:
    """
    단락 하나를 논리 줄들로. 각 줄은 (본문, 볼드조각들).

    `<w:br/>`로 줄을 나눈다 — 한쪽 언어만 세 단락을 줄바꿈으로 한 단락에
    몰아넣은 경우가 실제로 있어서, 이걸 풀지 않으면 그 구간이 통째로 버려진다.
    """
    lines: List[List[Tuple[str, bool]]] = [[]]
    for r in p.runs:
        if not r.text:
            continue
        for i, part in enumerate(r.text.split("\n")):
            if i:
                lines.append([])
            if part:
                lines[-1].append((part, bool(r.bold)))

    out: List[Tuple[str, List[str]]] = []
    for frags in lines:
        text = re.sub(r"\s+", " ", "".join(t for t, _ in frags)).strip()
        if not text:
            continue
        bolds, cur = [], []
        for t, is_bold in frags:
            if is_bold:
                cur.append(t)
            elif cur:
                bolds.append("".join(cur))
                cur = []
        if cur:
            bolds.append("".join(cur))
        out.append((text, [b for b in map(_docx_label, bolds) if b]))
    return out


@dataclass
class _Block:
    kind: str                 # h(제목) | p(본문) | tc(표 셀)
    level: int                # 제목 레벨, 그 외 0
    text: str
    bolds: List[str]
    cell: Optional[str] = None


def _docx_blocks(data: bytes) -> List[_Block]:
    """문서 순서대로 블록 추출 — 단락과 표 셀을 원래 순서 그대로."""
    from docx import Document                    # 무거워서 여기서만 import
    from docx.oxml.ns import qn

    doc = Document(io.BytesIO(data))
    pmap = {p._p: p for p in doc.paragraphs}
    tmap = {t._tbl: t for t in doc.tables}
    out: List[_Block] = []
    ti = 0
    for child in doc.element.body.iterchildren():
        if child.tag == qn("w:p"):
            p = pmap.get(child)
            if p is None:
                continue
            m = re.match(r"Heading (\d)", p.style.name or "")
            for text, bolds in _para_lines(p):
                out.append(_Block("h" if m else "p",
                                  int(m.group(1)) if m else 0, text, bolds))
        elif child.tag == qn("w:tbl"):
            tbl = tmap.get(child)
            if tbl is None:
                continue
            for ri, row in enumerate(tbl.rows):
                for ci, cell in enumerate(row.cells):
                    text = re.sub(r"\s+", " ", cell.text).strip()
                    if not text:
                        continue
                    bolds = [b for pp in cell.paragraphs
                             for _, bs in _para_lines(pp) for b in bs]
                    out.append(_Block("tc", 0, text, bolds, f"t{ti}r{ri}c{ci}"))
            ti += 1
    return out


def _cut_front(blocks: List[_Block]) -> Tuple[List[_Block], int]:
    """
    첫 Heading 1 이전을 버린다.

    앞표지·법적 고지·연락처는 **번역물이 아니라 지역별 별도 원고**다. 실제
    문서에서 국문은 "기술 지원/원격 지원/제품 문의", 영문은 "Sales
    information/Help desk/North America HQ"로 항목도 전화번호도 다르다.
    여기를 정렬하면 "기술 지원 ↔ Sales information" 같은 오염된 쌍이 들어간다.
    """
    for i, b in enumerate(blocks):
        if b.kind == "h" and b.level == 1:
            return blocks[i:], i
    return blocks, 0


def _sections(blocks: List[_Block]) -> List[Tuple[_Block, List[_Block]]]:
    secs: List[Tuple[_Block, List[_Block]]] = []
    for b in blocks:
        if b.kind == "h":
            secs.append((b, []))
        elif secs:
            secs[-1][1].append(b)
    return secs


def _match_bolds(pairs: List[Tuple[_Block, _Block]]) -> List[Tuple[str, str, bool]]:
    """
    정렬된 블록 쌍들에서 볼드 라벨 쌍을 뽑는다. (a, b, 직접확정여부)

    양쪽 1개씩인 블록은 그대로 확정. 여러 개인 블록은 어순이 뒤집힐 수 있으니
    문서 전체에서 함께 등장한 횟수를 세어 Dice 계수로 상호 최적 매칭한다.
    라벨은 매뉴얼에서 수십 번 반복되므로 이 통계가 잘 먹는다.
    """
    direct: List[Tuple[str, str]] = []
    cooc: Dict[str, Counter] = defaultdict(Counter)
    fa: Counter = Counter()
    fb: Counter = Counter()

    for x, y in pairs:
        ba, bb = x.bolds, y.bolds
        if not ba or not bb:
            continue
        if len(ba) == 1 and len(bb) == 1:
            direct.append((ba[0], bb[0]))
        else:
            for a in ba:
                for b in bb:
                    cooc[a][b] += 1
        for a in ba:
            fa[a] += 1
        for b in bb:
            fb[b] += 1

    # 1:1 단락에서 이미 확정된 쌍은 고정한다. 같은 라벨이 문서 어딘가에서
    # 단독으로 등장했다면 그게 가장 확실한 근거이므로, 애매한 단락에서
    # 다시 후보로 놓고 흔들 이유가 없다.
    locked_a: Dict[str, str] = {}
    for a, b in direct:
        locked_a.setdefault(a, b)
    locked_b = set(locked_a.values())

    best: Dict[str, Tuple[float, str]] = {}
    for a, cands in cooc.items():
        if a in locked_a:
            continue
        ranked = sorted(
            ((2 * n / (fa[a] + fb[b]), b) for b, n in cands.items() if b not in locked_b),
            reverse=True,
        )
        if not ranked:
            continue
        top, runner = ranked[0], (ranked[1] if len(ranked) > 1 else (0.0, ""))
        # 동점(=근거 없음)이면 채택하지 않는다
        if top[0] < DOCX_COOC_MIN or top[0] < runner[0] * DOCX_COOC_MARGIN:
            continue
        best[a] = top

    # 상호 최적일 때만 채택 — 여러 KO가 한 EN에 몰리는 걸 막는다
    rev: Dict[str, List[Tuple[float, str]]] = defaultdict(list)
    for a, (score, b) in best.items():
        rev[b].append((score, a))
    matched = [(a, b) for a, (score, b) in best.items() if max(rev[b])[1] == a]

    return ([(a, b, True) for a, b in direct]
            + [(a, b, False) for a, b in matched])


def parse_docx_pair(name_a: str, data_a: bytes,
                    name_b: str, data_b: bytes
                    ) -> Tuple[ParsedFile, ParsedFile, DocxAlignment]:
    """
    같은 버전의 Word 매뉴얼 두 개(국문본/영문본)를 key를 공유하는 ParsedFile
    쌍으로. 어느 쪽이 한국어인지는 여기서 정하지 않는다 — pick_languages()가
    한글 비율로 판정하므로 이 함수는 언어에 대해 대칭이다.
    """
    ba, front_a = _cut_front(_docx_blocks(data_a))
    bb, front_b = _cut_front(_docx_blocks(data_b))
    sa, sb = _sections(ba), _sections(bb)
    if not sa or not sb:
        empty = name_a if not sa else name_b
        raise ValueError(
            f"`{empty}` 에서 제목을 찾지 못했습니다. 두 매뉴얼 모두 Word의 "
            "**제목 1~4** 스타일이 적용된 제목이 있어야 구간을 맞출 수 있습니다. "
            "글자 크기만 키운 제목은 인식되지 않습니다."
        )
    align = DocxAlignment(
        name_a=name_a, name_b=name_b,
        headings_a=len(sa), headings_b=len(sb),
        front_dropped=front_a + front_b,
    )

    # 제목 시퀀스는 레벨 패턴으로 맞춘다 — 언어가 달라 본문 비교가 안 되기 때문.
    # 한쪽에만 있는 절이 있어도 SequenceMatcher가 그 구간만 건너뛴다.
    sm = SequenceMatcher(None, [h.level for h, _ in sa], [h.level for h, _ in sb])
    sec_pairs = []
    for i, j, n in sm.get_matching_blocks():
        for k in range(n):
            sec_pairs.append((sa[i + k], sb[j + k]))
    align.heading_matched = len(sec_pairs)
    align.sections = max(len(sa), len(sb))

    col_a: Dict[str, str] = {}
    col_b: Dict[str, str] = {}
    block_pairs: List[Tuple[_Block, _Block]] = []

    for si, ((ha, body_a), (hb, body_b)) in enumerate(sec_pairs):
        col_a[f"h{si}"] = ha.text
        col_b[f"h{si}"] = hb.text
        align.n_heading += 1
        if len(body_a) != len(body_b):
            align.mismatched.append((ha.text, len(body_a), hb.text, len(body_b)))
            continue
        align.sections_matched += 1
        for bi, (x, y) in enumerate(zip(body_a, body_b)):
            # 표 셀은 좌표까지 같을 때만 — 표 모양이 다르면 짝이 어긋난다
            if (x.kind == "tc") != (y.kind == "tc"):
                continue
            if x.kind == "tc" and x.cell != y.cell:
                continue
            key = f"s{si}b{bi}"
            col_a[key] = x.text
            col_b[key] = y.text
            align.n_block += 1
            block_pairs.append((x, y))

    for n, (a, b, is_direct) in enumerate(_match_bolds(block_pairs)):
        col_a[f"b{n}"] = a
        col_b[f"b{n}"] = b
        if is_direct:
            align.n_bold_direct += 1
        else:
            align.n_bold_cooc += 1

    pf_a = ParsedFile(name_a, "docx", "(블록)", {name_a: col_a}, len(col_a))
    pf_b = ParsedFile(name_b, "docx", "(블록)", {name_b: col_b}, len(col_b))
    return pf_a, pf_b, align


# ──────────────────────────────────────────────────────────────────────
# 언어 판정
# ──────────────────────────────────────────────────────────────────────

def korean_ratio(values: List[str], sample: int = 800) -> float:
    vals = [v for v in values if v.strip()][:sample]
    if not vals:
        return 0.0
    return sum(1 for v in vals if _KOREAN_RE.search(v)) / len(vals)


@dataclass
class LanguagePick:
    ko: Dict[str, str]
    en: Dict[str, str]
    ko_label: str
    en_label: str
    ratios: List[Tuple[str, float]] = field(default_factory=list)


def pick_languages(files: List[ParsedFile]) -> LanguagePick:
    """
    모든 (파일, 컬럼) 중에서 한국어 쪽과 영어 쪽을 고른다.

    파일명이 아니라 한글 비율로 판정한다 — 고객사 파일에 ko/en 표시가 없을 수
    있기 때문. 한국어는 비율이 가장 높은 컬럼, 영어는 가장 낮은 컬럼.
    """
    cand: List[Tuple[str, Dict[str, str], float]] = []
    for pf in files:
        for col, mapping in pf.columns.items():
            label = f"{pf.name}:{col}" if len(pf.columns) > 1 else pf.name
            cand.append((label, mapping, korean_ratio(list(mapping.values()))))
    if len(cand) < 2:
        raise ValueError(
            "언어 컬럼이 2개 이상 필요합니다. 파일을 2개 올리거나, "
            "한국어와 영어가 함께 든 파일을 올려주세요."
        )

    cand.sort(key=lambda c: c[2], reverse=True)
    ko_label, ko_map, ko_r = cand[0]
    en_label, en_map, en_r = cand[-1]
    if ko_r < 0.2:
        raise ValueError("한국어로 보이는 컬럼이 없습니다.")
    if en_r > 0.5:
        raise ValueError("영어로 보이는 컬럼이 없습니다.")
    return LanguagePick(
        ko=ko_map, en=en_map, ko_label=ko_label, en_label=en_label,
        ratios=[(lbl, r) for lbl, _, r in cand],
    )


# ──────────────────────────────────────────────────────────────────────
# 분석
# ──────────────────────────────────────────────────────────────────────

def is_label(ko: str, en: str) -> bool:
    """
    UI 라벨인가, 아니면 문장(=패턴)인가.

    영어 쪽 종결부호만 보면 안 된다. 실제 카탈로그에는
    "array의 초기화 여부를 검사하지 않습니다." ↔ "Skip array initialization checking"
    처럼 **한국어만 문장인** 항목이 흔하다. 한국어 종결어미와 플레이스홀더도 본다.
    """
    if _PLACEHOLDER_RE.search(en) or _PLACEHOLDER_RE.search(ko):
        return False
    if re.search(r"[.!?]$", en.strip()) or re.search(r"[.!?]$", ko.strip()):
        return False
    if _SENT_STRONG_RE.search(ko) or _SENT_TAIL_RE.search(_ko_core(ko)):
        return False
    return len(ko) <= LABEL_MAX_KO_CHARS and len(en.split()) <= LABEL_MAX_EN_WORDS


# ── 용어성(termhood) 판정 기준 ────────────────────────────────────────
# 처음에는 '길면 용어'로 뽑았는데 정반대였다. 길수록 오히려 설명문이고,
# 진짜 제품 용어는 **짧은 복합어가 문서 곳곳에 반복**되는 쪽이다.
#   길이 정렬:  '공격 대상이 될 수 있는 보안 취약점 검출부터'  ← 용어 아님
#   빈도 정렬:  'SQL 삽입', '크로스 사이트 스크립팅', '경로 조작'  ← 용어
# 그렇다고 빈도만 쓰면 '수 / 및 / 정보' 같은 일반어가 올라오므로, 복합어
# 조건(2어절 이상)을 함께 건다.
TERM_MIN_WORDS = 2
TERM_MAX_WORDS = 4
TERM_MAX_CHARS = 18
TERM_MIN_FREQ = 3          # 카탈로그 전체에서 이만큼은 반복돼야 용어로 본다

_WORD_RE = re.compile(r"[가-힣A-Za-z0-9]+")


def _build_word_index(texts: List[str]) -> Dict[str, set]:
    """어절 → 등장 항목 인덱스. 후보마다 3만 건을 전수 검색하지 않기 위함."""
    inv: Dict[str, set] = defaultdict(set)
    for i, t in enumerate(texts):
        for w in _WORD_RE.findall(t):
            inv[w].add(i)
    return inv


def _corpus_freq(ko: str, inv: Dict[str, set], texts: List[str]) -> int:
    """카탈로그 전체에서 이 표현이 (부분 문자열로) 몇 번 나오는가."""
    words = _WORD_RE.findall(ko)
    if not words:
        return 0
    sets = [inv.get(w, set()) for w in words]
    if not all(sets):
        return 0
    hit = set.intersection(*sets)
    return sum(1 for i in hit if ko in texts[i])


def _looks_like_term(ko: str) -> bool:
    """모양만으로 거르는 1차 조건 — 빈도는 별도로 본다."""
    if not _KOREAN_RE.search(ko):
        return False
    if _JOSA_TAIL_RE.search(_ko_core(ko)):
        return False
    if len(ko) > TERM_MAX_CHARS:
        return False
    return TERM_MIN_WORDS <= len(_WORD_RE.findall(ko)) <= TERM_MAX_WORDS


# 뽑아낼 최대 개수. 후보를 수천 개 쏟아내면 검토가 불가능하고, 글로서리는
# 모든 문단에 전수 치환되므로 커질수록 오탐도 늘어난다. '많이'가 아니라
# '핵심만'이 목표다.
DEFAULT_TERM_LIMIT = 100
DEFAULT_PATTERN_LIMIT = 30

# 문형 대표를 고를 때 이 길이를 넘는 문장은 예시로 부적합
PATTERN_MAX_CHARS = 60

_SENT_SHAPE_TAIL = re.compile(
    r"(습니다|합니다|입니다|하세요|하십시오|십시오|됩니다|있습니다|없습니다|"
    r"했습니다|합니까|입니까|하시겠습니까)[.?!]?$"
)


def _pattern_shape(ko: str) -> str:
    """
    문형 시그니처. 같은 말투의 문장을 한 덩어리로 묶기 위한 키다.

    패턴은 '이런 한국어 문형은 이렇게 영어로 쓴다'는 예시라서, 같은 어미의
    문장을 수백 개 넣어봐야 프롬프트만 길어지고 배우는 건 똑같다. 끝 어절
    두 개를 시그니처로 묶고 그룹마다 대표 하나만 남긴다.
    """
    words = _WORD_RE.findall(ko)
    if not words:
        return ko[-6:]
    tail = " ".join(words[-2:]) if len(words) >= 2 else words[-1]
    m = _SENT_SHAPE_TAIL.search(ko.strip())
    return f"{m.group(1)}|{tail}" if m else tail


def _pick_patterns(rows: List[dict], limit: int) -> pd.DataFrame:
    """문형별로 묶어 대표 문장만 남기고, 흔한 문형 순으로 자른다."""
    groups: Dict[str, List[dict]] = defaultdict(list)
    for r in rows:
        if len(r["KO"]) > PATTERN_MAX_CHARS:
            continue
        groups[_pattern_shape(r["KO"])].append(r)
    if not groups:
        return pd.DataFrame(rows[:limit])

    ranked = sorted(groups.items(), key=lambda kv: len(kv[1]), reverse=True)
    picked = []
    for shape, members in ranked[:limit]:
        # 그룹 대표 = 가장 짧은 문장 (예시로 읽기 쉬운 쪽)
        rep = min(members, key=lambda r: len(r["KO"]))
        rep = dict(rep)
        rep["문형 빈도"] = len(members)
        picked.append(rep)
    return pd.DataFrame(picked)


@dataclass
class ExtractResult:
    labels: pd.DataFrame
    terms: pd.DataFrame
    patterns: pd.DataFrame
    stats: Dict[str, int]


def analyze(pick: LanguagePick,
            existing_terms: Optional[Dict[str, set]] = None,
            term_limit: int = DEFAULT_TERM_LIMIT,
            pattern_limit: int = DEFAULT_PATTERN_LIMIT) -> ExtractResult:
    """
    KO/EN 맵을 라벨·용어·패턴 후보로 가른다.

    existing_terms: {ko: {en_lower, ...}} — 기존 글로서리. 신규/충돌/동일 판정용.
    """
    existing_terms = existing_terms or {}
    keys = set(pick.ko) & set(pick.en)

    agg: Dict[str, Counter] = defaultdict(Counter)      # ko -> EN 후보 빈도
    ns: Dict[str, set] = defaultdict(set)               # ko -> key 네임스페이스
    sentences: List[Tuple[str, str, str]] = []

    for k in keys:
        ko, en = pick.ko[k].strip(), pick.en[k].strip()
        if not ko or not en:
            continue
        if is_label(ko, en):
            agg[ko][en] += 1
            ns[ko].add(k.split(".")[0])
        else:
            sentences.append((ko, en, k))

    label_rows = []
    for ko, cands in agg.items():
        rep = cands.most_common(1)[0][0]        # 최빈 표기를 대표로
        distinct = {c.lower() for c in cands}   # 대소문자만 다른 건 같은 표기로 본다
        prevs = existing_terms.get(ko) or set()
        status, prev_en = "신규", ""
        if prevs:
            prev_en = " / ".join(sorted(prevs))
            status = ("동일" if rep.lower() in {str(p).lower() for p in prevs}
                      else "충돌(기존)")
        label_rows.append({
            "KO": ko,
            "EN": rep,
            "기존 EN": prev_en,
            "후보수": len(distinct),
            "EN 후보": " / ".join(sorted(cands)) if len(distinct) > 1 else "",
            "문맥(key)": " ".join(sorted(ns[ko])[:3]),
            "출현": int(sum(cands.values())),
            "기존대조": status,
            "DNT": ko == rep and not bool(_KOREAN_RE.search(ko)),
        })

    labels = pd.DataFrame(label_rows)
    if not labels.empty:
        labels = labels.sort_values(
            ["후보수", "출현"], ascending=[False, False]
        ).reset_index(drop=True)

        # 용어 후보 — 1:1로 대응되고, 복합어이고, 문서 곳곳에 반복되는 것.
        shortlist = labels[
            (labels["후보수"] == 1) & labels["KO"].map(_looks_like_term)
        ].copy()
        if not shortlist.empty:
            corpus = [v.strip() for v in pick.ko.values() if v.strip()]
            inv = _build_word_index(corpus)
            shortlist["빈도"] = shortlist["KO"].map(
                lambda k: _corpus_freq(k, inv, corpus)
            )
            terms = shortlist[shortlist["빈도"] >= TERM_MIN_FREQ].copy()
            terms = terms.sort_values("빈도", ascending=False).reset_index(drop=True)
            term_pool = len(terms)
            terms = terms.head(term_limit).reset_index(drop=True)
        else:
            terms, term_pool = shortlist, 0
    else:
        terms, term_pool = labels.copy(), 0

    # 문장형은 같은 KO가 반복되면 한 번만
    seen_sent: set = set()
    pat_rows = []
    for ko, en, k in sentences:
        if ko in seen_sent:
            continue
        seen_sent.add(ko)
        pat_rows.append({"KO": ko, "EN": en, "문맥(key)": k.split(".")[0]})
    pattern_pool = len(pat_rows)
    patterns = _pick_patterns(pat_rows, pattern_limit)

    stats = {
        "공통키": len(keys),
        "라벨 고유KO": len(agg),
        "충돌": int((labels["후보수"] > 1).sum()) if not labels.empty else 0,
        "용어후보": len(terms),
        "용어풀": term_pool,          # 상한을 걸기 전 개수
        "패턴후보": len(patterns),
        "패턴풀": pattern_pool,
        "DNT후보": int(labels["DNT"].sum()) if not labels.empty else 0,
        "신규": int((labels["기존대조"] == "신규").sum()) if not labels.empty else 0,
        "기존충돌": int((labels["기존대조"] == "충돌(기존)").sum()) if not labels.empty else 0,
    }
    return ExtractResult(labels=labels, terms=terms, patterns=patterns, stats=stats)


def existing_terms_index(terms_df: pd.DataFrame) -> Dict[str, set]:
    """
    load_terms() 결과를 {ko: {원본 EN, …}} 인덱스로.

    소문자로 뭉개지 않고 원본 표기를 그대로 담는다 — 충돌을 알릴 때
    "기존에는 무엇으로 등록돼 있는지"를 사람에게 보여줘야 하기 때문.
    비교는 읽는 쪽에서 대소문자 무시로 한다.
    """
    idx: Dict[str, set] = defaultdict(set)
    for _, r in terms_df.iterrows():
        ko = str(r.get("KO", "")).strip()
        en = str(r.get("EN", "")).strip()
        if ko and en:
            idx[ko].add(en)
    return dict(idx)


# ──────────────────────────────────────────────────────────────────────
# 표기 충돌 정리 — 여기만 LLM을 쓴다
#
# 같은 한국어에 영어가 여러 개인 항목은 두 종류가 섞여 있다.
#   ① 표기 불일치  '검출 규칙' → Detecting Rule / Detecting rules / Detection rule
#                  단복수·대소문자·활용형 차이일 뿐이라 하나로 통일하면 된다.
#   ② 문맥 분기    '확인' → OK / Check / Confirm
#                  버튼이냐 규칙이냐에 따라 진짜로 다른 단어다.
# ①을 사람이 하나씩 고르는 건 낭비고, ②를 기계가 정하는 건 위험하다.
# 그래서 LLM은 **둘을 가려내고 ①의 대표형만 제안**하며, ②는 사람에게 넘긴다.
# ──────────────────────────────────────────────────────────────────────

_RESOLVE_RE = re.compile(r"^\[(\d+)\]\s*(UNIFY|SPLIT)\s*\|([^|]*)\|(.*)$", re.MULTILINE)


def conflict_rows(labels: pd.DataFrame) -> pd.DataFrame:
    """표기가 갈리는 항목만."""
    if labels.empty:
        return labels
    return labels[labels["후보수"] > 1].copy()


def resolve_conflicts(client, rows: pd.DataFrame, model: str = "gpt-5.2",
                      batch_size: int = 25) -> Dict[str, dict]:
    """
    충돌 항목을 LLM이 '표기 통일'과 '문맥 분기'로 가른다.

    Returns {ko: {"kind": "UNIFY"|"SPLIT", "pick": str, "reason": str}}
    실패한 배치는 결과에서 빠진다 — 없는 항목은 사람이 직접 고르면 된다.
    """
    out: Dict[str, dict] = {}
    # itertuples는 'EN 후보'처럼 공백/한글이 든 컬럼명을 망가뜨리므로 dict로 뽑는다
    items: List[dict] = rows.to_dict("records")
    for start in range(0, len(items), batch_size):
        batch = items[start:start + batch_size]
        block = "\n".join(
            f"[{i}] KO: {it['KO']}\n"
            f"    candidates: {it.get('EN 후보') or it.get('EN', '')}\n"
            f"    used in key namespaces: {it.get('문맥(key)', '')}"
            for i, it in enumerate(batch)
        )
        prompt = f"""You are normalising a software UI string catalogue.

Each item is one Korean UI label that appears with SEVERAL different English
spellings across the product. Decide which of two situations it is:

UNIFY — the candidates are the SAME term written inconsistently
        (capitalisation, singular/plural, verb form, word order).
        Choose the single best canonical English form.
        Prefer: Title Case for UI labels, singular unless the label really
        denotes a list, and the noun form over the -ing form.

SPLIT  — the candidates genuinely mean DIFFERENT things depending on which
        screen the label is on (e.g. a confirm button vs a check rule).
        A human must decide per screen. Still name the single most likely
        candidate as a fallback, but mark the item SPLIT so a person reviews it.

Output EXACTLY one line per item, nothing else. ALWAYS fill the middle field:
[N] UNIFY|<chosen English>|<short reason in Korean>
[N] SPLIT|<most likely English>|<short reason in Korean>

Items:
{block}
""".strip()

        try:
            resp = client.responses.create(
                model=model,
                input=prompt,
                reasoning={"effort": "low"},
                text={"verbosity": "low"},
            )
            text = resp.output_text
        except Exception:
            continue

        for m in _RESOLVE_RE.finditer(text):
            idx, kind, pick, reason = m.groups()
            i = int(idx)
            if i >= len(batch):
                continue
            out[batch[i]["KO"]] = {
                "kind": kind,
                "pick": pick.strip(),
                "reason": reason.strip(),
            }
    return out


# ──────────────────────────────────────────────────────────────────────
# 등재 전 검수
#
# 카탈로그의 영어는 사람이 오랜 기간 나눠 쓴 것이라 표기가 고르지 않다.
# Title Case와 sentence case가 섞이고, 오타("Managmer")나 단복수 오류도
# 그대로 들어 있다. 그걸 검증 없이 글로서리에 넣으면 오류가 번역 결과에
# 그대로 전파된다. 그래서 **사용자가 고른 것만** 검수해서 고칠 것을
# 제안하고, 채택 여부는 사람이 정한다.
# ──────────────────────────────────────────────────────────────────────

_REVIEW_RE = re.compile(r"^\[(\d+)\]\s*([^|]*)\|(.*)$", re.MULTILINE)

_REVIEW_RULES = {
    "term": """These are glossary entries: a Korean term and the English used for it.

Flag an entry ONLY when the English has a real problem:
- Misspelling or a non-word (e.g. "Managmer" -> "Manager")
- Wrong capitalisation for a glossary entry. Use lower case for ordinary nouns
  because the term is substituted mid-sentence. Keep Title Case ONLY for product
  names, feature names and acronyms (e.g. "Virtual Drive", "SQL Injection").
- Singular/plural or article error
- Obvious mistranslation of the Korean

Do NOT flag an entry merely because you would have phrased it differently.""",
    "pattern": """These are translation examples: a Korean sentence and its English.

Flag an entry ONLY when the English has a real problem:
- Misspelling, non-word, or broken grammar
- Missing or doubled sentence-final punctuation compared with the Korean
- Singular/plural or article error
- Obvious mistranslation of the Korean

Do NOT flag an entry merely because you would have phrased it differently.""",
}


def review_entries(client, items: List[Tuple[str, str]], kind: str = "term",
                   model: str = "gpt-5.2", batch_size: int = 25) -> Dict[int, dict]:
    """
    등재 직전 검수. items = [(ko, en), …]

    Returns {index: {"suggest": 고친 영어, "reason": 사유}}
    — 문제가 없다고 판단한 항목은 결과에 들어가지 않는다.
    """
    out: Dict[int, dict] = {}
    for start in range(0, len(items), batch_size):
        batch = items[start:start + batch_size]
        block = "\n".join(
            f"[{i}] KO: {ko}\n    EN: {en}" for i, (ko, en) in enumerate(batch)
        )
        prompt = f"""{_REVIEW_RULES.get(kind, _REVIEW_RULES["term"])}

Output one line per entry that needs a change, nothing else:
[N] <corrected English>|<short reason in Korean>

If an entry is fine, omit it entirely. If nothing needs changing, output NONE.

Entries:
{block}
""".strip()

        try:
            resp = client.responses.create(
                model=model,
                input=prompt,
                reasoning={"effort": "low"},
                text={"verbosity": "low"},
            )
            text = resp.output_text
        except Exception:
            continue

        if text.strip().upper().startswith("NONE"):
            continue
        for m in _REVIEW_RE.finditer(text):
            idx, suggest, reason = m.groups()
            i = int(idx)
            suggest = suggest.strip()
            if i >= len(batch) or not suggest:
                continue
            if suggest == batch[i][1]:
                continue          # 제안이 원본과 같으면 무시
            out[start + i] = {"suggest": suggest, "reason": reason.strip()}
    return out


# ──────────────────────────────────────────────────────────────────────
# 내보내기
# ──────────────────────────────────────────────────────────────────────

def to_excel(terms: pd.DataFrame, patterns: pd.DataFrame, product: str = "ALL") -> bytes:
    """
    기존 마스터 엑셀과 **같은 시트/컬럼 구조**로 내보낸다.

    그래야 [Glossary 관리]의 기존 업로드 경로로 그대로 되먹일 수 있다.
    """
    empty = pd.Series(dtype=str)
    t = pd.DataFrame({
        "KO": terms.get("KO", empty),
        "EN": terms.get("EN", empty),
        "Product": product,
        "DNT": terms.get("DNT", pd.Series(dtype=bool)),
        "Case-sensitive": False,
        "Note": terms.get("문맥(key)", empty),
    })
    p = pd.DataFrame({
        "KO": patterns.get("KO", empty),
        "EN": patterns.get("EN", empty),
        "Note": patterns.get("문맥(key)", empty),
    })
    buf = io.BytesIO()
    with pd.ExcelWriter(buf, engine="openpyxl") as writer:
        t.to_excel(writer, sheet_name="glossary", index=False)
        p.to_excel(writer, sheet_name="pattern", index=False)
    return buf.getvalue()
